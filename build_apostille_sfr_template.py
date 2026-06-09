r"""
Pack 50.20-C — Генератор apostille_sfr_template.docx из эталона Минфина/СФР.

Берёт готовый эталон (Апостиль_минфина_СФР.docx), вставляет docxtpl-плейсхолдеры
в динамические ячейки, сохраняет как templates/docx/apostille_sfr_template.docx.

Динамические поля (по аналогии с apostille_template самозанятого):
  - R3 подписант СФР    → {{ apostille.signer_sfr_short }}
  - R7C4 дата           → {{ apostille.date_short }}
  - R9 номер            → {{ apostille.number }}
  - R8 подписант апостиля (Байрамов) → {{ apostille.signer_apostille_block }}

Остальное (отделение СФР, метки) — фиксировано в эталоне.

ВХОД: эталон рядом со скриптом (Апостиль_минфина_СФР.docx) ИЛИ путь аргументом.
ВЫХОД: templates/docx/apostille_sfr_template.docx

Запуск:
    python build_apostille_sfr_template.py [путь_к_эталону.docx]
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
if (SCRIPT_DIR / "templates" / "docx").exists():
    OUTPUT_DIR = SCRIPT_DIR / "templates" / "docx"
elif Path("D:/VISA/visa_kit/templates/docx").exists():
    OUTPUT_DIR = Path("D:/VISA/visa_kit/templates/docx")
else:
    OUTPUT_DIR = SCRIPT_DIR / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "apostille_sfr_template.docx"


def _find_etalon() -> Path:
    if len(sys.argv) > 1:
        p = Path(sys.argv[1])
        if p.exists():
            return p
    for name in ("Апостиль_минфина_СФР.docx", "apostille_sfr_etalon.docx"):
        cand = SCRIPT_DIR / name
        if cand.exists():
            return cand
    raise FileNotFoundError(
        "Эталон не найден. Положи 'Апостиль_минфина_СФР.docx' рядом со скриптом "
        "или укажи путь аргументом: python build_apostille_sfr_template.py <path>"
    )


def _set_cell_placeholder(cell, placeholder: str):
    """Заменяет текст ячейки на placeholder, сохраняя форматирование первого run.

    Чистит все параграфы кроме первого, в первом оставляет один run с плейсхолдером.
    """
    # очистим все параграфы, кроме первого
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    # сохраняем формат первого run, остальные удаляем
    runs = p.runs
    if runs:
        first = runs[0]
        first.text = placeholder
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(placeholder)


def main() -> int:
    etalon = _find_etalon()
    print(f"Эталон: {etalon}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_20_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")

    doc = Document(str(etalon))
    t = doc.tables[0]  # таблица APOSTILLE 11x5

    # R3 — подписант СФР (был "Высоцкая Ю.В."), ячейка значения = C2 (merge C2:C4)
    _set_cell_placeholder(t.rows[3].cells[2], "{{ apostille.signer_sfr_short }}")

    # R7C4 — дата (была "07.04.2026")
    _set_cell_placeholder(t.rows[7].cells[4], "{{ apostille.date_short }}")

    # R9C1 — номер (был "77-02759/26"), значение в C1 (merge C1:C2)
    _set_cell_placeholder(t.rows[9].cells[1], "{{ apostille.number }}")

    # R8C1 — подписант апостиля (был "Байрамов Н.А. \n Заместитель..."),
    # значение merge C1:C4. Делаем плейсхолдер с переносом строки внутри.
    # Используем jinja-переменную, которая придёт уже с \n -> через шаблон
    # вставим две строки: short + position. Для простоты — один плейсхолдер,
    # контекст соберёт многострочный текст.
    _set_cell_placeholder(t.rows[8].cells[1], "{{ apostille.signer_apostille_block }}")

    doc.save(str(OUTPUT_PATH))
    print(f"✅ Saved: {OUTPUT_PATH}  ({OUTPUT_PATH.stat().st_size/1024:.1f} KB)")

    # Проверка: вывести плейсхолдеры что реально попали
    import re
    check = Document(str(OUTPUT_PATH))
    found = set()
    for tt in check.tables:
        for row in tt.rows:
            for c in row.cells:
                for ph in re.findall(r"\{\{\s*(apostille\.\w+)\s*\}\}", c.text):
                    found.add(ph)
    print("Плейсхолдеры в шаблоне:", sorted(found))
    return 0


if __name__ == "__main__":
    sys.exit(main())
