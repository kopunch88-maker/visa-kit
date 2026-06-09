r"""
Pack 50.9-C-r6 — Генератор шаблона stdr_template.docx по эталону СФР.

Изменения r6 — точные размеры шрифтов из эталона:
  - Заголовок документа: 14pt (Title)
  - "Сведения о зарегистрированном лице:" + ФИО/ДР/СНИЛС: 9.5pt bold
  - "Подано заявление о...": 11pt (default Normal)
  - "Дата подачи": 5.5pt italic
  - ЭЦП блок: 7.5pt bold (стиль "a3")
  - ЭЦП блок увеличен по высоте (3.2 см) и размещён НИЖЕ на странице
  - Отступы ФИО (tab) увеличены до 6 см
  - "Документ подписан" — одна строка, "усиленной..." — две

Точное соответствие структуре эталона ЭТК_Орлов.docx:
  - 11 колонок (а не 12 как в v1)
  - 3-строчная шапка с vMerge + gridSpan
  - Колонка "Основание" с gridSpan=3 над (Код функции / Причины увольнения)
  - "Наименование документа" / "Дата" / "Номер документа" — самостоятельные колонки
    под общим "Наименование" (gridSpan=3) — НО на самом деле это под группой 5-7,
    а 8-10 — это отдельная группа "Документ-наименование/Дата/Номер" под именованием
    "Наименование" (тоже gridSpan=3).
    
Стоп — пересмотрел дамп эталона:
  Row 1: | Дата | Сведения | [gridSpan=3: "Наименование"]    | [gridSpan=3: "Основание"]      |
  Row 2: | merge| merge    | Труд.функция | Код функции | Причины | Наим.док | Дата | Номер  |
  
Значит "Наименование" — это группа из (Труд.функция / Код функции / Причины увольнения).
А под "Основание" идут (Наим.документа / Дата / Номер).

Ширины колонок из эталона (мм → cm):
  1: №       0.68
  2: Работ.  3.13
  3: Дата    2.45
  4: Свед.   2.45
  5: Труд.ф. 4.90
  6: Код     1.91
  7: Причины 4.09
  8: Док.    1.91
  9: Дата    1.64
  10: Номер  1.36
  11: Призн. 2.73
  Итого: 26.25 cm

Запуск:
    python build_stdr_template.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


# ============================================================================
# Constants
# ============================================================================

OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "stdr_template.docx"

# Альбомная А4: 29.7 × 21.0 cm
PAGE_WIDTH_CM = 29.7
PAGE_HEIGHT_CM = 21.0
MARGIN_LR_CM = 1.5
MARGIN_TB_CM = 1.0
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_LR_CM  # = 26.7

FONT_FAMILY = "Arial"
FONT_SIZE_PT = 7  # данные таблиц
FONT_SIZE_HEADER_PT = 7  # шапка таблиц
FONT_SIZE_LABEL_PT = 9.5  # Pack 50.9-r6: точно из эталона — ФИО/СНИЛС/ДР 9.5pt bold
FONT_SIZE_STATEMENT_PT = 11  # Pack 50.9-r6: "Подано заявление о..." 11pt
FONT_SIZE_DATE_LABEL_PT = 5.5  # Pack 50.9-r6: "Дата подачи" 5.5pt italic
FONT_SIZE_EP_PT = 7.5  # Pack 50.9-r6: текст ЭЦП 7.5pt bold (стиль a3)
FONT_SIZE_TITLE_PT = 12  # главный заголовок документа
FONT_SIZE_SECTION_PT = 10  # подзаголовки секций

STDR_TABLE1_SLOTS = 15
STDR_TABLE2_SLOTS = 8


# Ширины 11 колонок Таблицы 1 (по эталону), сумма ≈ 26.25 cm
# Растяну на 26.7 (наша usable ширина) пропорционально
T1_COL_WIDTHS_RAW = [0.68, 3.13, 2.45, 2.45, 4.90, 1.91, 4.09, 1.91, 1.64, 1.36, 2.73]
T1_TOTAL_RAW = sum(T1_COL_WIDTHS_RAW)  # 27.25... подожди, посчитаю
# Реально: 0.68+3.13+2.45+2.45+4.90+1.91+4.09+1.91+1.64+1.36+2.73 = 27.25
# Это больше usable 26.7 на 2%. Пропорционально уменьшу
T1_COL_WIDTHS = [w * USABLE_WIDTH_CM / T1_TOTAL_RAW for w in T1_COL_WIDTHS_RAW]


# Ширины 4 колонок Таблицы 2 — №, Работодатель, с, по
T2_COL_WIDTHS = [1.0, 18.0, 3.85, 3.85]  # = 26.7


# ============================================================================
# Helpers
# ============================================================================

def _set_run(run, *, size=FONT_SIZE_PT, bold=False, italic=False):
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Для русских символов
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    rFonts.set(qn("w:cs"), FONT_FAMILY)
    rFonts.set(qn("w:eastAsia"), FONT_FAMILY)


def _add_paragraph(doc, text, *, size=FONT_SIZE_PT, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before_pt=0, space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if text:
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, italic=italic)
    return p


def _set_cell_borders(cell, *, top=True, bottom=True, left=True, right=True,
                      size_val="4"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    sides = [("top", top), ("left", left), ("bottom", bottom), ("right", right)]
    for side_name, enabled in sides:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        border = OxmlElement(f"w:{side_name}")
        if enabled:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size_val)
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "nil")
        tcBorders.append(border)


def _set_table_borders_tbl(table):
    """Pack 50.9-r3: устанавливаем границы через tblBorders (как в эталоне СФР).

    Это правильный способ для таблиц с merged cells — Word корректно
    рендерит границы вокруг объединённых ячеек.
    Альтернатива (tcBorders на каждой ячейке) ломается на merged.
    """
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Удалим существующий tblBorders если есть
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_borders_all(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


def _cell_text(cell, text, *, size=FONT_SIZE_PT, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center"):
    cell.text = ""
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[vertical]
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold, italic=italic)


def _set_table_fixed_layout(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _apply_widths(table, widths_cm: list):
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _merge_horizontal(row, start_col: int, end_col: int):
    """Объединяет ячейки в строке от start_col до end_col включительно."""
    cells = row.cells
    merged = cells[start_col]
    for ci in range(start_col + 1, end_col + 1):
        merged = merged.merge(cells[ci])
    return merged


def _set_vmerge(cell, val):
    """Устанавливает w:vMerge на ячейку. val='restart' или None (continue)."""
    tcPr = cell._element.get_or_add_tcPr()
    vMerge = tcPr.find(qn("w:vMerge"))
    if vMerge is None:
        vMerge = OxmlElement("w:vMerge")
        tcPr.append(vMerge)
    if val == "restart":
        vMerge.set(qn("w:val"), "restart")
    elif val is None:
        # continue — без val атрибута
        if qn("w:val") in vMerge.attrib:
            del vMerge.attrib[qn("w:val")]


def _set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)


# ============================================================================
# Builder — section 1: applicant info
# ============================================================================

def _build_applicant_section(doc):
    """Шапка + Сведения о зарегистрированном лице."""
    # Главный заголовок
    _add_paragraph(
        doc,
        "Сведения о трудовой деятельности, предоставляемые из "
        "информационных ресурсов Фонда пенсионного и социального "
        "страхования Российской Федерации",
        size=FONT_SIZE_TITLE_PT,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    _add_paragraph(
        doc, "Сведения о зарегистрированном лице:",
        size=FONT_SIZE_LABEL_PT, bold=True,
        space_before_pt=4, space_after_pt=2,
    )

    # Inline: Фамилия | Имя строки. В эталоне это inline (без таблицы).
    # Я сделаю простой 2-колоночный список через таблицу без границ.
    t = doc.add_table(rows=5, cols=2)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, [4.0, USABLE_WIDTH_CM - 4.0])

    rows_data = [
        ("Фамилия", "{{ applicant_stdr.last_name_upper }}"),
        ("Имя", "{{ applicant_stdr.first_name_upper }}"),
        ("Отчество", "{{ applicant_stdr.middle_name_upper }}"),
        ("Дата Рождения", "{{ applicant_stdr.birth_date_long }}"),
        ("СНИЛС", "{{ applicant_stdr.snils }}"),
    ]
    # Pack 50.9-r6: ФИО как параграфы — точно как эталон.
    # 9.5pt bold, таб на ~6 см.
    from docx.enum.text import WD_TAB_ALIGNMENT as _WD_TAB_ALIGN

    LABEL_TAB_CM = 6.0  # Pack 50.9-r6: больше чем в r5 (3.5cm)
    VALUE_END_CM = USABLE_WIDTH_CM  # Underline тянется до правого края

    for label, value in rows_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Tab stop в начале значения (после label) и в конце строки (для underline до края)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(LABEL_TAB_CM), _WD_TAB_ALIGN.LEFT)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(VALUE_END_CM), _WD_TAB_ALIGN.RIGHT)

        # Метка слева — 9.5pt НЕ bold (label не жирный)
        r_label = p.add_run(label)
        _set_run(r_label, size=FONT_SIZE_LABEL_PT, bold=False)
        # Таб к значению
        r_tab = p.add_run("\t")
        _set_run(r_tab, size=FONT_SIZE_LABEL_PT)
        # Значение с underline — 9.5pt BOLD
        r_value = p.add_run(value)
        _set_run(r_value, size=FONT_SIZE_LABEL_PT, bold=True)
        rPr = r_value._element.get_or_add_rPr()
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        # Таб до правого края — с underline (продолжение линии)
        r_pad = p.add_run("\t")
        _set_run(r_pad, size=FONT_SIZE_LABEL_PT)
        rPr_pad = r_pad._element.get_or_add_rPr()
        u_pad = OxmlElement("w:u")
        u_pad.set(qn("w:val"), "single")
        rPr_pad.append(u_pad)

    # Удалю созданную ранее таблицу t — она пустая теперь
    t._element.getparent().remove(t._element)

    # Pack 50.9-r4: Строки "Подано заявление о..." — как параграфы с
    # табуляцией справа + underline на дате (как в эталоне СФР).
    # Эталон: параграф "Подано заявление о ... \t\t 08.05.2020" с underline на дате,
    # потом параграф "Дата подачи" мелким курсивом справа.
    from docx.shared import Mm

    def _add_statement_paragraph(label_text, placeholder):
        """Добавляет одну строку заявления с подчёркнутой датой справа.
        
        Pack 50.9-r6: размер 11pt (как Normal в эталоне).
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        # Tab stop справа на ширине usable
        from docx.enum.text import WD_TAB_ALIGNMENT
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(USABLE_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT
        )
        # Текст слева — 11pt
        r_label = p.add_run(label_text)
        _set_run(r_label, size=FONT_SIZE_STATEMENT_PT)
        # Таб
        r_tab = p.add_run("\t")
        _set_run(r_tab, size=FONT_SIZE_STATEMENT_PT)
        # Плейсхолдер даты — с underline, 11pt bold
        r_date = p.add_run(placeholder)
        _set_run(r_date, size=FONT_SIZE_STATEMENT_PT, bold=True)
        # underline
        rPr = r_date._element.get_or_add_rPr()
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    def _add_date_label_paragraph():
        """Параграф 'Дата подачи' справа, мелким курсивом 5.5pt."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Дата подачи")
        _set_run(r, size=FONT_SIZE_DATE_LABEL_PT, italic=True)

    _add_statement_paragraph(
        "Подано заявление о продолжении ведения трудовой книжки",
        "{{ stdr.statement_continue_date }}",
    )
    _add_date_label_paragraph()

    _add_statement_paragraph(
        "Подано заявление о предоставлении сведений о трудовой деятельности",
        "{{ stdr.statement_provide_date }}",
    )
    _add_date_label_paragraph()


# ============================================================================
# Builder — section 2: Table 1 (события с 2020+)
# ============================================================================

def _build_table1(doc):
    """Главная таблица Сведений о труд. деятельности — 11 колонок.
    
    Pack 50.9-r4: заголовок 'Сведения о трудовой деятельности зарегистрированного лица'
    убран — он уже есть В шапке таблицы как gridSpan=8 в row 0.
    """
    # Маленький отступ перед таблицей
    _add_paragraph(doc, "", space_before_pt=4)

    # 4 строки шапки + 15 слотов данных
    HEADER_ROWS = 4  # row 0,1,2: текстовая шапка; row 3: номера колонок 1..11
    t = doc.add_table(rows=HEADER_ROWS + STDR_TABLE1_SLOTS, cols=11)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, T1_COL_WIDTHS)

    # === ШАПКА ===
    # Row 0: № | Работодатель | [GS=8: Сведения о труд.деят.] | Признак отмены
    _cell_text(t.rows[0].cells[0], "№ п/п",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[0].cells[1],
               "Работодатель (наименование),\nрегистрационный номер в СФР",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Колонки 2..9 объединить и поставить "Сведения о труд.деятельности"
    sved_cell = _merge_horizontal(t.rows[0], 2, 9)
    _cell_text(sved_cell,
               "Сведения о трудовой деятельности зарегистрированного лица",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[0].cells[10],  # cell 10 после merge — это бывшая cell 10 (Признак)
               "Признак отмены записи сведений о приеме, переводе, увольнения",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1: пусто (vmerge) | пусто (vmerge) | Дата | Сведения | [GS=3: Наименование] | [GS=3: Основание] | пусто (vmerge)
    _cell_text(t.rows[1].cells[2],
               "Дата (число, месяц, год) приема,\nперевода, увольнения",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[1].cells[3],
               "Сведения о приеме, переводе,\nувольнении",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Колонки 4,5,6 — "Наименование"
    naim_cell = _merge_horizontal(t.rows[1], 4, 6)
    _cell_text(naim_cell, "Наименование",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Колонки 7,8,9 — "Основание"
    osn_cell = _merge_horizontal(t.rows[1], 7, 9)
    _cell_text(osn_cell, "Основание",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: пусто | пусто | пусто (vmerge с row 1) | пусто (vmerge с row 1) | Труд.функция | Код | Причины | Наим.док | Дата | Номер | пусто (vmerge)
    _cell_text(t.rows[2].cells[4],
               "Трудовая функция (должность, профессия, специальность, "
               "квалификация, конкретный вид поручаемой работы), структурное\nподразделение",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[2].cells[5],
               "Код\nвыполняемой функции (при наличии)",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[2].cells[6],
               "Причины увольнения, пункт, часть статьи, статья "
               "Трудового кодекса Российской Федерации,\nфедерального закона",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[2].cells[7], "Наименование документа",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[2].cells[8], "Дата",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[2].cells[9], "Номер документа",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # vMerge — устанавливаем для столбцов 0, 1, 10 через все 3 строки шапки.
    # Также для столбцов 2 и 3 — vmerge через row 1 и row 2 (под "Дата" и "Сведения").
    # Note: row 0 имеет restart, row 1 и row 2 — continue
    for col in [0, 1, 10]:
        _set_vmerge(t.rows[0].cells[col], "restart")
        _set_vmerge(t.rows[1].cells[col], None)
        _set_vmerge(t.rows[2].cells[col], None)
    # Колонки 2 и 3: restart в row 1, continue в row 2
    for col in [2, 3]:
        _set_vmerge(t.rows[1].cells[col], "restart")
        _set_vmerge(t.rows[2].cells[col], None)

    # Row 3: номера колонок 1..11
    for i in range(11):
        _cell_text(t.rows[3].cells[i], str(i + 1),
                   size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # === ДАННЫЕ (слоты с плейсхолдерами) ===
    for slot in range(STDR_TABLE1_SLOTS):
        row_idx = HEADER_ROWS + slot
        cells = t.rows[row_idx].cells
        # 11 полей:
        #   1: index            CENTER
        #   2: company_with_sfr LEFT
        #   3: event_date       CENTER
        #   4: event_type       CENTER
        #   5: position         LEFT
        #   6: okz_code         CENTER (Код выполняемой функции)
        #   7: dismissal_reason LEFT (только для УВОЛЬНЕНИЯ)
        #   8: doc_name         CENTER (Приказ)
        #   9: doc_date         CENTER
        #   10: doc_number      CENTER
        #   11: cancellation    CENTER (Признак отмены)
        slot_fields = [
            ("index",             WD_ALIGN_PARAGRAPH.CENTER),
            ("company_with_sfr",  WD_ALIGN_PARAGRAPH.LEFT),
            ("event_date",        WD_ALIGN_PARAGRAPH.CENTER),
            ("event_type",        WD_ALIGN_PARAGRAPH.CENTER),
            ("position",          WD_ALIGN_PARAGRAPH.LEFT),
            ("okz_code",          WD_ALIGN_PARAGRAPH.CENTER),
            ("dismissal_reason",  WD_ALIGN_PARAGRAPH.LEFT),
            ("doc_name",          WD_ALIGN_PARAGRAPH.CENTER),
            ("doc_date",          WD_ALIGN_PARAGRAPH.CENTER),
            ("doc_number",        WD_ALIGN_PARAGRAPH.CENTER),
            ("cancellation",      WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for col_idx, (field, align) in enumerate(slot_fields):
            placeholder = f"{{{{ stdr.table1_rows[{slot}].{field} }}}}"
            _cell_text(cells[col_idx], placeholder, align=align)

    # Pack 50.9-r3: используем tblBorders (на уровне таблицы) вместо tcBorders
    # на каждой ячейке — корректно работает с merged cells (как в эталоне СФР).
    _set_table_borders_tbl(t)


# ============================================================================
# Builder — section 3: Table 2 (периоды до 2019)
# ============================================================================

def _build_table2(doc):
    """Сведения о труд.деят. за периоды до 31.12.2019 — 4 колонки."""
    _add_paragraph(
        doc,
        "Сведения о трудовой деятельности зарегистрированного лица "
        "за периоды до 31 декабря 2019 года",
        size=FONT_SIZE_SECTION_PT, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=12, space_after_pt=4,
    )

    HEADER_ROWS = 2
    t = doc.add_table(rows=HEADER_ROWS + STDR_TABLE2_SLOTS, cols=4)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, T2_COL_WIDTHS)

    # Row 0: № | Работодатель | [GS=2: Периоды работы]
    _cell_text(t.rows[0].cells[0], "№ п/п",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[0].cells[1],
               "Работодатель (наименование),\nрегистрационный номер в СФР (при наличии)",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    periods_cell = _merge_horizontal(t.rows[0], 2, 3)
    _cell_text(periods_cell, "Периоды работы",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1: пусто (vmerge) | пусто (vmerge) | с дд.мм.гггг | по дд.мм.гггг
    _cell_text(t.rows[1].cells[2], "с дд.мм.гггг",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[1].cells[3], "по дд.мм.гггг",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # vMerge для cols 0,1 через 2 строки шапки
    for col in [0, 1]:
        _set_vmerge(t.rows[0].cells[col], "restart")
        _set_vmerge(t.rows[1].cells[col], None)

    # Слоты данных
    for slot in range(STDR_TABLE2_SLOTS):
        row_idx = HEADER_ROWS + slot
        cells = t.rows[row_idx].cells
        slot_fields = [
            ("index",            WD_ALIGN_PARAGRAPH.CENTER),
            ("company_with_sfr", WD_ALIGN_PARAGRAPH.LEFT),
            ("date_from",        WD_ALIGN_PARAGRAPH.CENTER),
            ("date_to",          WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for col_idx, (field, align) in enumerate(slot_fields):
            placeholder = f"{{{{ stdr.table2_rows[{slot}].{field} }}}}"
            _cell_text(cells[col_idx], placeholder, align=align)

    _set_table_borders_tbl(t)


# ============================================================================
# Builder — section 4: signature + ЭЦП
# ============================================================================

def _build_signature(doc):
    """Подпись + дата + М.П. + блок ЭЦП."""
    _add_paragraph(doc, "", space_before_pt=12)

    # 3 колонки: должность | подпись | расшифровка
    t = doc.add_table(rows=2, cols=3)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, [14.0, 6.35, 6.35])

    _cell_text(t.rows[0].cells[0],
               "Должность уполномоченного лица\nтерриториального органа СФР",
               size=FONT_SIZE_LABEL_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t.rows[0].cells[1], "Подпись",
               size=FONT_SIZE_LABEL_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[0].cells[2], "Расшифровка подписи",
               size=FONT_SIZE_LABEL_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_text(t.rows[1].cells[0], "{{ stdr.issue_date_long }}",
               size=FONT_SIZE_LABEL_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t.rows[1].cells[1], "(дата)",
               size=FONT_SIZE_HEADER_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.rows[1].cells[2], "М.П. (при наличии)",
               size=FONT_SIZE_LABEL_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Без границ
    for row in t.rows:
        for cell in row.cells:
            _set_cell_borders(cell, top=False, bottom=False, left=False, right=False)

    # Pack 50.9-r5: Блок ЭЦП — Drawing TextBox с СИНЕЙ рамкой и СИНИМ текстом
    # (цвет 1C6EB5 как в эталоне Орлова), расположенный справа на странице.
    # Якорь: behindDoc=0, относительно параграфа, обтекание wrapNone.
    # Размер: 7.83 × 2.52 см. Позиция справа.
    _add_ep_signature_drawing(doc)


def _add_ep_signature_drawing(doc):
    """Pack 50.9-r6: Drawing TextBox с электронной подписью (синяя рамка + текст).
    
    Точные параметры эталона Орлова:
      - размер: 7.83 × 2.52 см → R6: увеличен до 8.5 × 3.4 см чтобы текст помещался
      - шрифт: 7.5pt bold (стиль "a3" из эталона)
      - цвет: 1C6EB5
      - расположение: справа на странице, отступ снизу
    """
    # Создаю пустой параграф-якорь
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()

    # Размеры в EMU (1 cm ≈ 360000 EMU)
    # Pack 50.9-r6: УВЕЛИЧЕНО для размещения всего текста
    SHAPE_W_EMU = 3060000  # 8.5 см (было 7.83)
    SHAPE_H_EMU = 1220000  # 3.4 см  (было 2.52)
    # Позиция справа
    POS_H_EMU = 10692000 - SHAPE_W_EMU - 360000
    POS_V_EMU = 100000

    color_hex = "1C6EB5"
    # Pack 50.9-r6: шрифт 15 half-points = 7.5pt (как стиль a3 в эталоне)
    sz_val = "15"

    drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="500" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{POS_H_EMU}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="paragraph">
      <wp:posOffset>{POS_V_EMU}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="100" name="ЭЦП_печать"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks noChangeAspect="0"/>
    </wp:cNvGraphicFramePr>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr txBox="1"/>
          <wps:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:noFill/>
            <a:ln w="9525">
              <a:solidFill>
                <a:srgbClr val="{color_hex}"/>
              </a:solidFill>
              <a:prstDash val="solid"/>
            </a:ln>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="20" w:after="0" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="center"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>Документ подписан</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="0" w:after="40" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="center"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>усиленной квалифицированной электронной подписью.</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="40" w:after="0" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="left"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>Организация: ФОНД ПЕНСИОННОГО И СОЦИАЛЬНОГО СТРАХОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="0" w:after="0" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="left"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>Сертификат: 122b085f13c95875c1a786a93d9b4e6f</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="0" w:after="0" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="left"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>Издатель: Федеральное казначейство</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:pPr>
                  <w:spacing w:before="0" w:after="0" w:line="200" w:lineRule="auto"/>
                  <w:jc w:val="left"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>
                    <w:b/>
                    <w:color w:val="{color_hex}"/>
                    <w:sz w:val="{sz_val}"/>
                  </w:rPr>
                  <w:t>Действителен: c 02.09.2025 по 26.11.2026</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr wrap="square" lIns="91440" tIns="45720" rIns="91440" bIns="45720" anchor="t"/>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''
    from lxml import etree as _etree
    drawing_elem = _etree.fromstring(drawing_xml)
    run._element.append(drawing_elem)


# ============================================================================
# Main builder
# ============================================================================

def build_template() -> Document:
    doc = Document()
    _set_landscape(doc)

    # Шрифт по умолчанию
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_SIZE_PT)

    _build_applicant_section(doc)
    _build_table1(doc)
    _build_table2(doc)
    _build_signature(doc)

    return doc


def main() -> int:
    print(f"Building STDR template v6 → {OUTPUT_PATH}")
    print(f"Column widths (T1): {[round(w, 2) for w in T1_COL_WIDTHS]}")
    print(f"  Sum: {round(sum(T1_COL_WIDTHS), 2)} cm")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_9_C_r6_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")

    doc = build_template()
    doc.save(str(OUTPUT_PATH))

    size = OUTPUT_PATH.stat().st_size
    print(f"✅ Saved: {OUTPUT_PATH}  ({size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
