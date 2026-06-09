# -*- coding: utf-8 -*-
"""
Pack 48.1 — build-скрипт для DOCX-шаблона выписки ТБанка (BIK 044525974).

Создаёт `bank_statement_template_044525974.docx` рядом со скриптом. Этот файл
кладётся в `templates/docx/`, и resolver в docx_renderer.py (Pack 47.0)
автоматически выбирает его для applicants с bank.bik == "044525974".

Архитектура шаблона (по эталону PDF "Справка о движении средств" ТБанк):

  ┌─────────────────────────────────────────────────────────┐
  │ HEADER (1 раз, левая ячейка = лого, правая = реквизиты) │
  ├─────────────────────────────────────────────────────────┤
  │ TITLE "Справка о движении средств" (h1)                 │
  │ "Исх. № e2e769c8" слева | "DD.MM.YYYY" справа           │
  ├─────────────────────────────────────────────────────────┤
  │ CLIENT BLOCK                                            │
  │   ФИО (bold)                                            │
  │   Адрес места жительства: ...                           │
  │   Серия | Номер | Дата выдачи | Код подразделения       │
  │   Паспорт выдан: ...                                    │
  │   Адрес регистрации: ...                                │
  ├─────────────────────────────────────────────────────────┤
  │ "О продукте" (h2)                                       │
  │   Дата заключения договора: ...                         │
  │   Номер договора: ...                                   │
  │   Номер лицевого счета: ...                             │
  │   Сумма доступного остатка на DD.MM.YYYY: ...           │
  ├─────────────────────────────────────────────────────────┤
  │ "Движение средств за период с ... по ..." (h2)          │
  ├─────────────────────────────────────────────────────────┤
  │ TX TABLE (6 columns, 1 шаблонная строка с маркерами)    │
  │   Дата и время | Дата списания | Сумма в валюте операции│
  │   | Сумма в валюте карты | Описание | Номер карты       │
  │                                                         │
  │   __TX_DATE__  __TX_DATE_SETTLE__  __TX_AMOUNT__        │
  │   __TX_AMOUNT_CARD__  __TX_DESCRIPTION__  __TX_CARD__   │
  │                                                         │
  │   docx_renderer.py клонирует эту строку для каждой tx.  │
  ├─────────────────────────────────────────────────────────┤
  │ TOTALS (только на последней странице через page break)  │
  │   Пополнения: {{ bank.total_income_formatted }}         │
  │   Расходы:    {{ bank.total_expense_formatted }}        │
  ├─────────────────────────────────────────────────────────┤
  │ SIGNATURE BLOCK (inline PNG)                            │
  │   tbank_signature.png                                   │
  ├─────────────────────────────────────────────────────────┤
  │ FOOTER (каждая страница):                               │
  │   АО «ТБанк» универсальная лицензия... БИК... ИНН... КПП│
  │   стр. N                                                │
  └─────────────────────────────────────────────────────────┘

Цвета бренда ТБанка (вытащены из эталона):
  #FBDD2D — жёлтый бренда (лого)
  #030303 — основной текст (заголовки, ФИО, значения)
  #929292 — серые лейблы (заголовки колонок, "Серия:", "Дата выдачи:" и т.п.)
  #E5E5E5 — тонкие разделители между строками info-блока

Jinja-плейсхолдеры (заполняются docxtpl в Phase 1):
  applicant.full_name_native
  applicant.address_living_full / applicant.address_living_short
  applicant.passport_series, .passport_number, .passport_issue_date,
    .passport_department_code, .passport_issuer
  applicant.address_registered_full / applicant.address_registered_short
  bank.contract_date_formatted (Дата заключения договора)
  bank.contract_number (Номер договора)
  applicant.bank_account (Номер лицевого счёта)
  bank.statement_date_formatted (Сумма доступного остатка на DD.MM.YYYY)
  bank.closing_balance_formatted (значение остатка)
  bank.period_start_formatted / .period_end_formatted
  bank.outgoing_number (Исх. №)
  bank.statement_date_formatted (дата справа от Исх.№ — да, ту же дату)
  bank.total_income_formatted / .total_expense_formatted

Маркеры в tx-строке (заменяются на Фазе 2 в docx_renderer):
  __TX_DATE__         (содержит "\\n" — дата+время операции)
  __TX_DATE_SETTLE__  (содержит "\\n" — дата+время списания)
  __TX_AMOUNT__       (сумма в валюте операции, например "-964.00 ₽")
  __TX_AMOUNT_CARD__  (сумма в валюте карты, для рублёвых = то же что __TX_AMOUNT__)
  __TX_DESCRIPTION__  (описание операции, может быть длинным — переносится)
  __TX_CARD__         (4 цифры или "—")

Использование:
  cd D:\\VISA\\visa_kit
  python build_tbank_template_v1.py
  # → создан bank_statement_template_044525974.docx рядом со скриптом
"""
from __future__ import annotations
import sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Inches, Emu
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# === Константы ===

REPO = Path(__file__).resolve().parent
TBANK_LOGO_PATH = REPO / "tbank_logo.png"          # будет рядом со скриптом
TBANK_SIGNATURE_PATH = REPO / "tbank_signature.png"  # будет рядом со скриптом

OUTPUT_PATH = REPO / "bank_statement_template_044525974.docx"

# Цвета бренда (из эталона)
COLOR_BLACK = RGBColor(0x03, 0x03, 0x03)
COLOR_GRAY_LABEL = RGBColor(0x92, 0x92, 0x92)  # лейблы серые
COLOR_GRAY_LINE = RGBColor(0xE5, 0xE5, 0xE5)   # тонкие разделители
COLOR_GRAY_FOOTER = RGBColor(0x80, 0x80, 0x80)
COLOR_YELLOW_BRAND = RGBColor(0xFB, 0xDD, 0x2D)


# === Helpers ===

def _set_run_font(run, name="Arial", size_pt=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # rFonts ascii/hAnsi/cs — чтобы был один и тот же шрифт во всех зонах
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None,
                      color="auto", size=4):
    """Расставляет границы ячейки. None = не трогать. 'nil' = убрать."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is None:
            continue
        el = tcBorders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcBorders.append(el)
        if val == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), val)  # "single", "dashed", и т.п.
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color if isinstance(color, str) else f"{color.rgb}")


def _set_cell_margins(cell, *, top=80, bottom=80, left=120, right=120):
    """dxa: 1 inch = 1440 dxa. Стандарт ~120 = 2mm."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcMar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, hex_color):
    """Заливка ячейки."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_table_borders_none(table):
    """Убирает все границы у таблицы (рамки и внутренние)."""
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        return
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    else:
        # Очистить
        for child in list(tblBorders):
            tblBorders.remove(child)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)


def _set_paragraph_spacing(paragraph, *, before=0, after=0, line=None, line_rule="auto"):
    """Межстрочный интервал и отступы абзаца. before/after в dxa (1pt=20dxa)."""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), line_rule)


def _add_run_simple(p, text, *, size_pt=10, bold=False, color=None):
    """Добавляет run с текстом и форматированием."""
    r = p.add_run(text)
    _set_run_font(r, size_pt=size_pt, bold=bold, color=color)
    return r


def _add_bottom_separator(cell, color="E5E5E5"):
    """Тонкая нижняя граница серая — для info-блока."""
    _set_cell_borders(cell, bottom="single", color=color, size=4)


# === Сборка шаблона ===

def build_template() -> None:
    """Главная функция сборки."""
    if not TBANK_LOGO_PATH.exists():
        print(f"❌ Не найден tbank_logo.png по пути: {TBANK_LOGO_PATH}", file=sys.stderr)
        print(f"   Положи tbank_logo.png рядом со скриптом и повтори.", file=sys.stderr)
        sys.exit(1)
    if not TBANK_SIGNATURE_PATH.exists():
        print(f"❌ Не найден tbank_signature.png по пути: {TBANK_SIGNATURE_PATH}", file=sys.stderr)
        print(f"   Положи tbank_signature.png рядом со скриптом и повтори.", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # === Настройка страницы ===
    section = doc.sections[0]
    section.page_width = Mm(210)   # A4
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(22)

    # === Дефолтный стиль документа ===
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style.font.color.rgb = COLOR_BLACK

    # === 1. HEADER: лого + реквизиты банка ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    # Ширина страницы за вычетом полей = 210 - 40 = 170 mm
    header_table.columns[0].width = Mm(35)
    header_table.columns[1].width = Mm(135)
    _set_table_borders_none(header_table)

    cell_logo, cell_reqs = header_table.rows[0].cells
    cell_logo.width = Mm(35)
    cell_reqs.width = Mm(135)
    _set_cell_margins(cell_logo, top=0, bottom=0, left=0, right=0)
    _set_cell_margins(cell_reqs, top=0, bottom=0, left=0, right=0)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_reqs.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Лого — в первой ячейке, размер ~17×17мм
    p_logo = cell_logo.paragraphs[0]
    _set_paragraph_spacing(p_logo, before=0, after=0, line=240, line_rule="auto")
    r = p_logo.add_run()
    r.add_picture(str(TBANK_LOGO_PATH), width=Mm(17))

    # Реквизиты банка — текст справа
    bank_lines = [
        "АКЦИОНЕРНОЕ ОБЩЕСТВО «ТБАНК»",
        "РОССИЯ, 127287, МОСКВА, УЛ. 2-Я ХУТОРСКАЯ, Д. 38А, СТР. 26",
        "ТЕЛ.: +7 495 648-10-00, TBANK.RU",
    ]
    for i, line in enumerate(bank_lines):
        if i == 0:
            p = cell_reqs.paragraphs[0]
        else:
            p = cell_reqs.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(p, before=0, after=0, line=260, line_rule="auto")
        _add_run_simple(p, line, size_pt=8, bold=True, color=COLOR_BLACK)

    # Воздух после шапки
    p_air = doc.add_paragraph()
    _set_paragraph_spacing(p_air, before=240, after=0)

    # === 2. TITLE "Справка о движении средств" ===
    p_title = doc.add_paragraph()
    _set_paragraph_spacing(p_title, before=0, after=240, line=400, line_rule="auto")
    _add_run_simple(p_title, "Справка о движении средств",
                    size_pt=24, bold=True, color=COLOR_BLACK)

    # === 3. Исх. № + дата (двухколоночная таблица) ===
    exch_table = doc.add_table(rows=1, cols=2)
    exch_table.autofit = False
    exch_table.columns[0].width = Mm(85)
    exch_table.columns[1].width = Mm(85)
    _set_table_borders_none(exch_table)

    cell_exch_l, cell_exch_r = exch_table.rows[0].cells
    cell_exch_l.width = Mm(85)
    cell_exch_r.width = Mm(85)
    _set_cell_margins(cell_exch_l, top=0, bottom=0, left=0, right=0)
    _set_cell_margins(cell_exch_r, top=0, bottom=0, left=0, right=0)

    p_l = cell_exch_l.paragraphs[0]
    _set_paragraph_spacing(p_l, before=0, after=0)
    _add_run_simple(p_l, "Исх. № ", size_pt=9, color=COLOR_GRAY_LABEL)
    _add_run_simple(p_l, "{{ bank.outgoing_number }}", size_pt=9, color=COLOR_BLACK)

    p_r = cell_exch_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p_r, before=0, after=0)
    _add_run_simple(p_r, "{{ bank.statement_date_formatted }}",
                    size_pt=9, color=COLOR_BLACK)

    # === 4. CLIENT BLOCK ===
    # ФИО (bold, чёрный, ~11pt)
    p_air2 = doc.add_paragraph()
    _set_paragraph_spacing(p_air2, before=240, after=0)

    p_fio = doc.add_paragraph()
    _set_paragraph_spacing(p_fio, before=0, after=120, line=280, line_rule="auto")
    _add_run_simple(p_fio, "{{ applicant.full_name_native }}",
                    size_pt=10, bold=True, color=COLOR_BLACK)
    # Тонкая нижняя граница абзаца
    pPr = p_fio._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E5E5E5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Адрес места жительства
    p_adj = doc.add_paragraph()
    _set_paragraph_spacing(p_adj, before=120, after=120, line=280, line_rule="auto")
    _add_run_simple(p_adj, "Адрес места жительства: ",
                    size_pt=9, bold=True, color=COLOR_BLACK)
    _add_run_simple(p_adj, "{{ applicant.address_living_full }}",
                    size_pt=9, color=COLOR_BLACK)
    # нижняя граница
    pPr = p_adj._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E5E5E5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Паспорт: 4 поля в одной таблице (Серия | Номер | Дата выдачи | Код подразделения)
    pass_table = doc.add_table(rows=1, cols=4)
    pass_table.autofit = False
    # Распределение: Серия 30мм, Номер 40мм, Дата выдачи 50мм, Код 50мм
    widths = [Mm(30), Mm(40), Mm(50), Mm(50)]
    for i, w in enumerate(widths):
        pass_table.columns[i].width = w
    _set_table_borders_none(pass_table)
    # Снизу таблицы добавим тонкую линию
    tblPr = pass_table._element.find(qn("w:tblPr"))
    tblBorders = tblPr.find(qn("w:tblBorders"))
    bot = tblBorders.find(qn("w:bottom"))
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), "E5E5E5")

    cells = pass_table.rows[0].cells
    for c, w in zip(cells, widths):
        c.width = w
        _set_cell_margins(c, top=80, bottom=80, left=0, right=0)
    pass_labels_values = [
        ("Серия: ", "{{ applicant.passport_series }}"),
        ("Номер: ", "{{ applicant.passport_number }}"),
        ("Дата выдачи: ", "{{ applicant.passport_issue_date_formatted }}"),
        ("Код подразделения: ", "{{ applicant.passport_department_code }}"),
    ]
    for cell, (label, value) in zip(cells, pass_labels_values):
        p = cell.paragraphs[0]
        _set_paragraph_spacing(p, before=0, after=0, line=260, line_rule="auto")
        _add_run_simple(p, label, size_pt=9, bold=True, color=COLOR_BLACK)
        _add_run_simple(p, value, size_pt=9, color=COLOR_BLACK)

    # "Паспорт выдан:"
    p_pi = doc.add_paragraph()
    _set_paragraph_spacing(p_pi, before=120, after=120, line=280, line_rule="auto")
    _add_run_simple(p_pi, "Паспорт выдан: ",
                    size_pt=9, bold=True, color=COLOR_BLACK)
    _add_run_simple(p_pi, "{{ applicant.passport_issuer }}",
                    size_pt=9, color=COLOR_BLACK)
    pPr = p_pi._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E5E5E5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Адрес регистрации
    p_areg = doc.add_paragraph()
    _set_paragraph_spacing(p_areg, before=120, after=120, line=280, line_rule="auto")
    _add_run_simple(p_areg, "Адрес регистрации: ",
                    size_pt=9, bold=True, color=COLOR_BLACK)
    _add_run_simple(p_areg, "{{ applicant.address_registered_full }}",
                    size_pt=9, color=COLOR_BLACK)
    pPr = p_areg._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E5E5E5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # === 5. "О продукте" ===
    p_air3 = doc.add_paragraph()
    _set_paragraph_spacing(p_air3, before=240, after=0)

    p_h2 = doc.add_paragraph()
    _set_paragraph_spacing(p_h2, before=0, after=200, line=320, line_rule="auto")
    _add_run_simple(p_h2, "О продукте", size_pt=14, bold=False, color=COLOR_BLACK)

    product_rows = [
        ("Дата заключения договора: ", "{{ bank.contract_date_formatted }}"),
        ("Номер договора: ", "{{ bank.contract_number }}"),
        ("Номер лицевого счета: ", "{{ applicant.bank_account }}"),
        ("Сумма доступного остатка на {{ bank.statement_date_formatted }}: ",
         "{{ bank.closing_balance_formatted }}"),
    ]
    for label, value in product_rows:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=120, after=120, line=280, line_rule="auto")
        _add_run_simple(p, label, size_pt=9, bold=True, color=COLOR_BLACK)
        _add_run_simple(p, value, size_pt=9, color=COLOR_BLACK)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "E5E5E5")
        pBdr.append(bot)
        pPr.append(pBdr)

    # === 6. "Движение средств за период..." ===
    p_air4 = doc.add_paragraph()
    _set_paragraph_spacing(p_air4, before=240, after=0)

    p_h2_tx = doc.add_paragraph()
    _set_paragraph_spacing(p_h2_tx, before=0, after=200, line=320, line_rule="auto")
    _add_run_simple(
        p_h2_tx,
        "Движение средств за период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}",
        size_pt=14, bold=False, color=COLOR_BLACK
    )

    # === 7. TX TABLE (6 колонок) ===
    # Содержимое: header row + 1 шаблонная row с маркерами.
    # docx_renderer.py клонирует вторую строку для каждой tx.
    tx_table = doc.add_table(rows=2, cols=6)
    tx_table.autofit = False
    # Распределение колонок (mm), сумма = 170mm:
    #   Дата операции: 25
    #   Дата списания: 25
    #   Сумма опер.: 27
    #   Сумма карты: 27
    #   Описание: 46
    #   Номер карты: 20
    tx_widths = [Mm(25), Mm(25), Mm(27), Mm(27), Mm(46), Mm(20)]
    for i, w in enumerate(tx_widths):
        tx_table.columns[i].width = w
    _set_table_borders_none(tx_table)

    # === Header row ===
    header_cells = tx_table.rows[0].cells
    header_labels = [
        "Дата и время\nоперации",
        "Дата\nсписания",
        "Сумма в валюте\nоперации",
        "Сумма операции\nв валюте карты",
        "Описание\nоперации",
        "Номер\nкарты",
    ]
    for cell, label, w in zip(header_cells, header_labels, tx_widths):
        cell.width = w
        _set_cell_margins(cell, top=80, bottom=80, left=0, right=80)
        # Нижняя граница серая — отделяет заголовок от данных
        _set_cell_borders(cell, bottom="single", color="E5E5E5", size=4)
        # Каждая строка лейбла — отдельный параграф (т.к. \n в Word не работает в одном run)
        parts = label.split("\n")
        for i, part in enumerate(parts):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=0, line=240, line_rule="auto")
            _add_run_simple(p, part, size_pt=8, color=COLOR_GRAY_LABEL)

    # === Template row с маркерами ===
    tx_cells = tx_table.rows[1].cells
    # Дата операции (содержит \n → multiline в docx_renderer)
    # Маркер: только сам __TX_DATE__, _replace_marker_with_multiline на проде
    # развернёт "DD.MM.YYYY\nHH:MM" в 2 параграфа.
    tx_markers = [
        "__TX_DATE__",
        "__TX_DATE_SETTLE__",
        "__TX_AMOUNT__",
        "__TX_AMOUNT_CARD__",
        "__TX_DESCRIPTION__",
        "__TX_CARD__",
    ]
    for cell, marker, w in zip(tx_cells, tx_markers, tx_widths):
        cell.width = w
        _set_cell_margins(cell, top=100, bottom=100, left=0, right=80)
        # Нижняя граница серая — между tx-строками
        _set_cell_borders(cell, bottom="single", color="E5E5E5", size=4)
        p = cell.paragraphs[0]
        _set_paragraph_spacing(p, before=0, after=0, line=240, line_rule="auto")
        _add_run_simple(p, marker, size_pt=9, color=COLOR_BLACK)

    # === 8. TOTALS (Пополнения / Расходы) — после tx-таблицы ===
    p_air5 = doc.add_paragraph()
    _set_paragraph_spacing(p_air5, before=320, after=0)

    # Таблица 2 колонки: лейбл + значение
    tot_table = doc.add_table(rows=2, cols=2)
    tot_table.autofit = False
    tot_table.columns[0].width = Mm(35)
    tot_table.columns[1].width = Mm(60)
    _set_table_borders_none(tot_table)

    tot_rows_data = [
        ("Пополнения:", "{{ bank.total_income_formatted }}"),
        ("Расходы:", "{{ bank.total_expense_formatted }}"),
    ]
    for row_idx, (label, value) in enumerate(tot_rows_data):
        cells = tot_table.rows[row_idx].cells
        cells[0].width = Mm(35)
        cells[1].width = Mm(60)
        _set_cell_margins(cells[0], top=40, bottom=40, left=0, right=0)
        _set_cell_margins(cells[1], top=40, bottom=40, left=0, right=0)
        p_l = cells[0].paragraphs[0]
        _set_paragraph_spacing(p_l, before=0, after=0, line=260, line_rule="auto")
        _add_run_simple(p_l, label, size_pt=9, color=COLOR_BLACK)
        p_r = cells[1].paragraphs[0]
        _set_paragraph_spacing(p_r, before=0, after=0, line=260, line_rule="auto")
        _add_run_simple(p_r, value, size_pt=9, color=COLOR_BLACK)

    # === 9. SIGNATURE BLOCK (inline PNG) ===
    p_air6 = doc.add_paragraph()
    _set_paragraph_spacing(p_air6, before=480, after=0)

    p_sig = doc.add_paragraph()
    _set_paragraph_spacing(p_sig, before=0, after=0)
    r_sig = p_sig.add_run()
    # tbank_signature.png — 1937×360 px @ 300 dpi = ~164×30 mm.
    # Ставлю ширину 160мм (вписывается в 170мм рабочую зону страницы).
    r_sig.add_picture(str(TBANK_SIGNATURE_PATH), width=Mm(160))

    # === 10. FOOTER (на каждой странице) ===
    # Подвал через section.footer
    footer = section.footer
    # python-docx даёт нам один пустой параграф в footer'е
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_foot, before=0, after=0, line=240, line_rule="auto")
    _add_run_simple(
        p_foot,
        "АО «ТБанк» универсальная лицензия Банка России № 2673, "
        "к/с 30101810145250000974 в ГУ Банка России по ЦФО",
        size_pt=7, color=COLOR_GRAY_FOOTER
    )
    p_foot2 = footer.add_paragraph()
    p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_foot2, before=0, after=0, line=240, line_rule="auto")
    _add_run_simple(
        p_foot2,
        "БИК 044525974 ИНН 7710140679 КПП 771301001",
        size_pt=7, color=COLOR_GRAY_FOOTER
    )
    # Номер страницы — справа от подвала. Делаю через PAGE field.
    # Чтобы не ломать центрирование двух предыдущих параграфов — добавлю
    # ТРЕТИЙ параграф с tab-stop и PAGE field.
    p_pgnum = footer.add_paragraph()
    p_pgnum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p_pgnum, before=0, after=0, line=240, line_rule="auto")
    r_pgnum = p_pgnum.add_run()
    _set_run_font(r_pgnum, size_pt=7, color=COLOR_GRAY_FOOTER)
    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_pgnum._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    r_pgnum._element.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_pgnum._element.append(fld_end)

    # === SAVE ===
    doc.save(str(OUTPUT_PATH))
    size_kb = OUTPUT_PATH.stat().st_size / 1024
    print(f"✓ Шаблон собран: {OUTPUT_PATH} ({size_kb:.1f} KB)")
    print(f"  Положи его в templates/docx/ через apply-скрипт Pack 48.1.")


if __name__ == "__main__":
    build_template()
