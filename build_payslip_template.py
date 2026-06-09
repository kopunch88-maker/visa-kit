r"""
Pack 50.10-C — Генератор payslip_template.docx.

Структура по эталону "Расчетный_листок_за_<MM>_<YYYY>.docx":
  - Portrait A4 (21×29.7 cm)
  - Шапка из параграфов:
      P1: "РАСЧЕТНЫЙ ЛИСТОК ЗА {{ payslip.month_title_upper }}"  (10pt center)
      P2: пусто
      P3: "ФИО: {{ payslip.applicant_full_name }}"  (11pt)
      P4: "К выплате:        {{ payslip.payout_fmt }}"  (10pt)
      P5: "{{ payslip.company_full_name_upper }}"  (11pt bold)
      P6: "Должность:\t{{ payslip.position_title }}"  (11pt bold)
      P7: "Подразделение: {{ payslip.department }}\tОклад (тариф): {{ payslip.salary_oklad_raw }}"
      P8: "Долг предприятия на начало 0,00 Долг предприятия на конец 0,00 Общий
           облагаемый доход за {{ payslip.year }} год нарастающим итогом: {{ payslip.yearly_total_fmt }}"
  - Таблица 5×9
  - P: "Бухгалтер                                   {{ payslip.accountant_short }}"

Запуск:
    python build_payslip_template.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "payslip_template.docx"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LR_CM = 1.0
MARGIN_TB_CM = 1.5
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_LR_CM  # 19.0

FONT_FAMILY = "Arial"
FONT_SIZE_BASE = 10
FONT_SIZE_BIG = 11

# Ширины колонок таблицы (из эталона, нормализованы к 19.0 cm)
RAW_WIDTHS = [3.47, 1.98, 1.09, 1.09, 1.84, 1.97, 3.82, 1.63, 2.18]
TOTAL_RAW = sum(RAW_WIDTHS)  # 19.07
COL_WIDTHS = [w * USABLE_WIDTH_CM / TOTAL_RAW for w in RAW_WIDTHS]


def _set_run(run, *, size=FONT_SIZE_BASE, bold=False, italic=False):
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{k}"), FONT_FAMILY)


def _add_paragraph(doc, text="", *, size=FONT_SIZE_BASE, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before_pt=0, space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if text:
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, italic=italic)
    return p


def _cell_text(cell, text, *, size=FONT_SIZE_BASE, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center"):
    cell.text = ""
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[vertical]
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold, italic=italic)


def _set_cell_borders(cell, *, top=True, bottom=True, left=True, right=True, size_val="4"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    sides = [("top", top), ("left", left), ("bottom", bottom), ("right", right)]
    for side_name, enabled in sides:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        border = OxmlElement(f"w:{side_name}")
        if enabled:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size_val)
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "nil")
        tcBorders.append(border)


def _set_table_borders_tbl(table):
    """tblBorders на уровне таблицы — корректно работает с merged."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_fixed_layout(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _apply_widths(table, widths_cm: list):
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _merge_horizontal(row, start_col: int, end_col: int):
    cells = row.cells
    merged = cells[start_col]
    for ci in range(start_col + 1, end_col + 1):
        merged = merged.merge(cells[ci])
    return merged


def _set_vmerge(cell, val):
    tcPr = cell._element.get_or_add_tcPr()
    vMerge = tcPr.find(qn("w:vMerge"))
    if vMerge is None:
        vMerge = OxmlElement("w:vMerge")
        tcPr.append(vMerge)
    if val == "restart":
        vMerge.set(qn("w:val"), "restart")
    elif val is None:
        if qn("w:val") in vMerge.attrib:
            del vMerge.attrib[qn("w:val")]


def build_template() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_SIZE_BASE)

    # P1: Заголовок (10pt center)
    _add_paragraph(
        doc,
        "РАСЧЕТНЫЙ ЛИСТОК ЗА {{ payslip.month_title_upper }}",
        size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    # P3: ФИО (11pt)
    _add_paragraph(
        doc,
        "ФИО: {{ payslip.applicant_full_name }}",
        size=FONT_SIZE_BIG,
        space_before_pt=4,
    )

    # P4: К выплате
    _add_paragraph(
        doc,
        "К выплате:        {{ payslip.payout_fmt }}",
        size=FONT_SIZE_BASE,
    )

    # P5: Компания (11pt bold)
    _add_paragraph(
        doc,
        "{{ payslip.company_full_name_upper }}",
        size=FONT_SIZE_BIG, bold=True,
        space_before_pt=4,
    )

    # P6: Должность (11pt bold)
    _add_paragraph(
        doc,
        "Должность:\t{{ payslip.position_title }}",
        size=FONT_SIZE_BIG, bold=True,
    )

    # P7: Подразделение + Оклад (10pt)
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p7.paragraph_format.space_before = Pt(2)
    p7.paragraph_format.space_after = Pt(2)
    r = p7.add_run("Подразделение: {{ payslip.department }}\tОклад (тариф): {{ payslip.salary_oklad_raw }}")
    _set_run(r, size=FONT_SIZE_BASE)

    # P8: Долг + нарастающий итог
    _add_paragraph(
        doc,
        "Долг предприятия на начало\t0,00 Долг предприятия на конец                         "
        "0,00 Общий облагаемый доход за {{ payslip.year }} год нарастающим итогом: "
        "                                                                              "
        "{{ payslip.yearly_total_fmt }}",
        size=FONT_SIZE_BASE,
        space_after_pt=6,
    )

    # === ТАБЛИЦА 5×9 ===
    t = doc.add_table(rows=5, cols=9)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, COL_WIDTHS)

    # Row 0: шапка
    # Col 0,1: vMerge restart (Вид, Период)
    # Col 2,3: gridSpan=2 "Рабочие" в row 0
    # Col 4,5: vMerge restart (Оплачено, Сумма)
    # Col 6,7,8: vMerge restart (Вид, Период, Сумма)
    _cell_text(t.rows[0].cells[0], "Вид",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[1], "Период",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    rabochie_cell = _merge_horizontal(t.rows[0], 2, 3)
    _cell_text(rabochie_cell, "Рабочие",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    # cells[4] после merge — это бывшая cell 4 (Оплачено)
    _cell_text(t.rows[0].cells[3], "Оплачено",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[4], "Сумма",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[5], "Вид",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[6], "Период",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[7], "Сумма",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")

    # vMerge restart для всех кроме merged "Рабочие"
    # ВАЖНО: после merge_horizontal индексы ячеек смещаются.
    # row 0 после merge: [0=Вид, 1=Период, 2=Рабочие(merged), 3=Оплачено, 4=Сумма, 5=Вид, 6=Период, 7=Сумма]
    for ci in [0, 1, 3, 4, 5, 6, 7]:
        _set_vmerge(t.rows[0].cells[ci], "restart")

    # Row 1: подзаголовки таблицы "Рабочие"
    # Все ячейки кроме col 2 и 3 — vMerge=continue
    # row 1 (тут логические колонки 0..8)
    # 0: continue
    # 1: continue
    # 2: "Дни"
    # 3: "Часы"
    # 4-8: continue
    _cell_text(t.rows[1].cells[2], "Дни",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[1].cells[3], "Часы",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    for ci in [0, 1, 4, 5, 6, 7, 8]:
        _set_vmerge(t.rows[1].cells[ci], None)  # continue

    # Row 2: Начислено: (gridSpan=6, cols 0-5) | Удержано: (gridSpan=3, cols 6-8)
    nachisleno = _merge_horizontal(t.rows[2], 0, 5)
    _cell_text(nachisleno, "Начислено:",
               size=FONT_SIZE_BASE, bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    uderzhano = _merge_horizontal(t.rows[2], 1, 3)  # после первого merge cells сместились
    # После merge1 row 2 имеет [0=Начислено(GS=6), 1=cell6, 2=cell7, 3=cell8]
    # merge_horizontal(row, 1, 3) объединит cell6+cell7+cell8 → 1 ячейка
    _cell_text(uderzhano, "Удержано:                       {{ payslip.ndfl_fmt }}",
               size=FONT_SIZE_BASE, bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")

    # Row 3: данные
    _cell_text(t.rows[3].cells[0], "Оплата по окладу",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    _cell_text(t.rows[3].cells[1], "{{ payslip.month_short }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[2], "{{ payslip.working_days }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[3], "{{ payslip.working_hours }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[4], "{{ payslip.days_paid_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[5], "{{ payslip.salary_amount_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.RIGHT, vertical="center")
    _cell_text(t.rows[3].cells[6], "НДФЛ",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    _cell_text(t.rows[3].cells[7], "{{ payslip.month_short_dotted }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[8], "{{ payslip.ndfl_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.RIGHT, vertical="center")

    # Row 4: пусто 0-5 | "Выплачено:" gridSpan=3 (cols 6-8)
    vyplacheno = _merge_horizontal(t.rows[4], 6, 8)
    _cell_text(vyplacheno,
               "Выплачено:                                 {{ payslip.payout_fmt }}",
               size=FONT_SIZE_BASE, bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    # Cells 0..5 в row 4 — пусто
    for ci in range(6):
        _cell_text(t.rows[4].cells[ci], "", size=FONT_SIZE_BASE)

    # Border'ы — tblBorders на всю таблицу
    _set_table_borders_tbl(t)

    # Подпись внизу
    _add_paragraph(doc, "", space_before_pt=12)
    _add_paragraph(doc, "", space_before_pt=4)
    _add_paragraph(
        doc,
        "Бухгалтер                                              {{ payslip.accountant_short }}",
        size=FONT_SIZE_BASE,
    )

    return doc


def main() -> int:
    print(f"Building payslip template → {OUTPUT_PATH}")
    print(f"Column widths: {[round(w, 2) for w in COL_WIDTHS]}")
    print(f"  Sum: {round(sum(COL_WIDTHS), 2)} cm")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_10_C_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")

    doc = build_template()
    doc.save(str(OUTPUT_PATH))

    size = OUTPUT_PATH.stat().st_size
    print(f"✅ Saved: {OUTPUT_PATH}  ({size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
