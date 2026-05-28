r"""
Pack 50.12-C-r8 — Генератор soo_template.docx (Свидетельство об отъезде).

Версия r8 — финальная правка блока 3:
  "Трудовой договор" сдвинут вправо отступом 2см (под слово "ПРИКАЗ"),
  как в эталоне (там сдвиг сделан 21 ведущим пробелом). "Основание"
  остаётся у левого края. Так Приказ и Трудовой договор сгруппированы.

Версия r7 — приведение к эталону Регины по шрифту и воздуху:
  - ВЕСЬ документ 10pt (эталон везде sz=20=10pt, НЕ 8pt — заметка в handoff
    про "8pt + артефакт зума" оказалась ошибочной, проверено по document.xml)
  - больше воздуха: отступы ячеек блоков увеличены, между заголовком и блоком
    больше пространства
  - блок 2: пустые строки-разделители между смысловыми группами (как эталон)
  - блок 1 (п.1): Телефон/Эл.почта/Адрес-значения НЕ жирные, жирные только метки
  - блок 1 (п.2): "Фамилия" и значение выровнены по табу 0.63см под "1.1."
  - блок 3 (п.1): Основание и Трудовой договор выровнены друг под другом
  - блок 3: left-выравнивание (эталон не justify)

Версия r6 — пакет правок под эталон Регины (по скриншотам):
  0. (из r5) внутренние линии блоков убраны — только внешняя рамка
  1. блок 2 не растянут: межстрочный line=1.0, отступы ячеек top/bottom 40->20
  2. блок 4: Дата/Подпись прижаты вправо (align RIGHT) + не жирные;
     добавлен пустой параграф снизу (пространство под подписью)
  3. блок 3: "Трудовой договор" с отступом (как продолжение, с табом),
     "Основание" остаётся с новой строки
  4. блок 1: номера 1.1./1.2./1.3. НЕ жирные (метки слов остаются жирными)
  5. между заголовком раздела и блоком — пустой параграф-разделитель
     (через _section_header добавляется отступ space_after)

Версия r5 — убраны ВНУТРЕННИЕ линии внутри блоков-таблиц (разделы 1-4).
  Остаётся только внешняя рамка каждого блока (как у эталона Регины, где
  «строки» внутри визуально без разделителей). insideH/insideV = none,
  внешние top/left/bottom/right = single.
  Структура (таблицы-обёртки) сохранена — переписывать на pBdr-параграфы
  НЕ требуется.

Версия r4 — финальная доводка под эталон Регины (по скриншотам Word):
  - Раздел 1: убран вертикальный воздух в ячейках (top/bottom 15->0), строки плотные
  - Квадратики-номера 1./2./3./4.: компактные ~0.5см, высота строки под номер (не 0.55)
  - Таблицы остаются полноширинными (эталон tblW=19.05/18.81см — НЕ сужать,
    «отступ справа» на скрине эталона — артефакт короткого текста + зум 80%)
  - Колонка номера в разделах 2/4 ~0.95см (эталон 536dxa)

Версия r3 (база) — 1:1 с эталоном Регины:
  - Номера разделов 1./2./3./4. в МАЛЕНЬКИХ КВАДРАТНЫХ РАМКАХ слева
    (таблица 1×1) + заголовок текстом справа
  - РАЗДЕЛ 1 — таблица в рамке: метки Фамилия/Имя/Отчество по колонкам (жирн),
    значения под ними; строки Гражданство/Документ/Загран/ДР+СНИЛС/Адрес/Телефон
  - РАЗДЕЛ 2 — таблица 4×2 в рамке (узкая колонка номера + контент), просторно
  - РАЗДЕЛ 3 — таблица в рамке (3.1 + Основание)
  - РАЗДЕЛ 4 — таблица 3×1 в рамке (СФР, захардкожен по эталону Регины)
  - Без серой заливки. Times New Roman 8pt (блок 4 — 10pt).

Запуск:
    python build_soo_template_v3.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "soo_template.docx"

PAGE_W, PAGE_H = 21.0, 29.7
MARGIN_LR, MARGIN_T, MARGIN_B = 1.27, 1.2, 1.0
FONT = "Times New Roman"
SZ = 10
CONTENT_W_CM = 18.46  # ширина контента (A4 - поля)
CONTENT_W_DXA = int(CONTENT_W_CM * 567)


# ---------------------------------------------------------------------------
# Низкоуровневые хелперы
# ---------------------------------------------------------------------------
def _set_run(run, *, size=SZ, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{k}"), FONT)


def _p_runs(p, runs, *, size=SZ, align=None, space_before=0, space_after=0, line=1.0):
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        lines = text.split("\n")
        for li, line_txt in enumerate(lines):
            if li > 0:
                br = p.add_run()
                br.add_break()
            r = p.add_run(line_txt)
            _set_run(r, size=size, bold=bold)


def _add_p(doc, runs, *, size=SZ, align=WD_ALIGN_PARAGRAPH.LEFT,
           space_before=0, space_after=0, line=1.0):
    p = doc.add_paragraph()
    _p_runs(p, runs, size=size, align=align,
            space_before=space_before, space_after=space_after, line=line)
    return p


def _no_table_borders(table):
    """Убрать все границы таблицы (для таблиц-обёрток без видимой рамки)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def _set_cell_width(cell, cm):
    cell.width = Cm(cm)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top=20, bottom=20, left=60, right=60):
    """Внутренние отступы ячейки в твипах (просторность)."""
    tcPr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _cell(cell, runs, *, size=SZ, align=WD_ALIGN_PARAGRAPH.LEFT, valign="center",
          space_after=0, line=1.0):
    cell.text = ""
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP, "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[valign]
    p = cell.paragraphs[0]
    _p_runs(p, runs, size=size, align=align, space_after=space_after, line=line)


# ---------------------------------------------------------------------------
# Заголовок раздела: квадратик-номер + заголовок текстом
# ---------------------------------------------------------------------------
def _section_header(doc, num, title):
    """Номер в квадратной рамке (маленькая таблица 1×2): [N.] | заголовок."""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    # колонка 1 — квадратик (узкая), колонка 2 — заголовок (широкая)
    _set_cell_width(t.rows[0].cells[0], 0.5)
    _set_cell_width(t.rows[0].cells[1], CONTENT_W_CM - 0.5)

    # Квадратик с рамкой (8pt, как эталон)
    c0 = t.rows[0].cells[0]
    _set_cell_margins(c0, top=10, bottom=10, left=20, right=20)
    _cell(c0, [(num, False)], size=8, align=WD_ALIGN_PARAGRAPH.CENTER, valign="center")
    # рамка только вокруг квадратика
    tcPr = c0._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tcPr.append(borders)

    # Заголовок (без рамки, 8pt как эталон)
    c1 = t.rows[0].cells[1]
    _cell(c1, [(title, True)], size=8, align=WD_ALIGN_PARAGRAPH.LEFT, valign="center")

    # убрать остальные границы таблицы-обёртки (кроме квадратика)
    tblPr = t._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        tblBorders.append(e)
    tblPr.append(tblBorders)

    # высота строки — минимально по контенту (atLeast), не раздувать
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    t.rows[0].height = Cm(0.45)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Pack 50.12 r6/r7 (п.5): пространство между заголовком и блоком (увеличено)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(2)
    sep.paragraph_format.line_spacing = 1.0
    r = sep.add_run("")
    _set_run(r, size=4)
    return t


def _bordered_table(doc, rows, cols):
    """Таблица с границами: только ВНЕШНЯЯ рамка, без внутренних линий.

    Перекрывает стиль Table Grid: top/left/bottom/right = single,
    insideH/insideV = none (визуально как блок эталона Регины).
    """
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.autofit = False
    tblPr = t._tbl.tblPr
    # удалить унаследованный tblBorders, если есть
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    for edge in ("insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)
    return t


# ---------------------------------------------------------------------------
# Построение документа
# ---------------------------------------------------------------------------
def build_template() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(PAGE_W)
    s.page_height = Cm(PAGE_H)
    s.left_margin = Cm(MARGIN_LR)
    s.right_margin = Cm(MARGIN_LR)
    s.top_margin = Cm(MARGIN_T)
    s.bottom_margin = Cm(MARGIN_B)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SZ)

    # ============ ШАПКА (эталон: 8pt) ============
    _add_p(doc, [(
        "ДОГОВОР МЕЖДУ РОССИЙСКОЙ ФЕДЕРАЦИЕЙ И КОРОЛЕВСТВОМ ИСПАНИЯ О СОЦИАЛЬНОМ "
        "ОБЕСПЕЧЕНИИ ОТ 11 АПРЕЛЯ 1994 г. / АДМИНИСТРАТИВНОЕ СОГЛАШЕНИЕ О ПРИМЕНЕНИИ "
        "ДОГОВОРА МЕЖДУ РОССИЙСКОЙ ФЕДЕРАЦИЕЙ И КОРОЛЕВСТВОМ ИСПАНИЯ О СОЦИАЛЬНОМ "
        "ОБЕСПЕЧЕНИИ ОТ 12 МАЯ 1995 г.", True)],
        size=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)

    _add_p(doc, [("СВИДЕТЕЛЬСТВО ОБ ОТЪЕЗДЕ", True)],
           size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _add_p(doc, [("Статья 7 Договора, статья 3 Административного соглашения", False)],
           size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

    _add_p(doc, [("Дата ", True), ("{{ soo.date_long }}", True), ("г.", True)],
           size=8, space_after=1)
    _add_p(doc, [("№ ", True), ("{{ soo.number }}", True)], size=8, space_after=3)

    # ============ РАЗДЕЛ 1 ============
    _section_header(doc, "1.",
                    "ЗАСТРАХОВАННОЕ ЛИЦО, ОСУЩЕСТВЛЯЮЩЕЕ ТРУДОВУЮ ИЛИ ИНУЮ ДЕЯТЕЛЬНОСТЬ")

    # Таблица раздела 1: рамка вокруг всего. Внутри строки.
    # Строка 1: метки Фамилия/Имя/Отчество (3 колонки)
    # Строка 2: значения last/first/middle (3 колонки)
    # Строки 3-8: на всю ширину (gridSpan 3)
    t1 = _bordered_table(doc, rows=8, cols=3)
    col_w = [CONTENT_W_CM / 3] * 3
    for row in t1.rows:
        for ci, c in enumerate(row.cells):
            _set_cell_width(c, col_w[ci])
            _set_cell_margins(c, top=10, bottom=10, left=60, right=60)

    # R0 метки (п.4: номер 1.1. НЕ жирный, слово-метка жирное)
    _cell(t1.rows[0].cells[0], [("1.1. ", False), ("Фамилия", True)], valign="top")
    _cell(t1.rows[0].cells[1], [("Имя", True)], valign="top")
    _cell(t1.rows[0].cells[2], [("Отчество", True)], valign="top")
    # R1 значения (п.2: last_name под "Фамилия" — отступ слева на ширину "1.1. ")
    _cell(t1.rows[1].cells[0], [("{{ soo.last_name }}", False)], valign="top")
    t1.rows[1].cells[0].paragraphs[0].paragraph_format.left_indent = Cm(0.63)
    _cell(t1.rows[1].cells[1], [("{{ soo.first_name }}", False)], valign="top")
    _cell(t1.rows[1].cells[2], [("{{ soo.middle_name }}", False)], valign="top")

    # Хелпер: строка на всю ширину (merge 3 ячеек)
    def _full_row(row_idx, runs):
        row = t1.rows[row_idx]
        a = row.cells[0]
        merged = a.merge(row.cells[1]).merge(row.cells[2])
        _cell(merged, runs, valign="top")
        return merged

    _full_row(2, [("Гражданство ", True), ("{{ soo.nationality_ru }}", False)])
    _full_row(3, [
        ("Документ, удостоверяющий личность: паспорт ", True),
        ("{{ soo.passport_series }} № {{ soo.passport_number_only }}, выдан "
         "{{ soo.passport_issue_date }} {{ soo.passport_issuer }}, код подразделения "
         "{{ soo.division_code }}", False),
    ])
    _full_row(4, [
        ("Номер загранпаспорта ", True),
        ("{{ soo.foreign_passport_number }} выдан {{ soo.foreign_passport_issue_date }} "
         "{{ soo.foreign_passport_issuer }}", False),
    ])
    _full_row(5, [
        ("1.2. ", False), ("Дата рождения ", True), ("{{ soo.birth_date_long }}", False),
        ("\nСтраховой номер индивидуального лицевого счёта в Российской Федерации "
         "(СНИЛС) ", True), ("{{ soo.snils }}", False),
    ])
    _full_row(6, [
        ("1.3. ", False), ("Адрес ", True), ("{{ soo.home_address }}", False),
    ])
    _full_row(7, [
        ("Телефон ", True), ("{{ soo.phone }}", False),
        ("       Факс ", True),
        ("       Эл. почта: ", True), ("{{ soo.email }}", False),
    ])

    _add_p(doc, "", space_after=2)

    # ============ РАЗДЕЛ 2 ============
    _section_header(doc, "2.",
                    "СВЕДЕНИЯ О РАБОТОДАТЕЛЕ ИЛИ ИНОЙ ДЕЯТЕЛЬНОСТИ В ГОСУДАРСТВЕ, "
                    "ПРИМЕНЯЕМОГО ЗАКОНОДАТЕЛЬСТВА")

    t2 = _bordered_table(doc, rows=4, cols=2)
    w_num, w_body = 0.95, CONTENT_W_CM - 0.95
    for row in t2.rows:
        _set_cell_width(row.cells[0], w_num)
        _set_cell_width(row.cells[1], w_body)
        for c in row.cells:
            _set_cell_margins(c, top=40, bottom=40, left=60, right=60)

    _cell(t2.rows[0].cells[0], "2.1.", valign="top")
    _cell(t2.rows[0].cells[1], [("Наименование", True)], valign="top")
    _cell(t2.rows[1].cells[0], "", valign="top")
    _cell(t2.rows[1].cells[1], [("{{ soo.company_full_name }}", False)], valign="top")
    _cell(t2.rows[2].cells[0], "2.2.", valign="top")
    # Ячейка 2.2: Адрес + пустая строка + Адрес в Испании (воздух как эталон)
    c22 = t2.rows[2].cells[1]
    c22.text = ""
    c22.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _p_runs(c22.paragraphs[0], [
        ("Адрес: ", True), ("{{ soo.company_legal_address }}", False),
    ], line=1.0, space_after=6)
    _p_runs(c22.add_paragraph(), [
        ("Адрес удалённой работы в Испании: ", True), ("{{ soo.spain_address }}", False),
    ], line=1.0)
    _cell(t2.rows[3].cells[0], "", valign="top")
    _cell(t2.rows[3].cells[1], [
        ("Телефон: ", True), ("{{ soo.company_phone }}", False),
        ("   Факс:    Эл. почта: ", True), ("{{ soo.company_email }}", False),
    ], valign="top")

    _add_p(doc, "", space_after=2)

    # ============ РАЗДЕЛ 3 ============
    _section_header(doc, "3.", "ПРИМЕНИМОЕ ЗАКОНОДАТЕЛЬСТВО")

    t3 = _bordered_table(doc, rows=1, cols=1)
    _set_cell_width(t3.rows[0].cells[0], CONTENT_W_CM)
    _set_cell_margins(t3.rows[0].cells[0], top=40, bottom=40, left=80, right=80)
    c3 = t3.rows[0].cells[0]
    c3.text = ""
    p = c3.paragraphs[0]
    _p_runs(p, [
        ("3.1. В отношении застрахованного лица продолжает применяться "
         "законодательство Российской Федерации: с ", False),
        ("{{ soo.period_start }}", True), (" по ", False),
        ("{{ soo.period_end }}", True),
        (". Согласно ст. 7 Договора между Российской Федерацией и Королевством "
         "Испании о социальном обеспечении от 11.04.1994 г. и ст. 3 "
         "Административного соглашения о применении Договора между РФ и Королевством "
         "Испания о социальном обеспечении от 12.05.1995 подтверждаем, что с "
         "применимым международным стандартом координации социального обеспечения, "
         "подписанным Королевством Испании, ", False),
        ("{{ soo.full_name }}", False),
        (" будет являться застрахованным лицом в системе обязательного социального "
         "страхования Российской Федерации.", False),
    ], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, line=1.0)
    p2 = c3.add_paragraph()
    _p_runs(p2, [
        ("Основание: ПРИКАЗ № ", True), ("{{ soo.order_number }}", False),
        (" от ", True), ("{{ soo.order_date }}", False),
        ("г. о направлении работника в командировку.", True),
    ], align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)
    # r8: "Трудовой договор" сдвинут вправо под "ПРИКАЗ" (эталон: ведущие пробелы ~2см)
    p3 = c3.add_paragraph()
    p3.paragraph_format.left_indent = Cm(2.0)
    _p_runs(p3, [
        ("Трудовой договор № ", False), ("{{ soo.contract_number }}", False),
        (" от ", False), ("{{ soo.contract_date }}", False),
        ("г. с возможностью дистанционной работы.", False),
    ], align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)

    _add_p(doc, "", space_after=2)

    # ============ РАЗДЕЛ 4 (СФР, захардкожен) ============
    _section_header(doc, "4.",
                    "ДАННЫЕ КОМПЕТЕНТНОГО УЧРЕЖДЕНИЯ, ЗАПОЛНИВШЕГО СВИДЕТЕЛЬСТВО ОБ ОТЪЕЗДЕ")

    t4 = _bordered_table(doc, rows=3, cols=1)
    for row in t4.rows:
        _set_cell_width(row.cells[0], CONTENT_W_CM)
        _set_cell_margins(row.cells[0], top=30, bottom=30, left=80, right=80)

    _cell(t4.rows[0].cells[0], [("Наименование", True)], size=10, valign="top")
    _cell(t4.rows[1].cells[0], [
        ("Отделение Фонда пенсионного и социального страхования Российской Федерации "
         "по г. Москве и Московской области", False)], size=10, valign="top")
    # Блок 4 строка 3: адрес/тел жирные, Дата/Печать-Подпись вправо и НЕ жирные (п.2)
    c4 = t4.rows[2].cells[0]
    c4.text = ""
    c4.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    pa = c4.paragraphs[0]
    _p_runs(pa, [
        ("Адрес 107078, Россия, г. Москва, Орликов пер., д. 3, корп. А", True),
        ("\n………………………………………………………………………………………………………………………………………..", True),
        ("\nТелефон +74959862611", True),
    ], size=10, line=1.0)
    # Дата — вправо, не жирная
    pdate = c4.add_paragraph()
    _p_runs(pdate, [("Дата ……………………………………….", False)],
            size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0)
    # Печать (слева) + Подпись (прижата вправо правым табом) — не жирные, одна строка
    psig = c4.add_paragraph()
    from docx.enum.text import WD_TAB_ALIGNMENT
    psig.paragraph_format.tab_stops.add_tab_stop(
        Cm(CONTENT_W_CM - 0.3), WD_TAB_ALIGNMENT.RIGHT)
    _p_runs(psig, [("Печать\tПодпись …………………………………...", False)],
            size=10, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)
    # пространство снизу под подписью (п.2)
    pgap = c4.add_paragraph()
    pgap.paragraph_format.space_after = Pt(0)
    _p_runs(pgap, [("", False)], size=10, line=1.0)

    return doc


def main() -> int:
    print(f"Building soo_template (v3, 1:1 эталон) → {OUTPUT_PATH}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_12_C_r3_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")
    doc = build_template()
    doc.save(str(OUTPUT_PATH))
    print(f"✅ Saved: {OUTPUT_PATH}  ({OUTPUT_PATH.stat().st_size/1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
