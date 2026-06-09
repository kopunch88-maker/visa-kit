r"""
Pack 50.11-A — Генератор employer_letter_naim_template.docx.

Письмо работодателя для НАЙМА (трудовой договор), по образцу
"Письмо_Заказчика.docx". Отличается от письма самозанятого:
  - "работает в <компания>" (не "оказывает услуги")
  - "по трудовому договору" (не "договору оказания услуг")
  - "на неопределённый срок (бессрочно)"
  - НЕТ фраз про независимого подрядчика / вне рынка труда

Использует те же плейсхолдеры что build_context уже отдаёт:
  letter.number, letter.date, fmt_date_ru, fmt_date_long_ru
  applicant.full_name_native
  company.full_name_ru, company.short_name, company.director_short_ru
  position.title_ru_genitive, position.duties
  contract.number, contract.sign_date, contract.salary_rub, contract.salary_rub_words
  eur.amount_int, eur.amount_words_ru
  fmt_money

Запуск:
    python build_employer_letter_naim_template.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "employer_letter_naim_template.docx"

PAGE_W, PAGE_H = 21.0, 29.7
MARGIN = 1.7
FONT_FAMILY = "Times New Roman"
FONT_SIZE = 12


def _set_run(run, *, size=FONT_SIZE, bold=False, italic=False):
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{k}"), FONT_FAMILY)


def _add_p(doc, text="", *, size=FONT_SIZE, bold=False, italic=False,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, italic=italic)
    return p


def build_template() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(PAGE_W)
    s.page_height = Cm(PAGE_H)
    s.left_margin = Cm(MARGIN)
    s.right_margin = Cm(MARGIN)
    s.top_margin = Cm(MARGIN)
    s.bottom_margin = Cm(MARGIN)

    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_SIZE)

    # Шапка под фирменный бланк — несколько пустых строк (как в образце P0-P9)
    for _ in range(8):
        _add_p(doc, "", size=FONT_SIZE)

    # Исх. № ... от ...
    _add_p(doc,
           "Исх. № {{ letter.number }} от {{ fmt_date_ru(letter.date) }}г.",
           size=11, space_after=12)

    # P: Настоящим подтверждаем, что <ФИО> работает в <компания> с <дата ТД>
    # по настоящее время в должности <позиция род.> по трудовому договору №... от ...
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "Настоящим подтверждаем, что {{ applicant.full_name_native }} работает "
        "в {{ company.full_name_ru }} ({{ company.short_name }}) "
        "с {{ fmt_date_long_ru(contract.sign_date) }} по настоящее время "
        "в должности {{ position.title_ru_genitive }} по трудовому договору "
        "№ {{ contract.number }} от {{ fmt_date_ru(contract.sign_date) }}г."
    )
    _set_run(r, size=FONT_SIZE)

    # P: удалённый формат
    _add_p(doc,
           "{{ applicant.full_name_native }} выполняет свои обязанности "
           "исключительно в удалённом формате с использованием средств "
           "телекоммуникации и цифровых технологий без необходимости "
           "физического присутствия в офисе компании.",
           size=FONT_SIZE, space_after=6)

    # P: В должностные обязанности сотрудника входит:
    _add_p(doc, "В должностные обязанности сотрудника входит:",
           size=FONT_SIZE, space_after=2)

    # Цикл по duties
    _add_p(doc, "{%p for duty in position.duties or [] %}", size=FONT_SIZE)
    _add_p(doc, "{{ duty }}", size=FONT_SIZE)
    _add_p(doc, "{%p endfor %}", size=FONT_SIZE)

    # P: Размер вознаграждения
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "Размер вознаграждения {{ applicant.full_name_native }} составляет "
        "{{ fmt_money(contract.salary_rub) }} ({{ contract.salary_rub_words }}) "
        "рублей или {{ eur.amount_int }} ({{ eur.amount_words_ru }}) евро "
        "по курсу ЦБ РФ на день составления письма, выплачиваемых ежемесячно "
        "на регулярной основе в соответствии с условиями трудового договора."
    )
    _set_run(r, size=FONT_SIZE)

    # P: бессрочно
    _add_p(doc,
           "Трудовой договор с {{ applicant.full_name_native }} заключён "
           "на неопределённый срок (бессрочно).",
           size=FONT_SIZE, space_after=6)

    # P: разрешение работать из Испании
    _add_p(doc,
           "{{ company.short_name }} официально разрешает "
           "{{ applicant.full_name_native }} выполнение работ по данному "
           "договору находясь на территории Королевства Испания.",
           size=FONT_SIZE, space_after=6)

    # Пустые строки перед подписью
    for _ in range(3):
        _add_p(doc, "", size=FONT_SIZE)

    # Подпись
    _add_p(doc,
           "Генеральный директор                                         "
           "                _____________/ {{ company.director_short_ru }}",
           size=FONT_SIZE)

    return doc


def main() -> int:
    print(f"Building employer letter (naim) template → {OUTPUT_PATH}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_11_A_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")
    doc = build_template()
    doc.save(str(OUTPUT_PATH))
    print(f"✅ Saved: {OUTPUT_PATH}  ({OUTPUT_PATH.stat().st_size/1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
