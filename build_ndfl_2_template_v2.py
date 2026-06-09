"""
Pack 50.8-C — Построение шаблона templates/docx/ndfl_2_template.docx.

Форма КНД 1175018 (приказ ФНС России от 19.09.2023 № ЕД-7-11/649@,
в ред. от 09.01.2024 № ЕД-7-11/1@).

Структура (по эталону data.nalog.ru/.../pril4_14525884.pdf):
  §1 Налоговый агент (ОКТМО, Телефон, ИНН, КПП + строка "Налоговый агент")
  §2 Физлицо (ИНН + ФИО + статус/ДР/гражданство + код док-та + серия+номер)
  §3 Доходы по ставке 13 % — ДВЕ таблицы рядом, в каждой 1 шапка + 14 строк
      (всего 28 слотов на год; пустые слоты просто остаются пустыми)
  §4 Стандартные/социальные/имущественные вычеты (4 пары "Код / Сумма")
  §5 Общая сумма дохода и сумма налога (2 колонки label/value × 4 строки)
  §6 Неудержанный налог (2 строки)
  Подпись (директор + место для подписи + сноски)

Все ячейки без заливки (Pack 50.8-C-r2: убрана серая F2F2F2).

Запуск из корня репо:
    python build_ndfl_2_template.py
    # → templates/docx/ndfl_2_template.docx
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# Конфигурация
# ============================================================================

FONT = "Arial"
FONT_SIZE_PT = 9
FONT_SIZE_HEADER_PT = 10
FONT_SIZE_SMALL_PT = 7
FONT_SIZE_TITLE_PT = 12

MARGIN_CM = 1.0
PAGE_WIDTH_CM = 21.0  # A4
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM  # 19.0

# §3 — количество строк данных в каждой колонке (по эталону ФНС — 14)
NDFL_3_ROWS_PER_COL = 14
NDFL_3_TOTAL_SLOTS = NDFL_3_ROWS_PER_COL * 2  # 28


# ============================================================================
# Хелперы форматирования
# ============================================================================

def _set_run(run, *, size=None, bold=False, italic=False, color=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(doc_or_cell, text="", *, size=None, bold=False, italic=False,
                   align=None, space_before_pt=0, space_after_pt=0, line_spacing=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_run(run, size=size or FONT_SIZE_PT, bold=bold, italic=italic)
    return p


def _set_cell_borders(cell, *, top=True, bottom=True, left=True, right=True, size_eighths=4):
    """Задаёт границы ячейки таблицы. size_eighths = толщина в 1/8 пункта (4 = 0.5pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    sides = {"top": top, "bottom": bottom, "left": left, "right": right}
    for name, enabled in sides.items():
        el = tcBorders.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tcBorders.append(el)
        if enabled:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size_eighths))
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


def _cell_text(cell, text, *, size=None, bold=False, align=None, vertical="center"):
    """Заполняет ячейку одним абзацем (без заливки!)."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        _set_run(run, size=size or FONT_SIZE_PT, bold=bold)
    if vertical == "center":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif vertical == "top":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return cell


def _set_table_borders_all(table, size_eighths=4):
    """Все границы — single 0.5pt у всех ячеек."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, size_eighths=size_eighths)


def _set_table_no_borders(table):
    """Убирает все границы у всех ячеек (для таблиц-контейнеров)."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, top=False, bottom=False, left=False, right=False)


def _make_income_table(doc) -> None:
    """Pack 50.8-C-r4: §3 — ОДНА таблица 10 колонок × 15 строк (1 шапка + 14 данных).

    Левые 5 колонок: слоты 0..13. Правые 5 колонок: слоты 14..27.
    Гарантирует одинаковую ширину между половинами (избегаем вложенных таблиц).
    """
    # 10 колонок × (1 + 14) строк
    cols = 10
    table = doc.add_table(rows=1 + NDFL_3_ROWS_PER_COL, cols=cols)
    table.autofit = False

    # Ширины: 5 колонок × 2 половины. Каждая половина = 9.5 cm.
    # Месяц 1.2 | Код 1.4 | Сумма 3.0 | Код вычета 1.4 | Сумма вычета 2.5 = 9.5
    half_widths_cm = [1.2, 1.4, 3.0, 1.4, 2.5]  # = 9.5
    all_widths_cm = half_widths_cm + half_widths_cm  # 19.0
    for i, w in enumerate(all_widths_cm):
        table.columns[i].width = Cm(w)

    # Шапка
    headers_half = ["Месяц", "Код дохода", "Сумма дохода", "Код вычета", "Сумма вычета"]
    headers_all = headers_half + headers_half
    for i, h in enumerate(headers_all):
        _cell_text(
            table.rows[0].cells[i], h,
            size=FONT_SIZE_SMALL_PT,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Строки данных — левая половина (слоты 0..13), правая (14..27)
    for r in range(NDFL_3_ROWS_PER_COL):
        cells = table.rows[r + 1].cells
        # Левая половина: слот r
        slot_left = r
        _cell_text(cells[0], f"{{{{ ndfl_2.rows[{slot_left}].month }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[1], f"{{{{ ndfl_2.rows[{slot_left}].income_code }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[2], f"{{{{ ndfl_2.rows[{slot_left}].income_amount }}}}",
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(cells[3], f"{{{{ ndfl_2.rows[{slot_left}].deduction_code }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[4], f"{{{{ ndfl_2.rows[{slot_left}].deduction_amount }}}}",
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        # Правая половина: слот r + 14
        slot_right = r + NDFL_3_ROWS_PER_COL
        _cell_text(cells[5], f"{{{{ ndfl_2.rows[{slot_right}].month }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[6], f"{{{{ ndfl_2.rows[{slot_right}].income_code }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[7], f"{{{{ ndfl_2.rows[{slot_right}].income_amount }}}}",
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(cells[8], f"{{{{ ndfl_2.rows[{slot_right}].deduction_code }}}}",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(cells[9], f"{{{{ ndfl_2.rows[{slot_right}].deduction_amount }}}}",
                   align=WD_ALIGN_PARAGRAPH.RIGHT)

    _set_table_borders_all(table)


# ============================================================================
# Построение документа
# ============================================================================

def build_template(output_path: Path) -> None:
    doc = Document()

    # ---- Page setup: A4, узкие поля ----
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    # ---- Default font ----
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(FONT_SIZE_PT)

    # ========================================================================
    # ВЕРХНЯЯ ШАПКА (КНД + приложение)
    # ========================================================================
    t_head = doc.add_table(rows=1, cols=2)
    t_head.autofit = False
    t_head.columns[0].width = Cm(9.5)
    t_head.columns[1].width = Cm(9.5)
    _set_table_no_borders(t_head)

    # Левая ячейка — КНД
    left = t_head.rows[0].cells[0]
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Форма по КНД 1175018")
    _set_run(r, size=FONT_SIZE_SMALL_PT, bold=True)

    # Правая ячейка — приложение №4
    right = t_head.rows[0].cells[1]
    right.text = ""
    right_lines = [
        "Приложение № 4",
        "к приказу ФНС России",
        "от 19.09.2023 № ЕД-7-11/649@",
        "(в ред. Приказа ФНС России",
        "от 09.01.2024 № ЕД-7-11/1@)",
    ]
    for i, line in enumerate(right_lines):
        if i == 0:
            p = right.paragraphs[0]
        else:
            p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _set_run(r, size=FONT_SIZE_SMALL_PT)

    # ========================================================================
    # ЗАГОЛОВОК
    # ========================================================================
    _add_paragraph(
        doc, "СПРАВКА О ДОХОДАХ И СУММАХ НАЛОГА ФИЗИЧЕСКОГО ЛИЦА",
        size=FONT_SIZE_TITLE_PT, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=8, space_after_pt=2,
    )
    _add_paragraph(
        doc, "за {{ ndfl_2.year }} год от {{ ndfl_2.issue_date_str }}",
        size=FONT_SIZE_HEADER_PT, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    # ========================================================================
    # §1. СВЕДЕНИЯ О НАЛОГОВОМ АГЕНТЕ
    # ========================================================================
    _add_paragraph(
        doc, "1. Сведения о налоговом агенте",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=4, space_after_pt=2,
    )

    # Таблица 1×4: ОКТМО, Телефон, ИНН, КПП (label сверху, значение снизу)
    t1 = doc.add_table(rows=2, cols=4)
    t1.autofit = False
    col_widths = [Cm(4.8), Cm(4.8), Cm(4.7), Cm(4.7)]
    for i, w in enumerate(col_widths):
        t1.columns[i].width = w

    headers = ["Код по ОКТМО", "Телефон", "ИНН", "КПП"]
    values = [
        "{{ company.oktmo }}",
        "{{ company.phone }}",
        "{{ company.tax_id_primary }}",
        "{{ company.tax_id_secondary }}",
    ]
    for i, h in enumerate(headers):
        _cell_text(t1.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    for i, v in enumerate(values):
        _cell_text(t1.rows[1].cells[i], v, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_table_borders_all(t1)

    # Строка "Налоговый агент: ..."
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Налоговый агент: ")
    _set_run(r1, size=FONT_SIZE_PT)
    r2 = p.add_run("{{ company.full_name_ru }}")
    _set_run(r2, size=FONT_SIZE_PT)

    # Форма реорганизации (пусто)
    _add_paragraph(
        doc,
        "Форма реорганизации (ликвидации) (код): ____    "
        "ИНН/КПП реорганизованной организации: ____________ / ____________",
        size=FONT_SIZE_SMALL_PT,
        space_after_pt=2,
    )

    # ========================================================================
    # §2. СВЕДЕНИЯ О ФИЗИЧЕСКОМ ЛИЦЕ
    # ========================================================================
    _add_paragraph(
        doc, "2. Сведения о физическом лице — получателе дохода",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=6, space_after_pt=2,
    )

    # ИНН
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ИНН в Российской Федерации: ")
    _set_run(r, size=FONT_SIZE_PT)
    r = p.add_run("{{ applicant.inn }}")
    _set_run(r, size=FONT_SIZE_PT)

    # ФИО — таблица 1×3
    t2 = doc.add_table(rows=2, cols=3)
    t2.autofit = False
    for i, w in enumerate([Cm(6.3), Cm(6.3), Cm(6.4)]):
        t2.columns[i].width = w
    for i, h in enumerate(["Фамилия", "Имя", "Отчество*"]):
        _cell_text(t2.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t2.rows[1].cells[0], "{{ applicant.last_name_native }}")
    _cell_text(t2.rows[1].cells[1], "{{ applicant.first_name_native }}")
    _cell_text(t2.rows[1].cells[2], "{{ applicant.middle_name_native }}")
    _set_table_borders_all(t2)

    # Статус / ДР / Гражданство
    t3 = doc.add_table(rows=2, cols=3)
    t3.autofit = False
    for i, w in enumerate([Cm(6.3), Cm(6.3), Cm(6.4)]):
        t3.columns[i].width = w
    for i, h in enumerate(["Статус налогоплательщика", "Дата рождения", "Гражданство (код страны)"]):
        _cell_text(t3.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t3.rows[1].cells[0], "{{ ndfl_2.taxpayer_status }}")
    _cell_text(t3.rows[1].cells[1], "{{ ndfl_2.birth_date_str }}")
    _cell_text(t3.rows[1].cells[2], "{{ ndfl_2.country_code }}")
    _set_table_borders_all(t3)

    # Код документа + Серия и номер
    t4 = doc.add_table(rows=2, cols=2)
    t4.autofit = False
    for i, w in enumerate([Cm(9.5), Cm(9.5)]):
        t4.columns[i].width = w
    for i, h in enumerate(["Код документа, удостоверяющего личность", "Серия и номер документа"]):
        _cell_text(t4.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t4.rows[1].cells[0], "{{ ndfl_2.id_doc_code }}")
    _cell_text(t4.rows[1].cells[1], "{{ ndfl_2.passport_series_number }}")
    _set_table_borders_all(t4)

    # ========================================================================
    # §3. ДОХОДЫ ПО СТАВКЕ 13%
    #     Две таблицы рядом: 1 шапка + 14 строк данных в каждой = 28 слотов
    # ========================================================================
    _add_paragraph(
        doc, "3. Доходы, облагаемые по ставке 13 %",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=6, space_after_pt=2,
    )

    # Pack 50.8-C-r4: §3 — ОДНА таблица 10 колонок (5 слева + 5 справа),
    # 15 строк (1 шапка + 14 данных). Так колонки гарантированно одинаковые
    # между левой и правой половиной (избегаем вложенных таблиц).
    _make_income_table(doc)

    # ========================================================================
    # §4. СТАНДАРТНЫЕ, СОЦИАЛЬНЫЕ И ИМУЩЕСТВЕННЫЕ НАЛОГОВЫЕ ВЫЧЕТЫ
    # ========================================================================
    _add_paragraph(
        doc, "4. Стандартные, социальные и имущественные налоговые вычеты",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=6, space_after_pt=2,
    )

    # 4 пары "Код вычета / Сумма вычета"
    t_ded = doc.add_table(rows=2, cols=8)
    t_ded.autofit = False
    ded_w = Cm(CONTENT_WIDTH_CM / 8)
    for i in range(8):
        t_ded.columns[i].width = ded_w
    ded_headers = ["Код вычета", "Сумма вычета"] * 4
    for i, h in enumerate(ded_headers):
        _cell_text(t_ded.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(8):
        # MVP: вычетов нет → "0,00" в полях "Сумма вычета", пусто в "Код"
        _cell_text(t_ded.rows[1].cells[i], "0,00" if i % 2 == 1 else "",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_table_borders_all(t_ded)

    # ========================================================================
    # §5. ОБЩАЯ СУММА ДОХОДА И СУММА НАЛОГА
    # ========================================================================
    _add_paragraph(
        doc, "5. Общая сумма дохода и сумма налога",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=6, space_after_pt=2,
    )

    rows_section5 = [
        ("Общая сумма дохода", "{{ ndfl_2.total_income }}",
         "Налоговая база", "{{ ndfl_2.tax_base }}"),
        ("Сумма налога исчисленная", "{{ ndfl_2.tax_calculated }}",
         "Сумма фиксированных авансовых платежей", ""),
        ("Сумма налога на прибыль организаций, подлежащая зачету", "",
         "Сумма налога, исчисленная и уплаченная в иностранном государстве", ""),
        ("Сумма налога удержанная", "{{ ndfl_2.tax_withheld }}",
         "Сумма налога, излишне удержанная налоговым агентом", ""),
    ]
    t5 = doc.add_table(rows=len(rows_section5), cols=4)
    t5.autofit = False
    # Pack 50.8-C-r3: равные пропорции label/value по эталону ФНС
    # 5.5 + 4.0 + 5.5 + 4.0 = 19.0 cm (вся ширина страницы)
    for i, w in enumerate([Cm(5.5), Cm(4.0), Cm(5.5), Cm(4.0)]):
        t5.columns[i].width = w
    for ri, row_data in enumerate(rows_section5):
        for ci, val in enumerate(row_data):
            cell = t5.rows[ri].cells[ci]
            if ci % 2 == 0:
                _cell_text(cell, val, size=FONT_SIZE_SMALL_PT,
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                # Значения — НЕ жирные (как в эталоне ФНС)
                _cell_text(cell, val, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_table_borders_all(t5)

    # ========================================================================
    # §6. НЕУДЕРЖАННЫЙ НАЛОГ
    # ========================================================================
    _add_paragraph(
        doc, "6. Сумма дохода, с которого не удержан налог налоговым агентом, и сумма неудержанного налога",
        size=FONT_SIZE_PT, bold=True,
        space_before_pt=6, space_after_pt=2,
    )

    t6 = doc.add_table(rows=2, cols=2)
    t6.autofit = False
    for i, w in enumerate([Cm(13.0), Cm(6.0)]):
        t6.columns[i].width = w
    _cell_text(t6.rows[0].cells[0],
               "Сумма дохода, с которого не удержан налог налоговым агентом",
               size=FONT_SIZE_SMALL_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t6.rows[0].cells[1], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(t6.rows[1].cells[0], "Сумма неудержанного налога",
               size=FONT_SIZE_SMALL_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t6.rows[1].cells[1], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_table_borders_all(t6)

    # ========================================================================
    # ПОДПИСЬ
    # ========================================================================
    _add_paragraph(doc, "", space_before_pt=12)

    # ФИО + место для подписи
    t_sig = doc.add_table(rows=1, cols=2)
    t_sig.autofit = False
    for i, w in enumerate([Cm(11.0), Cm(8.0)]):
        t_sig.columns[i].width = w
    _set_table_no_borders(t_sig)
    _cell_text(t_sig.rows[0].cells[0], "{{ company.director_full_name_ru }}",
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="top")
    _cell_text(t_sig.rows[0].cells[1], "_______________________  (подпись)",
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="top")

    t_meta = doc.add_table(rows=1, cols=2)
    t_meta.autofit = False
    for i, w in enumerate([Cm(11.0), Cm(8.0)]):
        t_meta.columns[i].width = w
    _set_table_no_borders(t_meta)
    _cell_text(t_meta.rows[0].cells[0], "Налоговый агент (Ф. И. О.)*",
               size=FONT_SIZE_SMALL_PT, align=WD_ALIGN_PARAGRAPH.LEFT,
               vertical="top")
    _cell_text(t_meta.rows[0].cells[1], "", align=WD_ALIGN_PARAGRAPH.LEFT,
               vertical="top")

    _add_paragraph(
        doc, "*Отчество указывается при наличии.",
        size=FONT_SIZE_SMALL_PT, italic=True,
        space_before_pt=10,
    )

    # ---- Save ----
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ Saved: {output_path}  ({output_path.stat().st_size / 1024:.1f} KB)")


# ============================================================================
# Main
# ============================================================================

def main() -> int:
    # Резолв пути: скрипт лежит в корне репо, запускается оттуда.
    script_dir = Path(__file__).resolve().parent

    candidates = [script_dir]
    if (script_dir.parent / "templates" / "docx").exists():
        candidates.append(script_dir.parent)
    if Path("D:/VISA/visa_kit/templates/docx").exists():
        candidates.append(Path("D:/VISA/visa_kit"))

    root = None
    for c in candidates:
        if (c / "templates" / "docx").exists():
            root = c
            break

    if root is None:
        print("❌ Не найдена директория templates/docx/. Запускай из корня репо.")
        return 1

    output = root / "templates" / "docx" / "ndfl_2_template.docx"
    print(f"Building template → {output}")

    if output.exists():
        import shutil as _sh
        backup = output.with_suffix(
            output.suffix + f".bak_pre_pack50_8_C_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        _sh.copy2(output, backup)
        print(f"   backup: {backup.name}")

    build_template(output)

    if output.exists() and output.stat().st_size > 0:
        print(f"✅ OK — шаблон создан, размер {output.stat().st_size} байт")
        return 0
    print("❌ Шаблон не создан")
    return 1


if __name__ == "__main__":
    sys.exit(main())
