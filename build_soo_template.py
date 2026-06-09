r"""
Pack 50.12-C — Генератор soo_template.docx (Свидетельство об отъезде).

Структура (по эталону Орлов + блок 4 Регины):
  ШАПКА: договор РФ-Испания / СВИДЕТЕЛЬСТВО ОБ ОТЪЕЗДЕ / статьи / Дата / №
  РАЗДЕЛ 1 (застрахованное лицо): плашка + 1.1/1.2/1.3 (параграфы)
  РАЗДЕЛ 2 (работодатель): плашка + таблица 4×N
  РАЗДЕЛ 3 (применимое законодательство): плашка + 3.1 параграф
  РАЗДЕЛ 4 (СФР): плашка + таблица 3×1 (захардкожен, эталон Регины)

Все данные через плейсхолдеры {{ soo.* }}.

Запуск:
    python build_soo_template.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "soo_template.docx"

PAGE_W, PAGE_H = 21.0, 29.7
MARGIN_LR, MARGIN_T, MARGIN_B = 1.27, 1.2, 1.0
FONT = "Times New Roman"
SZ = 8       # основной размер (как эталон — 8pt)
SZ_DATA = 8  # данные тоже 8


def _set_run(run, *, size=SZ, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{k}"), FONT)


def _add_p(doc, runs, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    """runs = list of (text, bold) или строка."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        r = p.add_run(text)
        _set_run(r, size=SZ, bold=bold)
    return p


def _shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _section_header(doc, num_text, title_text):
    """Плашка-заголовок раздела: серый фон, номер + заголовок жирным."""
    # Таблица 1×1 с серой заливкой (имитация плашки)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade_cell(cell, "D9D9D9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{num_text}  {title_text}")
    _set_run(r, size=SZ, bold=True)
    return t


def _cell_runs(cell, runs, *, align=WD_ALIGN_PARAGRAPH.LEFT, valign="center"):
    cell.text = ""
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP, "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[valign]
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if isinstance(runs, str):
        runs = [(runs, False)]
    first = True
    for text, bold in runs:
        # поддержка многострочности через \n
        lines = text.split("\n")
        for li, line in enumerate(lines):
            if li > 0:
                r = p.add_run()
                r.add_break()
            r = p.add_run(line)
            _set_run(r, size=SZ, bold=bold)


def build_template() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(PAGE_W)
    s.page_height = Cm(PAGE_H)
    s.left_margin = Cm(MARGIN_LR)
    s.right_margin = Cm(MARGIN_LR)
    s.top_margin = Cm(MARGIN_T)
    s.bottom_margin = Cm(MARGIN_B)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SZ)

    # ============ ШАПКА ============
    _add_p(doc, [(
        "ДОГОВОР МЕЖДУ РОССИЙСКОЙ ФЕДЕРАЦИЕЙ И КОРОЛЕВСТВОМ ИСПАНИЯ О СОЦИАЛЬНОМ "
        "ОБЕСПЕЧЕНИИ ОТ 11 АПРЕЛЯ 1994 г. / АДМИНИСТРАТИВНОЕ СОГЛАШЕНИЕ О ПРИМЕНЕНИИ "
        "ДОГОВОРА МЕЖДУ РОССИЙСКОЙ ФЕДЕРАЦИЕЙ И КОРОЛЕВСТВОМ ИСПАНИЯ О СОЦИАЛЬНОМ "
        "ОБЕСПЕЧЕНИИ ОТ 12 МАЯ 1995 г.", True)],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    _add_p(doc, [("СВИДЕТЕЛЬСТВО ОБ ОТЪЕЗДЕ", True)],
           align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)

    _add_p(doc, [("Статья 7 Договора, статья 3 Административного соглашения", False)],
           align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    _add_p(doc, [("Дата ", True), ("{{ soo.date_long }}", True), ("г.", True)],
           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    _add_p(doc, [("№  ", True), ("{{ soo.number }}", True)],
           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    # ============ РАЗДЕЛ 1 — ЗАСТРАХОВАННОЕ ЛИЦО ============
    _section_header(doc, "1.",
                    "ЗАСТРАХОВАННОЕ ЛИЦО, ОСУЩЕСТВЛЯЮЩЕЕ ТРУДОВУЮ ИЛИ ИНУЮ ДЕЯТЕЛЬНОСТЬ")

    # 1.1 Фамилия/Имя/Отчество (метки + значения)
    _add_p(doc, [
        ("1.1. Фамилия ", False), ("{{ soo.last_name }}", False),
        ("   Имя ", False), ("{{ soo.first_name }}", False),
        ("   Отчество ", False), ("{{ soo.middle_name }}", False),
    ], space_before=2, space_after=1)

    _add_p(doc, [
        ("Гражданство ", False), ("{{ soo.nationality_ru }}", False),
    ], space_after=1)

    _add_p(doc, [
        ("Документ, удостоверяющий личность: паспорт ", False),
        ("{{ soo.passport_series }}", False), (" № ", False),
        ("{{ soo.passport_number_only }}", False), (", выдан ", False),
        ("{{ soo.passport_issue_date }}", False), (" ", False),
        ("{{ soo.passport_issuer }}", False),
        (", код подразделения ", False), ("{{ soo.division_code }}", False),
    ], space_after=1)

    _add_p(doc, [
        ("Номер загранпаспорта ", False), ("{{ soo.foreign_passport_number }}", False),
        (" выдан ", False), ("{{ soo.foreign_passport_issue_date }}", False),
        (" ", False), ("{{ soo.foreign_passport_issuer }}", False),
    ], space_after=1)

    # 1.2 Дата рождения + СНИЛС
    _add_p(doc, [
        ("1.2. Дата рождения ", False), ("{{ soo.birth_date_long }}", False),
        ("   Страховой номер индивидуального лицевого счёта в Российской Федерации "
         "(СНИЛС) ", False), ("{{ soo.snils }}", False),
    ], space_before=2, space_after=1)

    # 1.3 Адрес/телефон/факс/email
    _add_p(doc, [
        ("1.3. Адрес ", False), ("{{ soo.home_address }}", False),
    ], space_before=2, space_after=1)
    _add_p(doc, [
        ("Телефон ", False), ("{{ soo.phone }}", False),
        ("   Факс ", False),
        ("   Эл. почта: ", False), ("{{ soo.email }}", False),
    ], space_after=4)

    # ============ РАЗДЕЛ 2 — РАБОТОДАТЕЛЬ ============
    _section_header(doc, "2.",
                    "СВЕДЕНИЯ О РАБОТОДАТЕЛЕ ИЛИ ИНОЙ ДЕЯТЕЛЬНОСТИ В ГОСУДАРСТВЕ, "
                    "ПРИМЕНЯЕМОГО ЗАКОНОДАТЕЛЬСТВА")

    t2 = doc.add_table(rows=4, cols=2)
    t2.style = "Table Grid"
    t2.autofit = False
    # ширины: узкая (номер) + широкая (контент)
    widths = [Cm(1.0), Cm(17.46)]
    for row in t2.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = widths[ci]
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(widths[ci].cm * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    _cell_runs(t2.rows[0].cells[0], "2.1.", valign="top")
    _cell_runs(t2.rows[0].cells[1], [("Наименование", True)], valign="top")
    _cell_runs(t2.rows[1].cells[0], "", valign="top")
    _cell_runs(t2.rows[1].cells[1], [("{{ soo.company_full_name }}", False)], valign="top")
    _cell_runs(t2.rows[2].cells[0], "2.2.", valign="top")
    _cell_runs(t2.rows[2].cells[1], [
        ("Адрес: ", True), ("{{ soo.company_legal_address }}", False),
        ("\nАдрес удалённой работы в Испании: ", True),
        ("{{ soo.spain_address }}", False),
    ], valign="top")
    _cell_runs(t2.rows[3].cells[0], "", valign="top")
    _cell_runs(t2.rows[3].cells[1], [
        ("Телефон: ", True), ("{{ soo.company_phone }}", False),
        ("   Факс:    Эл. почта: ", True), ("{{ soo.company_email }}", False),
    ], valign="top")

    # ============ РАЗДЕЛ 3 — ПРИМЕНИМОЕ ЗАКОНОДАТЕЛЬСТВО ============
    _add_p(doc, "", space_before=2)
    _section_header(doc, "3.", "ПРИМЕНИМОЕ ЗАКОНОДАТЕЛЬСТВО")

    _add_p(doc, [
        ("3.1. В отношении застрахованного лица продолжает применяться "
         "законодательство Российской Федерации: с ", False),
        ("{{ soo.period_start }}", True),
        (" по ", False),
        ("{{ soo.period_end }}", True),
        (". Согласно ст. 7 Договора между Российской Федерацией и Королевством "
         "Испании о социальном обеспечении от 11.04.1994 г. и ст. 3 "
         "Административного соглашения о применении Договора между РФ и Королевством "
         "Испания о социальном обеспечении от 12.05.1995 подтверждаем, что с "
         "применимым международным стандартом координации социального обеспечения, "
         "подписанным Королевством Испании, ", False),
        ("{{ soo.full_name }}", False),
        (" будет являться застрахованным лицом в системе обязательного социального "
         "страхования Российской Федерации.", False),
    ], align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=2, space_after=2)

    _add_p(doc, [
        ("Основание: ПРИКАЗ № ", False), ("{{ soo.order_number }}", False),
        (" от ", False), ("{{ soo.order_date }}", False),
        ("г. о направлении работника в командировку. Трудовой договор № ", False),
        ("{{ soo.contract_number }}", False), (" от ", False),
        ("{{ soo.contract_date }}", False),
        ("г. с возможностью дистанционной работы.", False),
    ], align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    # ============ РАЗДЕЛ 4 — КОМПЕТЕНТНОЕ УЧРЕЖДЕНИЕ (СФР) ============
    # Захардкожен по эталону Регины
    _section_header(doc, "4.",
                    "ДАННЫЕ КОМПЕТЕНТНОГО УЧРЕЖДЕНИЯ, ЗАПОЛНИВШЕГО СВИДЕТЕЛЬСТВО ОБ ОТЪЕЗДЕ")

    t4 = doc.add_table(rows=3, cols=1)
    t4.style = "Table Grid"
    t4.autofit = False
    w4 = Cm(18.46)
    for row in t4.rows:
        cell = row.cells[0]
        cell.width = w4
        tcPr = cell._element.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(w4.cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    # SZ 10pt в блоке 4 (как эталон Регины)
    def _cell4(cell, runs):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        first = True
        for text, bold in runs:
            lines = text.split("\n")
            for li, line in enumerate(lines):
                if li > 0 or not first:
                    if li > 0:
                        r = p.add_run(); r.add_break()
                r = p.add_run(line)
                r.font.name = FONT
                r.font.size = Pt(10)
                r.font.bold = bold
                rPr = r._element.get_or_add_rPr()
                rF = rPr.find(qn("w:rFonts"))
                if rF is None:
                    rF = OxmlElement("w:rFonts"); rPr.append(rF)
                for k in ("ascii","hAnsi","cs"): rF.set(qn(f"w:{k}"), FONT)
            first = False

    _cell4(t4.rows[0].cells[0], [("Наименование", True)])
    _cell4(t4.rows[1].cells[0], [
        ("Отделение Фонда пенсионного и социального страхования Российской Федерации "
         "по г. Москве и Московской области", False)])
    _cell4(t4.rows[2].cells[0], [
        ("Адрес 107078, Россия, г. Москва, Орликов пер., д. 3, корп. А", True),
        ("\n………………………………………………………………………………………………………………………………………..", True),
        ("\nТелефон +74959862611", True),
        ("\n                                                                Дата ……………………………………….", False),
        ("\nПечать                                                  Подпись …………………………………...", True),
    ])

    return doc


def main() -> int:
    print(f"Building soo_template → {OUTPUT_PATH}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_12_C_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")
    doc = build_template()
    doc.save(str(OUTPUT_PATH))
    print(f"✅ Saved: {OUTPUT_PATH}  ({OUTPUT_PATH.stat().st_size/1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
