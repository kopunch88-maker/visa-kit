r"""
Pack 50.10-C-r2 — Генератор payslip_template.docx, точно по эталону.

Изменения r2 относительно r1:
  - Шрифт Times New Roman (как в эталоне), не Arial
  - "К выплате" → выравнивание RIGHT на отдельной строке (как эталон)
  - "Долг предприятия..." перенесён ПОД таблицу (как эталон)
  - В таблице Row 2 ("Начислено:" / "Удержано:") добавлены СУММЫ с табстопом справа

Структура по эталону "Расчетный_листок_за_<MM>_<YYYY>.docx":
  - Portrait A4 (21×29.7 cm)
  - Шапка из параграфов:
      P1: "РАСЧЕТНЫЙ ЛИСТОК ЗА {{ payslip.month_title_upper }}"  (10pt center)
      P2: пусто
      P3: "ФИО: {{ payslip.applicant_full_name }}"  (11pt)
      P4: "К выплате:        {{ payslip.payout_fmt }}"  (10pt)
      P5: "{{ payslip.company_full_name_upper }}"  (11pt bold)
      P6: "Должность:\t{{ payslip.position_title }}"  (11pt bold)
      P7: "Подразделение: {{ payslip.department }}\tОклад (тариф): {{ payslip.salary_oklad_raw }}"
      P8: "Долг предприятия на начало 0,00 Долг предприятия на конец 0,00 Общий
           облагаемый доход за {{ payslip.year }} год нарастающим итогом: {{ payslip.yearly_total_fmt }}"
  - Таблица 5×9
  - P: "Бухгалтер                                   {{ payslip.accountant_short }}"

Запуск:
    python build_payslip_template.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "payslip_template.docx"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LR_CM = 1.0
MARGIN_TB_CM = 1.5
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_LR_CM  # 19.0

FONT_FAMILY = "Times New Roman"
FONT_SIZE_BASE = 10
FONT_SIZE_BIG = 11

# Ширины колонок таблицы (из эталона, нормализованы к 19.0 cm)
RAW_WIDTHS = [3.47, 1.98, 1.09, 1.09, 1.84, 1.97, 3.82, 1.63, 2.18]
TOTAL_RAW = sum(RAW_WIDTHS)  # 19.07
COL_WIDTHS = [w * USABLE_WIDTH_CM / TOTAL_RAW for w in RAW_WIDTHS]


def _set_run(run, *, size=FONT_SIZE_BASE, bold=False, italic=False):
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{k}"), FONT_FAMILY)


def _add_paragraph(doc, text="", *, size=FONT_SIZE_BASE, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before_pt=0, space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if text:
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, italic=italic)
    return p


def _cell_text(cell, text, *, size=FONT_SIZE_BASE, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center"):
    cell.text = ""
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[vertical]
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold, italic=italic)


def _set_cell_borders(cell, *, top=True, bottom=True, left=True, right=True, size_val="4"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    sides = [("top", top), ("left", left), ("bottom", bottom), ("right", right)]
    for side_name, enabled in sides:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        border = OxmlElement(f"w:{side_name}")
        if enabled:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size_val)
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "nil")
        tcBorders.append(border)


def _set_table_borders_tbl(table):
    """tblBorders на уровне таблицы — корректно работает с merged."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_fixed_layout(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _apply_widths(table, widths_cm: list):
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _merge_horizontal(row, start_col: int, end_col: int):
    cells = row.cells
    merged = cells[start_col]
    for ci in range(start_col + 1, end_col + 1):
        merged = merged.merge(cells[ci])
    return merged


def _set_vmerge(cell, val):
    tcPr = cell._element.get_or_add_tcPr()
    vMerge = tcPr.find(qn("w:vMerge"))
    if vMerge is None:
        vMerge = OxmlElement("w:vMerge")
        tcPr.append(vMerge)
    if val == "restart":
        vMerge.set(qn("w:val"), "restart")
    elif val is None:
        if qn("w:val") in vMerge.attrib:
            del vMerge.attrib[qn("w:val")]


def build_template() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)

    # Default font — Times New Roman
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_SIZE_BASE)

    # P1: Заголовок (10pt center)
    _add_paragraph(
        doc,
        "РАСЧЕТНЫЙ ЛИСТОК ЗА {{ payslip.month_title_upper }}",
        size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    # P2: пусто
    _add_paragraph(doc, "", size=FONT_SIZE_BASE)

    # P3: ФИО (10pt) — выравнивание влево
    _add_paragraph(
        doc,
        "ФИО:   {{ payslip.applicant_full_name }}",
        size=FONT_SIZE_BASE,
    )

    # P4: "К выплате" — RIGHT alignment, 10pt bold (как эталон)
    p_k_vyplate = doc.add_paragraph()
    p_k_vyplate.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_k_vyplate.paragraph_format.space_before = Pt(0)
    p_k_vyplate.paragraph_format.space_after = Pt(2)
    r1 = p_k_vyplate.add_run("К выплате:           ")
    _set_run(r1, size=FONT_SIZE_BASE, bold=True)
    r2 = p_k_vyplate.add_run("{{ payslip.payout_fmt }}")
    _set_run(r2, size=FONT_SIZE_BASE, bold=True)

    # P5: Компания (10pt bold)
    _add_paragraph(
        doc,
        "{{ payslip.company_full_name_upper }}",
        size=FONT_SIZE_BASE, bold=True,
        space_before_pt=4,
    )

    # P6: Должность (10pt bold)
    _add_paragraph(
        doc,
        "Должность:    {{ payslip.position_title }}",
        size=FONT_SIZE_BASE, bold=True,
    )

    # P7: Подразделение + Оклад (10pt) - с табом справа для "Оклад (тариф)"
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p7.paragraph_format.space_before = Pt(2)
    p7.paragraph_format.space_after = Pt(6)
    from docx.enum.text import WD_TAB_ALIGNMENT as _TAB
    p7.paragraph_format.tab_stops.add_tab_stop(Cm(10.5), _TAB.LEFT)
    r = p7.add_run("Подразделение:   {{ payslip.department }}")
    _set_run(r, size=FONT_SIZE_BASE)
    r = p7.add_run("\t")
    _set_run(r, size=FONT_SIZE_BASE)
    r = p7.add_run("Оклад (тариф):   {{ payslip.salary_oklad_raw }}")
    _set_run(r, size=FONT_SIZE_BASE)

    # === ТАБЛИЦА 5×9 ===
    t = doc.add_table(rows=5, cols=9)
    t.autofit = False
    _set_table_fixed_layout(t)
    _apply_widths(t, COL_WIDTHS)

    # Row 0: шапка
    _cell_text(t.rows[0].cells[0], "Вид",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[1], "Период",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    rabochie_cell = _merge_horizontal(t.rows[0], 2, 3)
    _cell_text(rabochie_cell, "Рабочие",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[3], "Оплачено",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[4], "Сумма",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[5], "Вид",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[6], "Период",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[0].cells[7], "Сумма",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    for ci in [0, 1, 3, 4, 5, 6, 7]:
        _set_vmerge(t.rows[0].cells[ci], "restart")

    # Row 1: Дни / Часы под "Рабочие"
    _cell_text(t.rows[1].cells[2], "Дни",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[1].cells[3], "Часы",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    for ci in [0, 1, 4, 5, 6, 7, 8]:
        _set_vmerge(t.rows[1].cells[ci], None)

    # Row 2: Начислено: ___ | Удержано: ___
    # PACK 50.10-C-r2: добавлены СУММЫ в эту строку
    # ВАЖНО: сначала мержим ОБА диапазона, ПОТОМ заполняем (чтобы индексы не сбивались).
    # Логические колонки: 0-5 = Начислено, 6-8 = Удержано.
    row2 = t.rows[2]
    # Merge правую группу ПЕРВОЙ (cols 6-8), т.к. левый merge сдвинет индексы.
    uderzhano = _merge_horizontal(row2, 6, 8)
    # Merge левую группу (cols 0-5)
    nachisleno = _merge_horizontal(row2, 0, 5)

    # Заполняем "Начислено:" слева + сумма справа через табстоп
    nachisleno.text = ""
    nachisleno.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_n = nachisleno.paragraphs[0]
    p_n.alignment = WD_ALIGN_PARAGRAPH.LEFT
    nachisleno_width_cm = sum(COL_WIDTHS[:6])
    p_n.paragraph_format.tab_stops.add_tab_stop(Cm(nachisleno_width_cm - 0.3), _TAB.RIGHT)
    r = p_n.add_run("Начислено:")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_n.add_run("\t")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_n.add_run("{{ payslip.salary_amount_fmt }}")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)

    # Заполняем "Удержано:"
    uderzhano.text = ""
    uderzhano.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_u = uderzhano.paragraphs[0]
    p_u.alignment = WD_ALIGN_PARAGRAPH.LEFT
    uderzhano_width_cm = sum(COL_WIDTHS[6:9])
    p_u.paragraph_format.tab_stops.add_tab_stop(Cm(uderzhano_width_cm - 0.3), _TAB.RIGHT)
    r = p_u.add_run("Удержано:")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_u.add_run("\t")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_u.add_run("{{ payslip.ndfl_fmt }}")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)

    # Row 3: данные
    _cell_text(t.rows[3].cells[0], "Оплата по окладу",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    _cell_text(t.rows[3].cells[1], "{{ payslip.month_short }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[2], "{{ payslip.working_days }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[3], "{{ payslip.working_hours }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[4], "{{ payslip.days_paid_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[5], "{{ payslip.salary_amount_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.RIGHT, vertical="center")
    _cell_text(t.rows[3].cells[6], "НДФЛ",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center")
    _cell_text(t.rows[3].cells[7], "{{ payslip.month_short_dotted }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.CENTER, vertical="center")
    _cell_text(t.rows[3].cells[8], "{{ payslip.ndfl_fmt }}",
               size=FONT_SIZE_BASE, align=WD_ALIGN_PARAGRAPH.RIGHT, vertical="center")

    # Row 4: Выплачено (только cols 6-8)
    vyplacheno = _merge_horizontal(t.rows[4], 6, 8)
    vyplacheno.text = ""
    vyplacheno.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_v = vyplacheno.paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_v.paragraph_format.tab_stops.add_tab_stop(Cm(uderzhano_width_cm - 0.3), _TAB.RIGHT)
    r = p_v.add_run("Выплачено:")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_v.add_run("\t")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    r = p_v.add_run("{{ payslip.payout_fmt }}")
    _set_run(r, size=FONT_SIZE_BASE, bold=True)
    for ci in range(6):
        _cell_text(t.rows[4].cells[ci], "", size=FONT_SIZE_BASE)

    _set_table_borders_tbl(t)

    # === ПОСЛЕ ТАБЛИЦЫ ===
    # Pack 50.10-C-r2: "Долг предприятия..." перенесён сюда (под таблицу)
    p_dolg = doc.add_paragraph()
    p_dolg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_dolg.paragraph_format.space_before = Pt(2)
    p_dolg.paragraph_format.space_after = Pt(2)
    # Tab stops для выравнивания "0,00" на правом крае
    p_dolg.paragraph_format.tab_stops.add_tab_stop(Cm(9.0), _TAB.LEFT)  # после "Долг на начало"
    p_dolg.paragraph_format.tab_stops.add_tab_stop(Cm(USABLE_WIDTH_CM - 0.3), _TAB.RIGHT)  # после "0,00 конец"
    r = p_dolg.add_run("Долг предприятия на начало")
    _set_run(r, size=FONT_SIZE_BASE)
    r = p_dolg.add_run("\t0,00   Долг предприятия на конец\t0,00")
    _set_run(r, size=FONT_SIZE_BASE)

    # Параграф с нарастающим итогом — отдельная строка
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_total.paragraph_format.space_before = Pt(0)
    p_total.paragraph_format.space_after = Pt(8)
    p_total.paragraph_format.tab_stops.add_tab_stop(Cm(USABLE_WIDTH_CM - 0.3), _TAB.RIGHT)
    r = p_total.add_run("Общий облагаемый доход за {{ payslip.year }} год нарастающим итогом:")
    _set_run(r, size=FONT_SIZE_BASE)
    r = p_total.add_run("\t{{ payslip.yearly_total_fmt }}")
    _set_run(r, size=FONT_SIZE_BASE)

    # Подпись внизу
    _add_paragraph(doc, "", space_before_pt=12)
    _add_paragraph(doc, "", space_before_pt=4)
    _add_paragraph(
        doc,
        "Бухгалтер                                              {{ payslip.accountant_short }}",
        size=FONT_SIZE_BASE,
    )

    return doc


def main() -> int:
    print(f"Building payslip template → {OUTPUT_PATH}")
    print(f"Column widths: {[round(w, 2) for w in COL_WIDTHS]}")
    print(f"  Sum: {round(sum(COL_WIDTHS), 2)} cm")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_10_C_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")

    doc = build_template()
    doc.save(str(OUTPUT_PATH))

    size = OUTPUT_PATH.stat().st_size
    print(f"✅ Saved: {OUTPUT_PATH}  ({size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
