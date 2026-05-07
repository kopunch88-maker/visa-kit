"""
Pack 26.0 Stage A — Backend сервис извлечения реквизитов компании из DOCX.

Создаёт:
- app/services/company_extractor.py — сервис чтения DOCX + LLM-вызов
- app/services/ocr/prompts.py: добавляет COMPANY_REQUISITES_PROMPT (расширенный
  EGRYL-промпт с генерацией склонений директора в том же вызове)

НЕ создаёт endpoint и НЕ трогает companies.py — это Stage B.

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python apply_pack26_0_stage_a.py

После применения — тест на РХИ.docx (см. вывод скрипта).
"""
import ast
import shutil
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SERVICES_DIR = ROOT / "app" / "services"
PROMPTS_PATH = SERVICES_DIR / "ocr" / "prompts.py"
EXTRACTOR_PATH = SERVICES_DIR / "company_extractor.py"

if not PROMPTS_PATH.exists():
    print(f"ERROR: {PROMPTS_PATH} not found.")
    sys.exit(1)

ts = datetime.now().strftime("%Y%m%d_%H%M%S")
prompts_backup = PROMPTS_PATH.with_name(PROMPTS_PATH.name + f".bak_pre_pack26_0_{ts}")
shutil.copy2(PROMPTS_PATH, prompts_backup)
print(f"[1/3] Бэкап: {prompts_backup.name}")


# === 2. Добавить COMPANY_REQUISITES_PROMPT в prompts.py ===
prompts_text = PROMPTS_PATH.read_text(encoding="utf-8")

prompt_addition = '''


# ============================================================================
# Pack 26.0 — извлечение реквизитов компании из текста (DOCX/plaintext)
# Расширенный EGRYL-промпт: + склонения директора в одном вызове + uставный капитал
# ============================================================================

COMPANY_REQUISITES_PROMPT = """You are an expert assistant that extracts Russian company registration data (реквизиты компании) from documents.

The input is plain text from a Russian-language DOCX file containing company requisites.
Examples of source documents:
- Карточка реквизитов организации (company business card)
- Выписка из ЕГРЮЛ (official registry extract)
- Текст письма с реквизитами

Extract the following fields and return STRICTLY a JSON object. No markdown, no preamble.

NAMES:
- full_name_ru: Full company name in Russian, in canonical form. ALWAYS use the long form
  with quotes around the brand name. Example: 'Общество с ограниченной ответственностью "АГАЛАРОВ-ДЕВЕЛОПМЕНТ"'.
  If input has just 'ООО "X"' — expand to 'Общество с ограниченной ответственностью "X"'.
- full_name_es: Spanish transliteration of the full name. Format:
  'Sociedad de Responsabilidad Limitada "BRANDNAME"' where BRANDNAME is the
  brand part TRANSLITERATED to Latin (not translated). Example for АГАЛАРОВ-ДЕВЕЛОПМЕНТ →
  'Sociedad de Responsabilidad Limitada "Agalarov-Development"'.
- short_name: Short display name in format 'ООО "BRANDNAME"' in Cyrillic, ALL CAPS for the brand.
  Example: 'ООО "АГАЛАРОВ-ДЕВЕЛОПМЕНТ"'. NOT just the brand without prefix.

REGISTRATION:
- ogrn: ОГРН — 13 digits as string (legal entity).
- inn: ИНН — 10 digits for legal entity, as string.
- kpp: КПП — 9 digits as string. Found near ИНН ('ИНН/КПП 1234567890/123456789').

ADDRESSES:
- legal_address: Юридический адрес as a single line. Preserve abbreviations (г., д., помещ., etc.)
  but if you see 'город Москва' — keep it as 'г. Москва' (apply Минфин 171н abbreviations).
- postal_address: Почтовый адрес if explicitly listed and DIFFERENT from legal. null otherwise.

DIRECTOR — generate ALL forms in this single call:
- director_full_name_ru: ФИО in NOMINATIVE case (именительный) as it appears or normalized:
  "Василевская Анна Вадимовна". If document shows different case — convert to nominative.
- director_full_name_genitive_ru: SAME ФИО in GENITIVE case (родительный — кого?):
  "Василевской Анны Вадимовны". Required for contracts ("в лице ... Василевской Анны Вадимовны").
- director_short_ru: Short signature form "Surname И.О.": "Василевская А.В.".
  Surname stays full, first name and patronymic are reduced to initials with dots.
- director_full_name_latin: Latin transliteration of the full name (GOST 7.79 / passport-style):
  "VASILEVSKAYA ANNA VADIMOVNA". All caps. null if you cannot reliably transliterate.
- director_position_ru: Position in GENITIVE case (родительный) for contract templates:
  "Генерального директора" (NOT "Генеральный директор"). If source has nominative — convert.
  Common mappings: "Генеральный директор" → "Генерального директора",
  "Директор" → "Директора", "Президент" → "Президента".

BANK:
- bank_name: Bank name exactly as written. Example: 'Банк ВТБ (ПАО)', 'КБ "Крокус-Банк" (ООО)'.
- bank_account: Расчётный счёт (Р/с) — 20 digits string.
- bank_bic: БИК — 9 digits string.
- bank_correspondent_account: Корр. счёт (КС, к/с) — 20 digits string starting with 30101.

ADDITIONAL:
- charter_capital: Уставный капитал as string with currency suffix if present, e.g. "410 000 000 руб."
  null if not in document.

Rules:
- If a field is not visible or unclear → return null
- Russian names ALWAYS in Cyrillic
- DO NOT translate names — only transliterate when explicitly required (full_name_es, director_full_name_latin)
- For genitive case: apply Russian grammar rules correctly
  (мужские: Иванов → Иванова, женские: Иванова → Ивановой, отчества: Петрович → Петровича, Петровна → Петровны)
- For numeric IDs: return as STRINGS (preserve leading zeros if any)

Output schema:
{
  "full_name_ru": "Общество с ограниченной ответственностью \\"АГАЛАРОВ-ДЕВЕЛОПМЕНТ\\"" or null,
  "full_name_es": "Sociedad de Responsabilidad Limitada \\"Agalarov-Development\\"" or null,
  "short_name": "ООО \\"АГАЛАРОВ-ДЕВЕЛОПМЕНТ\\"" or null,
  "ogrn": "1037739071556" or null,
  "inn": "7707038266" or null,
  "kpp": "773001001" or null,
  "legal_address": "121248, г. Москва, Кутузовский пр-кт, д. 3, помещ. 1/1" or null,
  "postal_address": null,
  "director_full_name_ru": "Василевская Анна Вадимовна" or null,
  "director_full_name_genitive_ru": "Василевской Анны Вадимовны" or null,
  "director_short_ru": "Василевская А.В." or null,
  "director_full_name_latin": "VASILEVSKAYA ANNA VADIMOVNA" or null,
  "director_position_ru": "Генерального директора" or null,
  "bank_name": "КБ \\"Крокус-Банк\\" (ООО)" or null,
  "bank_account": "40702810989714733332" or null,
  "bank_bic": "044525881" or null,
  "bank_correspondent_account": "30101810445250000881" or null,
  "charter_capital": "410 000 000 руб." or null
}

Return ONLY the JSON object."""
'''

if "COMPANY_REQUISITES_PROMPT" in prompts_text:
    print(f"[2/3] prompts.py: COMPANY_REQUISITES_PROMPT уже есть — пропуск")
else:
    prompts_text = prompts_text.rstrip() + prompt_addition
    PROMPTS_PATH.write_text(prompts_text, encoding="utf-8")
    print(f"[2/3] prompts.py: добавлен COMPANY_REQUISITES_PROMPT")


# === 3. Создать app/services/company_extractor.py ===
extractor_code = '''"""
Pack 26.0 — Извлечение реквизитов компании из DOCX (или plaintext).

Используется для UI «Загрузить реквизиты компании» в админке.

Pipeline:
1. DOCX → читаем все параграфы и таблицы python-docx → собираем plaintext
2. Plaintext → отправляем LLM с COMPANY_REQUISITES_PROMPT
3. Возвращаем dict готовый к создаию/обновлению Company через CompanyCreate/CompanyPatch

Преимущества vs Vision OCR:
- Текст из DOCX чистый (без OCR-шума, ошибок типа 7707038236 vs 7707038266)
- Дешевле в 5-10× (~500 vs ~3000 токенов)
- Быстрее (~2с vs ~10с)

Только DOCX в первой итерации. PDF/JPG — в следующих пакетах через Vision-путь.
"""
import io
import json
import logging
import re
from typing import Optional

from app.services.llm import get_llm_client
from app.services.ocr.prompts import COMPANY_REQUISITES_PROMPT

log = logging.getLogger(__name__)


class CompanyExtractError(Exception):
    """Не удалось извлечь реквизиты компании."""
    pass


def _read_docx_to_text(docx_bytes: bytes) -> str:
    """
    Читает DOCX и возвращает plaintext всех параграфов и ячеек таблиц.

    Сохраняет порядок чтения (как видит пользователь): параграфы → таблицы по строкам.
    Игнорирует пустые строки.
    """
    try:
        from docx import Document
    except ImportError:
        raise CompanyExtractError(
            "python-docx not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise CompanyExtractError(f"Failed to read DOCX: {e}")

    parts: list[str] = []

    # Параграфы верхнего уровня
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Таблицы — каждая ячейка как отдельная "строка"
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts[-3:]:  # дедупликация ячеек, повторяющихся в merged
                    parts.append(t)

    text = "\\n".join(parts).strip()

    if not text:
        raise CompanyExtractError("DOCX file appears to be empty")

    if len(text) < 30:
        raise CompanyExtractError(
            f"DOCX content too short ({len(text)} chars) — likely not a requisites document"
        )

    return text


def _strip_json_fence(s: str) -> str:
    """Убирает ```json ... ``` обёртку если LLM её добавил, и ищет {...}."""
    s = s.strip()
    if s.startswith("```"):
        # снять ```json или ``` в начале и ``` в конце
        lines = s.split("\\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        s = "\\n".join(lines).strip()

    first_brace = s.find("{")
    last_brace = s.rfind("}")
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        s = s[first_brace : last_brace + 1]
    return s


def _normalize_extracted_fields(data: dict) -> dict:
    """
    Минимальная нормализация полей после LLM:
    - тримминг строк
    - удаление пробелов внутри ИНН/КПП/ОГРН/счетов/БИК
    - конверсия пустых строк в null

    LLM может возвращать "  " вместо null, или ИНН с пробелом — чистим.
    """
    cleaned: dict = {}
    for key, value in data.items():
        if value is None:
            cleaned[key] = None
            continue
        if isinstance(value, str):
            v = value.strip()
            if v == "":
                cleaned[key] = None
                continue
            # Числовые поля — убрать пробелы и не-цифры
            if key in ("ogrn", "inn", "kpp", "bank_account", "bank_bic", "bank_correspondent_account"):
                v = re.sub(r"[^0-9]", "", v)
                if v == "":
                    cleaned[key] = None
                    continue
            cleaned[key] = v
        else:
            cleaned[key] = value
    return cleaned


async def extract_company_from_docx(docx_bytes: bytes) -> dict:
    """
    Главная функция Pack 26.0.

    Args:
        docx_bytes: содержимое .docx файла

    Returns:
        dict с полями совместимыми с CompanyCreate (плюс доп. charter_capital в notes)

    Raises:
        CompanyExtractError при любой проблеме (плохой файл, плохой ответ LLM, и т.д.)
    """
    text = _read_docx_to_text(docx_bytes)

    log.info(f"Pack 26.0: extracted {len(text)} chars from DOCX, calling LLM...")

    client = get_llm_client()

    try:
        response_text = await client.complete(
            system="You are a precise data extraction assistant. Always return strict JSON.",
            user=f"{COMPANY_REQUISITES_PROMPT}\\n\\n--- DOCUMENT TEXT ---\\n{text}\\n--- END ---",
            max_tokens=2048,
            temperature=0.0,
        )
    except Exception as e:
        log.error(f"LLM error in company extraction: {e}", exc_info=True)
        raise CompanyExtractError(f"LLM error: {e}")

    cleaned_response = _strip_json_fence(response_text)

    try:
        parsed = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        log.error(f"Failed to parse LLM JSON: {e}\\nResponse head: {response_text[:500]}")
        raise CompanyExtractError(
            f"LLM returned invalid JSON. Head: {response_text[:200]}"
        )

    if not isinstance(parsed, dict):
        raise CompanyExtractError(f"LLM returned non-dict: {type(parsed).__name__}")

    normalized = _normalize_extracted_fields(parsed)

    log.info(
        f"Pack 26.0: extracted {sum(1 for v in normalized.values() if v is not None)}"
        f"/{len(normalized)} fields"
    )

    return normalized
'''

if EXTRACTOR_PATH.exists():
    print(f"[3/3] [!] {EXTRACTOR_PATH.name} уже существует — заменяю")
EXTRACTOR_PATH.write_text(extractor_code, encoding="utf-8")
print(f"[3/3] Создан: app/services/company_extractor.py ({len(extractor_code.splitlines())} строк)")


# === Финальная проверка синтаксиса ===
errors = []
for p in (PROMPTS_PATH, EXTRACTOR_PATH):
    try:
        ast.parse(p.read_text(encoding="utf-8"))
    except SyntaxError as e:
        errors.append(f"{p.name}: {e}")

if errors:
    print(f"\n[FAIL] синтаксические ошибки:")
    for e in errors:
        print(f"  - {e}")
    print(f"\nОткат:")
    print(f"  Copy-Item -Force '{prompts_backup}' '{PROMPTS_PATH}'")
    print(f"  Remove-Item '{EXTRACTOR_PATH}'")
    sys.exit(1)

print(f"\n[OK] оба файла валидны")
print(f"\n=== Pack 26.0 Stage A применён ===\n")

# === Тестовая команда ===
print("Тест на твоём ООО_РХИ.docx:")
print()
print("  $env:DATABASE_URL = \"postgresql://postgres:uxGVZsShKKnbOZuHyiFpEaufEAnKjYXI@switchyard.proxy.rlwy.net:34408/railway\"")
print("  $env:PYTHONIOENCODING = \"utf-8\"")
print("  cd D:\\VISA\\visa_kit\\backend")
print()
print("  @'")
print("  import asyncio, json")
print("  from pathlib import Path")
print("  from app.services.company_extractor import extract_company_from_docx")
print()
print("  async def main():")
print("      docx_bytes = Path(r\"C:\\Users\\<USER>\\Downloads\\ООО_РХИ.docx\").read_bytes()")
print("      # ИЛИ путь где у тебя лежит файл")
print("      result = await extract_company_from_docx(docx_bytes)")
print("      print(json.dumps(result, ensure_ascii=False, indent=2))")
print()
print("  asyncio.run(main())")
print("  '@ | Out-File -Encoding utf8 _test_pack26.py")
print("  python _test_pack26.py")
print("  Remove-Item _test_pack26.py")
print()
print("Если JSON правильный — едем в Stage B (endpoint + frontend).")
print()
print(f"Откат:")
print(f"  Copy-Item -Force '{prompts_backup}' '{PROMPTS_PATH}'")
print(f"  Remove-Item '{EXTRACTOR_PATH}'")
