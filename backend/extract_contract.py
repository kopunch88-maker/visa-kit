"""
Временный скрипт — извлекает весь текст из шаблона договора
для подготовки списка замен на переменные.
После использования удалить.
"""
import sys
import io

# Принудительно UTF-8 для вывода — иначе Windows валится на ×, ё, и т.п.
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

doc = Document("../templates/docx/contract_template.docx")

print("=== ПАРАГРАФЫ ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print()
print("=== ТАБЛИЦЫ ===")
for ti, table in enumerate(doc.tables):
    rows_count = len(table.rows)
    cols_count = len(table.columns)
    print(f"\n--- Таблица {ti} ({rows_count} строк, {cols_count} колонок) ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                if len(text) > 300:
                    text = text[:300] + "..."
                print(f"  [{ri},{ci}] {text}")