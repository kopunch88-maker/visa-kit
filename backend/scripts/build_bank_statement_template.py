"""
Превращает _bank_statement_original.docx в bank_statement_template.docx.

Скрипт делает простые замены в:
- Шапке выписки (период, балансы, итоги)

Таблицу транзакций нужно подправить РУКАМИ в Word после запуска скрипта.
См. инструкцию в PACK4_README.md.

Запуск:
    python scripts/build_bank_statement_template.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_bank_statement_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "bank_statement_template.docx"


# Замены только в шапке. Таблицу транзакций будем превращать в цикл вручную.
REPLACEMENTS = [
    # Период выписки
    ("За период с  20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    # На случай если в файле другой формат пробелов
    ("За период с 20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),

    # Балансы и итоги
    ("301 018,66 RUR", "{{ bank.opening_balance_formatted }}"),
    ("900 000,00 RUR", "{{ bank.total_income_formatted }}"),
    ("1 171 778,54 RUR", "{{ bank.total_expense_formatted }}"),
    ("29 240,12 RUR", "{{ bank.closing_balance_formatted }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    total = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)

    doc.save(str(TARGET))
    print(f"Saved: {TARGET.name}")
    print(f"Header replacements: {total}")
    print()
    print("==== СЛЕДУЮЩИЙ ШАГ ====")
    print("Откройте bank_statement_template.docx в Word и подправьте")
    print("таблицу транзакций вручную — см. PACK4_README.md, секция")
    print("'Вставка цикла в таблицу транзакций'.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
