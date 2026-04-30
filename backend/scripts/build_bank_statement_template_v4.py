"""
Версия 4 шаблона выписки — план Б.

Стратегия:
1. Берём _bank_statement_original.docx
2. Заменяем шапку (период, балансы, итоги) на Jinja-переменные
3. УДАЛЯЕМ всю таблицу транзакций целиком
4. Вставляем placeholder {{ bank_transactions_subdoc }} вместо таблицы
5. При рендере docx_renderer вставит туда готовую таблицу через subdoc

Это надёжнее любых {%tr%} циклов, потому что таблицу строим программно
с полным контролем структуры и стилей.

Запуск:
    python scripts/build_bank_statement_template_v4.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_bank_statement_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "bank_statement_template.docx"


REPLACEMENTS = [
    ("За период с  20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("За период с 20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("301 018,66 RUR", "{{ bank.opening_balance_formatted }}"),
    ("900 000,00 RUR", "{{ bank.total_income_formatted }}"),
    ("1 171 778,54 RUR", "{{ bank.total_expense_formatted }}"),
    ("29 240,12 RUR", "{{ bank.closing_balance_formatted }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def find_transactions_table(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Дата проводки" in first_row_text and "Код операции" in first_row_text:
            return table
    return None


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    # 1. Замены в шапке
    header_count = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                header_count += process_paragraphs(cell.paragraphs)
    print(f"  Header replacements: {header_count}")

    # 2. Удаляем таблицу транзакций целиком
    tx_table = find_transactions_table(doc)
    if tx_table is None:
        print("[ERROR] Transactions table not found")
        return 1

    # Получаем родителя таблицы (обычно body)
    table_element = tx_table._element
    parent = table_element.getparent()

    # Создаём параграф с placeholder перед таблицей
    # Этот placeholder docxtpl заменит на subdoc с программно собранной таблицей
    placeholder_paragraph = doc.paragraphs[0]._element.makeelement(
        qn("w:p"), {}
    )
    pr = doc.paragraphs[0]._element.makeelement(qn("w:r"), {})
    pt = doc.paragraphs[0]._element.makeelement(qn("w:t"), {})
    pt.text = "{{p bank_table_subdoc}}"  # {{p ...}} - это синтаксис docxtpl для inline-вставки subdoc
    pr.append(pt)
    placeholder_paragraph.append(pr)

    # Вставляем placeholder перед таблицей и удаляем саму таблицу
    parent.insert(list(parent).index(table_element), placeholder_paragraph)
    parent.remove(table_element)
    print(f"  Replaced transactions table with placeholder")

    doc.save(str(TARGET))
    print(f"[OK] Saved: {TARGET.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
