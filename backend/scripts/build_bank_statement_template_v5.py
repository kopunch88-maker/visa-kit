"""
Версия 5 шаблона выписки — финал.

Стратегия: клонирование строки-образца.

1. Берём _bank_statement_original.docx
2. Заменяем шапку (балансы, период) на Jinja-переменные
3. В таблице транзакций ОСТАВЛЯЕМ заголовочную строку и одну строку-образец
4. В строке-образце текст ячеек заменяем на МАРКЕРЫ (НЕ Jinja, чтобы docxtpl
   их не пытался интерпретировать). Маркеры:
       __TX_DATE__
       __TX_CODE__
       __TX_DESCRIPTION__
       __TX_AMOUNT__
5. Удаляем все остальные 14 строк (оригинальные транзакции Алиева)

При рендере docx_renderer:
- Подставит Jinja-переменные в шапке (стандартный docxtpl)
- Найдёт строку-образец по маркерам
- Клонирует её N раз через deepcopy(_tr)
- В каждой копии заменит маркеры на реальные данные транзакции
- Удалит оригинальную строку-образец

Запуск:
    python scripts/build_bank_statement_template_v5.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_bank_statement_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "bank_statement_template.docx"


HEADER_REPLACEMENTS = [
    ("За период с  20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("За период с 20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("301 018,66 RUR", "{{ bank.opening_balance_formatted }}"),
    ("900 000,00 RUR", "{{ bank.total_income_formatted }}"),
    ("1 171 778,54 RUR", "{{ bank.total_expense_formatted }}"),
    ("29 240,12 RUR", "{{ bank.closing_balance_formatted }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs, replacements) -> int:
    count = 0
    for p in paragraphs:
        for old, new in replacements:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def find_transactions_table(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Дата проводки" in first_row_text and "Код операции" in first_row_text:
            return table
    return None


def replace_cell_text_with_marker(cell, marker: str):
    """
    Заменяет ВЕСЬ текст в ячейке на одну строку-маркер,
    сохраняя стили первого run первого параграфа.
    """
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    # Первый параграф оставляем, остальные удаляем
    first_para = paragraphs[0]
    for p in paragraphs[1:]:
        p_element = p._element
        p_element.getparent().remove(p_element)

    # В первом параграфе оставляем только один run с маркером
    if first_para.runs:
        first_run = first_para.runs[0]
        first_run.text = marker
        # Удаляем все runs кроме первого
        for run in first_para.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        # Если runs не было — добавляем
        first_para.add_run(marker)


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    # 1. Замены в шапке (paragraphs and other tables)
    header_count = process_paragraphs(doc.paragraphs, HEADER_REPLACEMENTS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                header_count += process_paragraphs(cell.paragraphs, HEADER_REPLACEMENTS)
    print(f"  Header replacements: {header_count}")

    # 2. Перестраиваем таблицу транзакций
    tx_table = find_transactions_table(doc)
    if tx_table is None:
        print("[ERROR] Transactions table not found")
        return 1

    n_rows = len(tx_table.rows)
    print(f"  Transactions table has {n_rows} rows")

    if n_rows < 2:
        print("[ERROR] Table has no template row to use")
        return 1

    # Удаляем все строки кроме первой (заголовок) и второй (образец)
    tbl_element = tx_table._tbl
    rows_to_delete = list(tx_table.rows)[2:]
    for row in rows_to_delete:
        tbl_element.remove(row._tr)
    print(f"  Removed {len(rows_to_delete)} extra transaction rows")

    # 3. В строке-образце (теперь это row[1]) заменяем содержимое на маркеры
    template_row = tx_table.rows[1]
    cells = template_row.cells
    n_cells = len(cells)
    print(f"  Template row has {n_cells} cells")

    # Маркеры идут в порядке колонок: дата, код, описание, сумма
    markers = ["__TX_DATE__", "__TX_CODE__", "__TX_DESCRIPTION__", "__TX_AMOUNT__"]
    for i, marker in enumerate(markers):
        if i < n_cells:
            replace_cell_text_with_marker(cells[i], marker)

    print(f"  Replaced text in template row with markers")

    doc.save(str(TARGET))
    print(f"[OK] Saved: {TARGET.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
