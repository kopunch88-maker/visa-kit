"""
Полностью пересобирает cv_template.docx программно.

В отличие от ручной правки в Word, этот скрипт:
1. Читает _cv_original.docx
2. Программно перестраивает структуру — паграфы с Jinja-выражениями
3. Все блоки опционально — если у клиента нет education / work_history /
   languages, то соответствующая секция просто не отрисовывается

Запуск:
    python scripts/build_cv_template_v2.py
"""
import sys
import io
from copy import deepcopy
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_cv_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "cv_template.docx"


# Содержимое всего CV-шаблона.
# Каждый элемент — это (текст_параграфа, стиль_исходного_параграфа_по_номеру)
# Структура повторяет исходное резюме, но всё содержимое заменено
# на Jinja-выражения с условиями.

# Параграфы из исходника, на стиль которых будем опираться:
# P0   — имя (большое жирное)
# P2-7 — личные данные
# P9   — заголовок "Образование"
# P10  — текст про образование
# P17  — заголовок "Опыт работы"
# P19  — даты работы
# P20  — название компании
# P21  — должность
# P22  — обязанность

# Текст шаблона. Используем простые маркеры стилей:
#   H1 — большой заголовок (имя)
#   H2 — заголовок секции
#   N  — обычный текст
#   B  — жирный обычный
#   I  — курсив

CV_STRUCTURE = [
    # Шапка
    ("H1", "{{ applicant.full_name_native|upper }}"),
    ("N",  ""),
    ("N",  "Дата рождения\t{{ fmt_date_ru(applicant.birth_date) }}"),
    ("N",  "Гражданство\t{{ applicant.nationality_ru }}"),
    ("N",  "Образование\tВысшее"),
    ("N",  "Телефон\t{{ applicant.phone }}"),
    ("N",  "E-mail\t{{ applicant.email }}"),
    ("N",  ""),

    # Образование
    ("H2", "Образование"),
    ("N",  "{% if applicant.education %}{% for edu in applicant.education %}"),
    ("N",  "{{ edu.institution }}"),
    ("N",  "Год окончания: {{ edu.graduation_year }}"),
    ("N",  "{{ edu.degree }}"),
    ("N",  "Специальность: {{ edu.specialty }}"),
    ("N",  ""),
    ("N",  "{% endfor %}{% else %}(не указано){% endif %}"),
    ("N",  ""),

    # Опыт работы
    ("H2", "Опыт работы"),
    ("N",  "{% if applicant.work_history %}{% for job in applicant.work_history %}"),
    ("B",  "{{ job.period_start }} – {{ job.period_end }}"),
    ("B",  "{{ job.company }}"),
    ("I",  "{{ job.position }}"),
    ("N",  "{% for duty in job.duties %}{{ duty }};"),
    ("N",  "{% endfor %}"),
    ("N",  ""),
    ("N",  "{% endfor %}{% else %}(не указано){% endif %}"),
    ("N",  ""),

    # Доп. информация
    ("H2", "Дополнительная информация"),
    ("N",  "{% if applicant.languages %}Иностранные языки: {{ applicant.languages|join(', ') }}{% endif %}"),
]


def build_cv():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    # Сохраняем стили из первых параграфов исходного документа,
    # чтобы новые параграфы имели те же шрифты/размеры
    style_samples = {}
    if len(doc.paragraphs) > 0:
        style_samples["H1"] = deepcopy(doc.paragraphs[0])  # имя в шапке
    if len(doc.paragraphs) > 9:
        style_samples["H2"] = deepcopy(doc.paragraphs[9])  # "Образование"
    if len(doc.paragraphs) > 10:
        style_samples["N"] = deepcopy(doc.paragraphs[10])  # обычный текст
    if len(doc.paragraphs) > 19:
        style_samples["B"] = deepcopy(doc.paragraphs[19])  # дата работы (жирная)
    if len(doc.paragraphs) > 21:
        style_samples["I"] = deepcopy(doc.paragraphs[21])  # должность

    # Удаляем все параграфы документа
    body = doc.element.body
    for p in list(doc.paragraphs):
        p_element = p._element
        p_element.getparent().remove(p_element)

    # Также удаляем таблицы (в CV их нет, но на всякий случай)
    for table in list(doc.tables):
        t_element = table._element
        t_element.getparent().remove(t_element)

    # Создаём параграфы из CV_STRUCTURE
    for style_key, text in CV_STRUCTURE:
        p = doc.add_paragraph()
        # Скопируем стиль абзаца из эталона если есть
        if style_key in style_samples:
            sample = style_samples[style_key]
            # Копируем pPr (свойства параграфа)
            sample_pPr = sample._element.find(qn("w:pPr"))
            if sample_pPr is not None:
                p._element.insert(0, deepcopy(sample_pPr))

        run = p.add_run(text)
        # Применяем простое форматирование
        if style_key == "H1":
            run.bold = True
            run.font.size = Pt(16)
        elif style_key == "H2":
            run.bold = True
            run.font.size = Pt(13)
        elif style_key == "B":
            run.bold = True
        elif style_key == "I":
            run.italic = True

    doc.save(str(TARGET))
    print(f"[OK] Saved: {TARGET.name}")
    print(f"     Paragraphs: {len(CV_STRUCTURE)}")
    return 0


if __name__ == "__main__":
    sys.exit(build_cv())
