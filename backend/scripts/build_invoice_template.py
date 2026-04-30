"""
Превращает _invoice_original.docx в invoice_template.docx с переменными {{ }}.

Запуск:
    python scripts/build_invoice_template.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_invoice_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "invoice_template.docx"


# ВАЖНО: порядок имеет значение. Длинные строки сначала, потом более короткие.
REPLACEMENTS = [
    # === Шапка — получатель (сам исполнитель) ===
    ("Алиев Джафар Надирович", "{{ applicant.full_name_native }}"),
    ("352919, Краснодарский край, г. Армавир, ул. 11-я Линия, д. 31 кв. 2",
     "{{ applicant.home_address }}"),

    # === Заголовок счёта ===
    ("Счет № 1 от 31 января 2026 года",
     "Счет № {{ invoice.sequence_number }} от {{ fmt_date_human_ru(invoice.document_date) }}"),

    # === Плательщик в отдельной строке ===
    ('ООО «Строительная компания СК10», ИНН 6168006148, КПП 616401001',
     "{{ company.short_name }}, ИНН {{ company.tax_id_primary }}, КПП {{ company.tax_id_secondary }}"),

    # Адрес плательщика — двойной пробел в исходнике, частая опечатка
    ("344002, г. Ростов-на-дону,  ул. Московская, зд. 73/29а, ком. 7",
     "{{ company.legal_address }}"),
    ("344002, г. Ростов-на-дону, ул. Московская, зд. 73/29а, ком. 7",
     "{{ company.legal_address }}"),

    # === Реквизиты получателя в первой таблице ===
    ("ИНН 230217957801", "ИНН {{ applicant.inn }}"),
    ("40803840441563809831", "{{ applicant.bank_account }}"),
    ("АО «АЛЬФА-БАНК», г. Москва", "{{ applicant.bank_name }}"),
    ("044525593", "{{ applicant.bank_bic }}"),
    ("30101810200000000593", "{{ applicant.bank_correspondent_account }}"),

    # === Услуга в таблице товаров ===
    (
        "Услуги инженера-геодезиста (камеральщик) за период 01.01.2026-31.01.2026 по Договору №004/09/25 от 05.09.2025г.",
        "Услуги {{ position.title_ru_genitive }} за период {{ fmt_date_ru(invoice.period_start) }}-{{ fmt_date_ru(invoice.period_end) }} по Договору №{{ contract.number }} от {{ fmt_date_ru(contract.sign_date) }}г.",
    ),

    # === Сумма "300 000" — встречается несколько раз. Используем jinja `{{ }}` напрямую ===
    # ВНИМАНИЕ: docx может разбить эти числа на несколько runs, поэтому fallback на простую замену
    ("Триста тысяч рублей 00 копеек",
     "{{ contract.salary_rub_words|capitalize }} рублей 00 копеек"),
    ("Всего одно наименование на сумму 300 000 руб.",
     "Всего одно наименование на сумму {{ fmt_money(invoice.salary_rub) }} руб."),
    # Цифру 300 000 в столбце «Сумма» табличный счётчик заменит несколько раз
    ("300 000", "{{ fmt_money(invoice.salary_rub) }}"),

    # === Подпись ===
    ("Алиев Д.Н.", "{{ applicant.initials_native }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))
    total = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)

    doc.save(str(TARGET))
    print(f"Saved: {TARGET.name}")
    print(f"Replacements: {total}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
