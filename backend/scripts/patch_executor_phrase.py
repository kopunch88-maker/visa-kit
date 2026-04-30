"""
Pack 8.5 — финальная правка формулировки исполнителя в шаблоне договора.

ДО:
    Г-н республики {{ applicant.nationality_ru_genitive }} {{ applicant.full_name_native }}, 
    именуемый в дальнейшем "Исполнитель"

ПОСЛЕ:
    {{ applicant.citizen_phrase }} {{ applicant.full_name_native }}, 
    именуем{{ applicant.named_suffix }} в дальнейшем "Исполнитель"

Где:
    citizen_phrase: "Гражданин Российской Федерации" / "Гражданка Республики Армения" / ...
    named_suffix:  "ый" (мужской) или "ая" (женский)

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python scripts\\patch_executor_phrase.py
"""
import sys
import io
import shutil
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

BACKEND_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = BACKEND_ROOT.parent / "templates" / "docx" / "contract_template.docx"

if not TEMPLATE_PATH.exists():
    print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
    sys.exit(1)

backup = TEMPLATE_PATH.with_suffix(".docx.bak3")
shutil.copy(TEMPLATE_PATH, backup)
print(f"[OK] Backup saved: {backup}")

from docx import Document

doc = Document(str(TEMPLATE_PATH))

# Что заменяем (две независимые подстановки)
REPLACEMENTS = [
    (
        # ДО: "Г-н республики {{ applicant.nationality_ru_genitive }}"
        "Г-н республики {{ applicant.nationality_ru_genitive }}",
        # ПОСЛЕ: новая комбинированная переменная
        "{{ applicant.citizen_phrase }}",
    ),
    (
        # ДО: "именуемый"
        # ПОСЛЕ: динамическое окончание
        "именуемый в дальнейшем",
        "именуем{{ applicant.named_suffix }} в дальнейшем",
    ),
]

patched_count = 0
for para in doc.paragraphs:
    text = para.text
    new_text = text
    changed_in_this_para = False
    
    for old_str, new_str in REPLACEMENTS:
        if old_str in new_text:
            new_text = new_text.replace(old_str, new_str)
            changed_in_this_para = True
    
    if changed_in_this_para:
        if len(para.runs) > 0:
            first_run = para.runs[0]
            for r in para.runs:
                r.text = ""
            first_run.text = new_text
            patched_count += 1
            print(f"[OK] Patched paragraph: {text[:80]}...")

if patched_count == 0:
    print("[WARN] No replacements made. Maybe template was already patched?")
    sys.exit(1)

print(f"[INFO] Total {patched_count} paragraph(s) patched")

doc.save(str(TEMPLATE_PATH))
print("[DONE] Template saved")

# Verification
print("\n=== Verification ===")
doc = Document(str(TEMPLATE_PATH))
shown = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text.strip():
        print(f"P{i}: {text}")
        shown += 1
        if shown >= 6:
            break

print()
print("=" * 70)
print("СЛЕДУЮЩИЙ ШАГ — обновить context.py")
print("=" * 70)
print()
print("Откройте D:\\VISA\\visa_kit\\backend\\app\\templates_engine\\context.py")
print()
print("1. Найдите словарь COUNTRY_GENITIVE (если он есть) или место где готовятся")
print("   данные applicant.")
print()
print("2. Замените существующий маппинг nationality_ru_genitive на полный словарь")
print("   с правильными юридическими названиями (см. инструкцию в чате).")
print()
print("3. В блок где готовятся данные applicant добавьте поля:")
print("   - citizen_phrase  (Гражданин/Гражданка + страна в родительном)")
print("   - named_suffix    ('ый' или 'ая')")
