"""
Превращает _cv_original.docx в cv_template.docx.

CV особенный — в нём переменное количество мест работы. Поэтому делаем
Jinja-цикл {% for job in applicant.work_history %}.

Структура исходного резюме:
  - Шапка: имя, дата рождения, гражданство
  - Образование (одно или несколько — пока поддерживаем одно)
  - Опыт работы — несколько блоков (5 у Алиева)
  - Дополнительная информация: права, языки

Запуск:
    python scripts/build_cv_template.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_cv_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "cv_template.docx"


# Простые замены (имя, дата рождения, гражданство, образование)
REPLACEMENTS = [
    # === Шапка ===
    ("АЛИЕВ ДЖАФАР НАДИРОВИЧ", "{{ applicant.full_name_native|upper }}"),

    # Дата рождения (формат "Дата рождения   09.03.1963")
    ("09.03.1963", "{{ fmt_date_ru(applicant.birth_date) }}"),

    # Гражданство
    ("Азербайджан", "{{ applicant.nationality_ru_genitive }}"),

    # === Образование ===
    (
        "Государственное образовательное учреждение           высшего профессионального образования «Ростовский государственный строительный университет»,",
        "{{ applicant.education[0].institution if applicant.education else '' }},",
    ),
    ("Год окончания: 2010",
     "Год окончания: {{ applicant.education[0].graduation_year if applicant.education else '' }}"),
    ("Инженер", "{{ applicant.education[0].degree if applicant.education else '' }}"),
    ("Прикладная геодезия",
     "{{ applicant.education[0].specialty if applicant.education else '' }}"),

    # === Языки и права (в самом конце) ===
    ("Иностранные языки: Английский В1",
     "Иностранные языки: {{ applicant.languages|join(', ') }}"),
]


# Маркеры начала и конца блока опыта работы — для замены на Jinja-цикл
WORK_HISTORY_START_MARKER = "Сентябрь 2025 – по настоящее время"
WORK_HISTORY_END_MARKER_BEFORE = "обработка данных в программе AutoCAD."  # последний пункт последней работы

# Текст, которым заменим всю секцию опыта работы — Jinja-цикл
WORK_HISTORY_REPLACEMENT = """{% for job in applicant.work_history %}{{ job.period_start }} – {{ job.period_end }}
{{ job.company }}
{{ job.position }}
{% for duty in job.duties %}{{ duty }};
{% endfor %}
{% endfor %}"""


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def replace_work_history_block(doc):
    """
    Находит блок 'Опыт работы' с реальными местами работы Алиева и
    заменяет его на Jinja-цикл, который перебирает applicant.work_history.

    Стратегия: находим параграф со словом "Опыт работы" (заголовок),
    после него идут блоки работ. Заменяем их все на один параграф с циклом.
    """
    paragraphs = doc.paragraphs

    # Находим индексы начала и конца блока опыта
    start_idx = None
    end_idx = None

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if text == "Опыт работы" and start_idx is None:
            start_idx = i + 1
        if WORK_HISTORY_END_MARKER_BEFORE in text:
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print("[WARN] Could not find work history block markers")
        return 0

    # Очищаем содержимое всех параграфов в блоке кроме первого
    # В первом ставим Jinja-цикл
    for j in range(start_idx, end_idx + 1):
        p = paragraphs[j]
        if j == start_idx:
            # В первый параграф ставим открытие цикла
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "{% for job in applicant.work_history %}{{ job.period_start }} – {{ job.period_end }}"
        elif j == start_idx + 1:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "{{ job.company }}"
        elif j == start_idx + 2:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "{{ job.position }}"
        elif j == start_idx + 3:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "{% for duty in job.duties %}{{ duty }};{% endfor %}"
        elif j == end_idx:
            # Последний — закрываем цикл
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "{% endfor %}"
        else:
            # Промежуточные — очищаем
            for run in p.runs:
                run.text = ""

    return end_idx - start_idx + 1


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    # Шаг 1: простые замены (имя, образование, языки)
    total = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)

    # Шаг 2: блок опыта работы → Jinja-цикл
    work_replaced = replace_work_history_block(doc)
    print(f"Simple replacements: {total}")
    print(f"Work history paragraphs replaced: {work_replaced}")

    doc.save(str(TARGET))
    print(f"Saved: {TARGET.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
