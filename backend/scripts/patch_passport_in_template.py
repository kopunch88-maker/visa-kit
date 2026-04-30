"""
Pack 8.5 — финальная правка паспорта в шаблоне договора.

В блоке "8. Адреса и реквизиты Сторон" в таблице сейчас захардкожено:
    Паспорт RUS {{ applicant.passport_number }}, 
    выдан {{ applicant.passport_issue_date }}г. {{ applicant.passport_issuer }}

Заменяем на:
    Паспорт {{ applicant.passport_formatted }}, 
    выдан {{ applicant.passport_issuer }} {{ applicant.passport_issue_date_str }}

Где passport_formatted уже содержит правильное "серии 4503 № 123456" для РФ
или "№ C01366076" для иностранцев.

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python scripts\\patch_passport_in_template.py
"""
import sys
import io
import shutil
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

BACKEND_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = BACKEND_ROOT.parent / "templates" / "docx" / "contract_template.docx"

if not TEMPLATE_PATH.exists():
    print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
    sys.exit(1)

backup = TEMPLATE_PATH.with_suffix(".docx.bak4")
shutil.copy(TEMPLATE_PATH, backup)
print(f"[OK] Backup saved: {backup}")

from docx import Document

doc = Document(str(TEMPLATE_PATH))


def patch_text_in_paragraph(para, old_str, new_str):
    """Заменяет old_str на new_str в параграфе, сохраняя стиль первого run."""
    if old_str not in para.text:
        return False
    new_text = para.text.replace(old_str, new_str)
    if len(para.runs) > 0:
        first_run = para.runs[0]
        for r in para.runs:
            r.text = ""
        first_run.text = new_text
        return True
    return False


# Замены — список пар (что искать, на что менять)
# Идём от более длинного к короткому, чтобы не было частичных совпадений
REPLACEMENTS = [
    # Старый формат паспорта целиком на новый
    (
        "Паспорт RUS {{ applicant.passport_number }}",
        "Паспорт {{ applicant.passport_formatted }}",
    ),
    # На случай если без RUS, но с старым форматом даты
    (
        "выдан {{ applicant.passport_issue_date }}г.",
        "выдан {{ applicant.passport_issuer }} {{ applicant.passport_issue_date_str }}",
    ),
    # Возможные другие варианты — на всякий случай
    (
        "Паспорт {{ applicant.passport_country_code }} {{ applicant.passport_number }}",
        "Паспорт {{ applicant.passport_formatted }}",
    ),
]


total_patched = 0

# Сначала проходим по всем параграфам в документе
for para in doc.paragraphs:
    for old, new in REPLACEMENTS:
        if patch_text_in_paragraph(para, old, new):
            total_patched += 1
            print(f"[OK] Patched in paragraph: '{old[:50]}...'")

# Затем по всем таблицам — там тоже параграфы
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in REPLACEMENTS:
                    if patch_text_in_paragraph(para, old, new):
                        total_patched += 1
                        print(f"[OK] Patched in table cell: '{old[:50]}...'")

if total_patched == 0:
    print()
    print("[WARN] Ничего не заменено. Возможно шаблон уже пропатчен")
    print("       или формулировка отличается от ожидаемой.")
    print()
    print("Покажу текущее содержимое всех параграфов с 'паспорт' или 'passport':")
    print()
    
    # Показываем что есть в параграфах
    for i, para in enumerate(doc.paragraphs):
        if "паспорт" in para.text.lower() or "passport" in para.text.lower():
            print(f"P{i}: {para.text[:200]}")
    
    # И в таблицах
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if "паспорт" in para.text.lower() or "passport" in para.text.lower():
                        print(f"Table{ti} Row{ri} Cell{ci}: {para.text[:200]}")
    sys.exit(1)

print(f"\n[INFO] Всего заменено: {total_patched}")

doc.save(str(TEMPLATE_PATH))
print("[DONE] Template saved")

# Verification
print("\n=== Verification (fragments with 'паспорт') ===")
doc = Document(str(TEMPLATE_PATH))
for i, para in enumerate(doc.paragraphs):
    if "паспорт" in para.text.lower():
        print(f"P{i}: {para.text[:200]}")

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                if "паспорт" in para.text.lower():
                    print(f"Table{ti} Row{ri} Cell{ci}: {para.text[:200]}")
