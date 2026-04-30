"""
Версия 3 — с правильной структурой 4-колоночной таблицы.

Колонки выписки:
1. Дата проводки
2. Код операции
3. Описание
4. Сумма в валюте счета

(В предыдущих версиях я ошибочно делал 5 колонок — Сумма + в валюте счета
читалось как две.)

Запуск:
    python scripts/build_bank_statement_template_v3.py
"""
import sys
import io
from copy import deepcopy
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_bank_statement_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "bank_statement_template.docx"


REPLACEMENTS = [
    ("За период с  20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("За период с 20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("301 018,66 RUR", "{{ bank.opening_balance_formatted }}"),
    ("900 000,00 RUR", "{{ bank.total_income_formatted }}"),
    ("1 171 778,54 RUR", "{{ bank.total_expense_formatted }}"),
    ("29 240,12 RUR", "{{ bank.closing_balance_formatted }}"),
]


# 4 ячейки — по числу колонок в реальной таблице
ROW_CELL_CONTENTS = [
    "{%tr for transaction in bank.transactions %}{{ transaction.date_formatted }}",
    "{{ transaction.code }}",
    "{{ transaction.description }}",
    "{{ transaction.amount_formatted }}{%tr endfor %}",
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def find_transactions_table(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Дата проводки" in first_row_text and "Код операции" in first_row_text:
            return table
    return None


def rebuild_transactions_table(table) -> int:
    """
    Стратегия: берём первую строку с реальной транзакцией (которая в исходнике
    имеет правильные стили и структуру) и используем её как образец. Удаляем
    все строки кроме заголовка и этой одной. Заменяем содержимое ячеек на
    Jinja-переменные. Это сохраняет форматирование исходника.
    """
    if len(table.rows) < 2:
        print("[ERROR] Table has no transaction rows to use as template")
        return 0

    tbl_element = table._tbl

    # Сохраняем заголовок (row 0) и первую транзакцию (row 1)
    # Удаляем всё начиная с row 2
    rows_to_remove = list(table.rows)[2:]
    for row in rows_to_remove:
        tbl_element.remove(row._tr)

    # Теперь оставшаяся row 1 — наш образец строки. Подменяем её содержимое.
    template_row = table.rows[1]
    cells = template_row.cells

    # Проверяем число колонок
    n_cells = len(cells)
    print(f"  Template row has {n_cells} cells")

    if n_cells < 4:
        print(f"[WARN] Expected at least 4 cells, got {n_cells}")

    # Заменяем содержимое первых 4 ячеек на Jinja-переменные
    # Если ячеек больше 4 — оставшиеся очищаем
    for i, cell in enumerate(cells):
        # Удаляем все параграфы внутри ячейки (там был текст транзакции Алиева)
        for para in list(cell.paragraphs):
            p_element = para._element
            p_element.getparent().remove(p_element)

        if i < len(ROW_CELL_CONTENTS):
            content = ROW_CELL_CONTENTS[i]
        else:
            content = ""  # пустая для лишних ячеек

        # Добавляем новый параграф с нашей переменной
        cell.add_paragraph(content)

    return len(rows_to_remove)


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    header_count = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                header_count += process_paragraphs(cell.paragraphs)
    print(f"  Header replacements: {header_count}")

    tx_table = find_transactions_table(doc)
    if tx_table is None:
        print("[ERROR] Transactions table not found")
        return 1

    removed = rebuild_transactions_table(tx_table)
    print(f"  Removed {removed} extra transaction rows")
    print(f"  Replaced first transaction row with Jinja loop template")

    doc.save(str(TARGET))
    print(f"[OK] Saved: {TARGET.name}")
    print()
    print("Now run:")
    print("  python scripts\\render_test_bank_statement.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())