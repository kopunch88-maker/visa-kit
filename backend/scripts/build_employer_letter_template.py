"""
Превращает _employer_letter_original.docx в employer_letter_template.docx.

Письмо от компании содержит EUR-эквивалент зарплаты, который подставляется
автоматически из курса ЦБ РФ на дату письма (employer_letter_date).

Запуск:
    python scripts/build_employer_letter_template.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_employer_letter_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "employer_letter_template.docx"


REPLACEMENTS = [
    # === Исходящий номер и дата письма ===
    ("Исх. №544 от 17.04.2026г.",
     "Исх. №{{ letter.number }} от {{ fmt_date_ru(letter.date) }}г."),

    # === Компания: длинная и короткая форма ===
    ("Общество с ограниченной ответственностью «Строительная компания СК10» (ООО «СК10»), ИНН 6168006148",
     "{{ company.full_name_ru }} ({{ company.short_name }}), ИНН {{ company.tax_id_primary }}"),

    # Дальше "ООО «СК10»" встречается отдельно несколько раз
    ("ООО «СК10»", "{{ company.short_name }}"),

    # === Заявитель ===
    ("Алиев Джафар Надирович", "{{ applicant.full_name_native }}"),

    # === Договор ===
    ("№ 004/09/25 от «05» сентября 2025 года",
     "№ {{ contract.number }} от {{ fmt_date_long_ru(contract.sign_date) }}"),

    # === Должность (родительный падеж) ===
    ("инженера-геодезиста (камеральщик)", "{{ position.title_ru_genitive }}"),

    # === Сумма с EUR-эквивалентом ===
    # Это ключевая строка — содержит RUB и автоматически рассчитанный EUR
    (
        "300 000 (триста тысяч рублей) в месяц эквивалентно по курсу на дату письма 3.355 евро (три тысячи триста пятьдесят пять) евро",
        "{{ fmt_money(contract.salary_rub) }} ({{ contract.salary_rub_words }} рублей) в месяц эквивалентно по курсу на дату письма {{ eur.amount_int }} евро ({{ eur.amount_words_es }}) евро",
    ),

    # === Срок действия договора ===
    ("31 августа 2029 года",
     "{{ contract.end_date.day }} {{ ['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября','декабря'][contract.end_date.month - 1] }} {{ contract.end_date.year }} года"),

    # === Подпись ===
    ("Тараскин Ю.А.", "{{ company.director_short_ru }}"),
    ("Генеральный директор", "Генеральный директор"),  # placeholder, no change

    # ИНН компании (на случай если встречается отдельно)
    ("6168006148", "{{ company.tax_id_primary }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))
    total = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)

    doc.save(str(TARGET))
    print(f"Saved: {TARGET.name}")
    print(f"Replacements: {total}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
