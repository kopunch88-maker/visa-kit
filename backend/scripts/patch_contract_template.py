"""
Pack 8.5 — Финал C: чиним шаблон договора.

ЧТО ДЕЛАЕТ:
1. Создаёт backup contract_template.docx
2. В шаблоне находит захардкоженную строку с "Ростов-на-Дону" и заменяет
   на jinja-переменные:
   "г. {{ contract.sign_city }}     «{{ contract.sign_date_str }}»"
3. Печатает SQL-инструкцию для проверки что в context.py готовится sign_date_str

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python scripts\\patch_contract_template.py

Если хотите откатить:
    Скопируйте contract_template.docx.bak обратно в contract_template.docx
"""
import sys
import io
import shutil
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

BACKEND_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = BACKEND_ROOT.parent / "templates" / "docx" / "contract_template.docx"

if not TEMPLATE_PATH.exists():
    print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
    sys.exit(1)

print(f"Patching template: {TEMPLATE_PATH}")

# Backup
backup = TEMPLATE_PATH.with_suffix(".docx.bak")
shutil.copy(TEMPLATE_PATH, backup)
print(f"[OK] Backup saved: {backup}")

from docx import Document

doc = Document(str(TEMPLATE_PATH))

# Используем единое поле sign_date_str — менее хрупко чем 3 отдельные переменные
REPLACEMENT = (
    "г. {{ contract.sign_city }}"
    "                                                                                                             "
    "«{{ contract.sign_date_str }}»"
)

found = False
for para in doc.paragraphs:
    text = para.text
    # Простой поиск — наличие "Ростов-на-Дону" и года достаточно характерны
    if "Ростов-на-Дону" in text and "г." in text:
        if len(para.runs) > 0:
            first_run = para.runs[0]
            # Очистка всех runs
            for r in para.runs:
                r.text = ""
            first_run.text = REPLACEMENT
            found = True
            print(f"[OK] Patched: {text[:80]}")
            break

if not found:
    print("[WARN] Hardcoded line not found. Either template was already patched")
    print("       or the format is different than expected.")
    sys.exit(1)

doc.save(str(TEMPLATE_PATH))
print(f"[DONE] Template saved")

# Verification
print("\n=== Verification (first 5 non-empty paragraphs) ===")
doc = Document(str(TEMPLATE_PATH))
shown = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text.strip():
        print(f"P{i}: {text}")
        shown += 1
        if shown >= 5:
            break

print()
print("=" * 70)
print("СЛЕДУЮЩИЙ ШАГ — добавить sign_date_str в context.py")
print("=" * 70)
print()
print("Откройте D:\\VISA\\visa_kit\\backend\\app\\templates_engine\\context.py")
print("Найдите блок где формируется 'contract', обычно так:")
print()
print('    "contract": {')
print('        "number": application.contract_number or "",')
print('        "sign_date": application.contract_sign_date,')
print('        "sign_city": application.contract_sign_city or "",')
print("        ...")
print()
print("ДОБАВЬТЕ строку sign_date_str (форматированную дату):")
print()
print('        "sign_date_str": _format_date_ru(application.contract_sign_date),')
print()
print("И добавьте функцию _format_date_ru если её нет:")
print()
print("def _format_date_ru(d):")
print('    """04.05.2025 → "04 мая 2025 г."""')
print("    if not d:")
print('        return ""')
print('    months = {1: "января", 2: "февраля", 3: "марта", 4: "апреля",')
print('              5: "мая", 6: "июня", 7: "июля", 8: "августа",')
print('              9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"}')
print('    return f"{d.day:02d} {months[d.month]} {d.year} г."')
