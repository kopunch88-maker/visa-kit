"""
Полностью пересобирает bank_statement_template.docx программно.

В отличие от build_bank_statement_template.py (который правил оригинал),
этот скрипт:
1. Читает _bank_statement_original.docx
2. Заменяет шапку на переменные
3. УДАЛЯЕТ все 15 строк-транзакций из таблицы
4. Создаёт ОДНУ строку с правильным {%tr for ... %} ... {%tr endfor %} циклом
5. Сохраняет как bank_statement_template.docx

Без необходимости открывать Word и без автозамен.

Запуск:
    python scripts/build_bank_statement_template_v2.py
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_bank_statement_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "bank_statement_template.docx"


# === Замены в шапке (как раньше) ===

REPLACEMENTS = [
    ("За период с  20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("За период с 20.01.2026 по 19.04.2026",
     "За период с {{ bank.period_start_formatted }} по {{ bank.period_end_formatted }}"),
    ("301 018,66 RUR", "{{ bank.opening_balance_formatted }}"),
    ("900 000,00 RUR", "{{ bank.total_income_formatted }}"),
    ("1 171 778,54 RUR", "{{ bank.total_expense_formatted }}"),
    ("29 240,12 RUR", "{{ bank.closing_balance_formatted }}"),
]


# === Содержимое строки-шаблона для цикла транзакций ===
# Используем {%tr ... %} — это специальный синтаксис docxtpl для повторения
# строки таблицы целиком.

ROW_CELL_CONTENTS = [
    "{%tr for transaction in bank.transactions %}{{ transaction.date_formatted }}",
    "{{ transaction.code }}",
    "{{ transaction.description }}",
    "{{ transaction.amount_formatted }}",
    "{{ transaction.amount_formatted }}{%tr endfor %}",
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def find_transactions_table(doc):
    """
    Находит таблицу с транзакциями — это таблица с колонкой 'Дата проводки'.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Дата проводки" in first_row_text and "Код операции" in first_row_text:
            return table
    return None


def rebuild_transactions_table(table) -> int:
    """
    Удаляет все строки кроме заголовочной, добавляет одну строку-шаблон
    с Jinja-циклом.
    """
    if not table.rows:
        return 0

    # Сохраняем заголовочную строку
    header_row = table.rows[0]

    # Удаляем все строки кроме первой через прямую правку XML таблицы
    # python-docx не имеет высокоуровневого API для удаления строки —
    # делаем через XML
    tbl_element = table._tbl
    rows_to_remove = list(table.rows)[1:]  # все строки кроме заголовка
    for row in rows_to_remove:
        tbl_element.remove(row._tr)

    # Добавляем одну новую строку с тем же количеством ячеек что в header
    n_cols = len(header_row.cells)
    new_row = table.add_row()
    cells = new_row.cells

    if len(cells) != n_cols:
        print(f"  [WARN] New row has {len(cells)} cells, expected {n_cols}")

    # Заполняем ячейки шаблонными переменными
    for i, content in enumerate(ROW_CELL_CONTENTS):
        if i >= len(cells):
            break
        # Очищаем существующее содержимое ячейки
        cell = cells[i]
        # Удаляем все параграфы
        for para in cell.paragraphs:
            p_element = para._element
            p_element.getparent().remove(p_element)
        # Добавляем один новый параграф с нужным текстом
        new_para = cell.add_paragraph(content)

    return len(rows_to_remove)


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))

    # 1. Замены в шапке
    header_count = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                header_count += process_paragraphs(cell.paragraphs)
    print(f"  Header replacements: {header_count}")

    # 2. Перестраиваем таблицу транзакций
    tx_table = find_transactions_table(doc)
    if tx_table is None:
        print("[ERROR] Transactions table not found")
        return 1
    removed = rebuild_transactions_table(tx_table)
    print(f"  Removed {removed} transaction rows, added 1 template row with Jinja loop")

    # 3. Сохраняем
    doc.save(str(TARGET))
    print(f"[OK] Saved: {TARGET.name}")
    print()
    print("Now run:")
    print("  python scripts\\render_test_bank_statement.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())