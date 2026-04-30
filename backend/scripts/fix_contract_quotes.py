"""
Pack 8.5 — поправка кавычек в шаблоне договора.

После первой правки строка стала:
    «{{ contract.sign_date_str }}»

Но кавычки уже вшиты в саму sign_date_str (поправили в _format_date_ru),
поэтому нужно убрать наружные кавычки из шаблона.

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python scripts\\fix_contract_quotes.py
"""
import sys
import io
import shutil
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

BACKEND_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = BACKEND_ROOT.parent / "templates" / "docx" / "contract_template.docx"

if not TEMPLATE_PATH.exists():
    print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
    sys.exit(1)

# Backup
backup = TEMPLATE_PATH.with_suffix(".docx.bak2")
shutil.copy(TEMPLATE_PATH, backup)
print(f"[OK] Backup saved: {backup}")

from docx import Document

doc = Document(str(TEMPLATE_PATH))

# Поправляем строку: убираем «» вокруг {{ contract.sign_date_str }}
OLD_FRAGMENT = "«{{ contract.sign_date_str }}»"
NEW_FRAGMENT = "{{ contract.sign_date_str }}"

found = False
for para in doc.paragraphs:
    text = para.text
    if OLD_FRAGMENT in text:
        # Заменяем во всём параграфе. Сохраняем форматирование первого run.
        new_text = text.replace(OLD_FRAGMENT, NEW_FRAGMENT)
        if len(para.runs) > 0:
            first_run = para.runs[0]
            for r in para.runs:
                r.text = ""
            first_run.text = new_text
            found = True
            print(f"[OK] Patched paragraph")
            break

if not found:
    print("[WARN] Pattern «{{ contract.sign_date_str }}» not found.")
    print("       Maybe the template was already fixed?")
    sys.exit(1)

doc.save(str(TEMPLATE_PATH))
print("[DONE] Template saved")

# Verification
print("\n=== Verification ===")
doc = Document(str(TEMPLATE_PATH))
for i, para in enumerate(doc.paragraphs[:5]):
    text = para.text
    if text.strip():
        print(f"P{i}: {text}")
