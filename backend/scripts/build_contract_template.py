"""
Превращает реальный договор СК10/Алиев в шаблон docxtpl с переменными {{ }}.

Берёт _contract_original.docx → заменяет конкретные значения на jinja-теги →
сохраняет как contract_template.docx.

Запуск:
    cd D:\\VISA\\visa_kit\\backend
    python scripts\\build_contract_template.py

После запуска contract_template.docx можно открыть в Word и проверить —
это будет нормальный читаемый документ, только с {{ переменными }} вместо
конкретных Алиева/СК10/300000.
"""
import sys
import io
from pathlib import Path

# UTF-8 для вывода
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

# Пути относительно корня проекта
PROJECT_ROOT = Path(__file__).resolve().parents[2]  # D:/VISA/visa_kit
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_contract_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "contract_template.docx"

# ============================================================================
# Список замен
# Порядок ВАЖЕН — длинные строки сначала, чтобы не сработали короткие подстроки.
# ============================================================================

REPLACEMENTS = [
    # === Шапка договора ===
    ("№004/09/25", "№{{ contract.number }}"),
    ('"05" сентября 2025г.', '{{ fmt_date_long_ru(contract.sign_date) }}'),
    ('"05" сентября 2025 г.', '{{ fmt_date_long_ru(contract.sign_date) }}'),

    # ВАЖНО: г. Ростов-на-Дону в шапке (ГОРОД ПОДПИСАНИЯ) и в почтовом адресе
    # компании (ПОЧТОВЫЙ ГОРОД) — это РАЗНЫЕ значения. Поэтому НЕ делаем
    # глобальную замену "г. Ростов-на-Дону". Заменяем ровно ту строку из шапки.
    # Шапка договора начинается с "г. Ростов-на-Дону\t\t..." с табами и пробелами.
    # Пробуем оба варианта: с разным числом пробелов вокруг.

    # === Кандидат / исполнитель ===
    ("Алиев Джафар Надирович", "{{ applicant.full_name_native }}"),
    ("Алиев Д.Н.", "{{ applicant.initials_native }}"),
    ("республики Азербайджана", "республики {{ applicant.nationality_ru_genitive }}"),
    ("Паспорт AZE C01366076", "Паспорт {{ applicant.passport_country_code }} {{ applicant.passport_number }}"),
    ("выдан 24.03.2017г. МИД Азербайджана", "выдан {{ fmt_date_ru(applicant.passport_issue_date) }}г. {{ applicant.passport_issuer }}"),
    ("ИНН 230217957801", "ИНН {{ applicant.inn }}"),

    # Адрес исполнителя — теперь две части (разбит на 2 строки в файле)
    ("352919, Краснодарский край, г. Армавир, ", "{{ applicant.home_address_line1 }} "),
    ("ул. 11-я Линия, д. 31 кв. 2", "{{ applicant.home_address_line2 }}"),

    # === Компания / заказчик ===
    ('Общество с ограниченной ответственностью "Строительная компания СК10"', "{{ company.full_name_ru }}"),
    ("ООО «Строительная компания СК10»", "{{ company.short_name }}"),
    ("Тараскина Юрия Александровича", "{{ company.director_full_name_genitive_ru }}"),
    ("Тараскин Ю.А.", "{{ company.director_short_ru }}"),
    ("Генерального директора", "{{ company.director_position_ru }}"),

    # ИНН/КПП компании
    ("ИНН 6168006148, КПП 616401001", "ИНН {{ company.tax_id_primary }}, КПП {{ company.tax_id_secondary }}"),

    # Юридический адрес компании — две строки
    ("Юрид. адрес: 344002, г. Ростов-на-дону, ", "Юрид. адрес: {{ company.legal_address_line1 }} "),
    ("ул. Московская, зд. 73/29а, ком. 7", "{{ company.legal_address_line2 }}"),

    # Почтовый адрес компании — две строки
    # ПЕРВЫЙ переход: было g. Ростов-на-Дону (теперь там {{ contract.sign_city }} из-за прошлой замены).
    # Откатываем и сразу заменяем правильно на postal_address_line1
    ("344022, г. {{ contract.sign_city }},", "{{ company.postal_address_line1 }}"),
    # На случай если переменная еще не вставилась (запуск с чистого исходника) — фолбэк
    ("344022, г. Ростов-на-Дону,", "{{ company.postal_address_line1 }}"),
    ("ул. Нижнебульварная  6, БЦ «5 морей»", "{{ company.postal_address_line2 }}"),

    # Банковские реквизиты компании
    ('Филиал "ЦЕНТРАЛЬНЫИ " БАНКА ВТБ (ПАО)', "{{ company.bank_name }}"),
    ("40702810206640002909", "{{ company.bank_account }}"),
    ("044525411", "{{ company.bank_bic }}"),
    ("30101810145250000411", "{{ company.bank_correspondent_account }}"),

    # === Должность и услуги ===
    ("инженера-геодезист (камеральщик)", "{{ position.title_ru }}"),

    # === Деньги ===
    (
        "300 000 (триста тысяч) рублей в месяц",
        "{{ fmt_money(contract.salary_rub) }} ({{ contract.salary_rub_words }}) рублей в месяц",
    ),

    # === Сроки ===
    ('"31" августа 2029 г.', '{{ fmt_date_long_ru(contract.end_date) }}'),

    # === Реквизиты исполнителя в Альфа-банке (его личный счёт) ===
    ("40803840441563809831", "{{ applicant.bank_account }}"),
    # Заодно чиним опечатку в банке (закрытая кавычка)
    ("АО «АЛЬФА-БАНК, г. Москва", "{{ applicant.bank_name }}"),
    # БИК и К\с — НЕ заменяем, потому что у других клиентов могут быть другие банки.
    # Вместо этого делаем поля applicant.bank_bic и applicant.bank_correspondent_account
    ("044525593", "{{ applicant.bank_bic }}"),
    ("30101810200000000593", "{{ applicant.bank_correspondent_account }}"),

    # === Шапка с городом подписания — оставляем последним, чтобы не зацепить
    # почтовый адрес. К этому моменту почтовый уже заменён на postal_address_*.
    ("г. Ростов-на-Дону", "г. {{ contract.sign_city }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """
    Заменяет old на new в параграфе целиком.

    Если строка разбита по нескольким `runs` (что часто бывает в Word при разном
    форматировании внутри строки) — мы аккуратно «склеиваем» runs в первый,
    сохраняя его форматирование. Остальные runs очищаем.

    Возвращает True если что-то заменилось.
    """
    full_text = paragraph.text
    if old not in full_text:
        return False

    new_text = full_text.replace(old, new)

    if not paragraph.runs:
        return False

    # Все символы кладём в первый run, остальные обнуляем
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def main():
    if not SOURCE.exists():
        print(f"ОШИБКА: исходный файл не найден: {SOURCE}")
        print("Скопируйте ваш реальный договор как _contract_original.docx")
        return 1

    print(f"Читаем исходник: {SOURCE.name}")
    doc = Document(str(SOURCE))
    total = 0

    # Параграфы верхнего уровня
    total += process_paragraphs(doc.paragraphs)

    # Параграфы внутри таблиц (включая вложенные)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            total += process_paragraphs(ncell.paragraphs)

    # Заголовки и футеры
    for section in doc.sections:
        if section.header:
            total += process_paragraphs(section.header.paragraphs)
        if section.footer:
            total += process_paragraphs(section.footer.paragraphs)

    doc.save(str(TARGET))
    print(f"Сохранено: {TARGET.name}")
    print(f"Сделано замен: {total}")
    print()
    print("Теперь откройте contract_template.docx в Word и проверьте:")
    print("- все имена/числа/даты заменены на {{ переменные }}")
    print("- форматирование (шрифты, таблицы) сохранилось")
    print("- если что-то пошло не так — исходник цел в _contract_original.docx")
    return 0


if __name__ == "__main__":
    sys.exit(main())