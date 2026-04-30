"""
Превращает _act_original.docx в act_template.docx с переменными {{ }}.

Запуск:
    python scripts/build_act_template.py

Перед запуском:
- _act_original.docx должен лежать в templates/docx/
- Word должен быть закрыт (иначе PermissionError)
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE = PROJECT_ROOT / "templates" / "docx" / "_act_original.docx"
TARGET = PROJECT_ROOT / "templates" / "docx" / "act_template.docx"


REPLACEMENTS = [
    # === Шапка акта ===
    ("АКТ №1/26", "АКТ №{{ act.sequence_number }}/{{ act.year_suffix }}"),
    ("Договор №004/09/25 от 05.09.2025г.",
     "Договор №{{ contract.number }} от {{ fmt_date_ru(contract.sign_date) }}г."),

    # Город (оставляем Ростов-на-Дону, у СК10 город всегда такой)
    # Дата подписания акта = последний день месяца
    ("«31» января 2026г.", "{{ fmt_date_long_ru(act.document_date) }}"),

    # === Преамбула ===
    ("Алиев Джафар Надирович", "{{ applicant.full_name_native }}"),
    ('Общество с ограниченной ответственностью "Строительная компания СК10"',
     "{{ company.full_name_ru }}"),
    ("Тараскина Юрия Александровича", "{{ company.director_full_name_genitive_ru }}"),
    ("Генерального директора", "{{ company.director_position_ru }}"),

    # Гражданство
    ("республики Азербайджана", "республики {{ applicant.nationality_ru_genitive }}"),

    # === Период оказания услуг ===
    ("с 01.01.2026г. по 31.01.2026г.",
     "с {{ fmt_date_ru(act.period_start) }}г. по {{ fmt_date_ru(act.period_end) }}г."),

    # === Должность в родительном падеже ===
    ("инженера-геодезист (камеральщик)", "{{ position.title_ru_genitive }}"),

    # === Сумма ===
    (
        "300 000 рублей 00 копеек (триста тысяч рублей 00 копеек)",
        "{{ fmt_money(act.salary_rub) }} рублей 00 копеек ({{ act.salary_rub_words }} рублей 00 копеек)",
    ),

    # === Подписи ===
    ("Тараскин Ю.А.", "{{ company.director_short_ru }}"),
    ("Алиев Д.Н.", "{{ applicant.initials_native }}"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def process_paragraphs(paragraphs) -> int:
    count = 0
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                count += 1
    return count


def main():
    if not SOURCE.exists():
        print(f"[ERROR] Source not found: {SOURCE}")
        return 1

    print(f"Reading: {SOURCE.name}")
    doc = Document(str(SOURCE))
    total = process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_paragraphs(cell.paragraphs)

    doc.save(str(TARGET))
    print(f"Saved: {TARGET.name}")
    print(f"Replacements: {total}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
