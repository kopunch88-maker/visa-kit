"""
Универсальный экстрактор текста из DOCX.

Использование:
    python scripts/extract_docx_text.py <имя_файла_без_пути>

Файл должен лежать в templates/docx/.
Вывод сохраняется в scripts/_dumps/<имя>.txt в UTF-8.

Примеры:
    python scripts/extract_docx_text.py _act_original.docx
    python scripts/extract_docx_text.py _bank_statement_original.docx
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

BACKEND_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = BACKEND_ROOT.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates" / "docx"
DUMPS_DIR = BACKEND_ROOT / "scripts" / "_dumps"


def main():
    if len(sys.argv) != 2:
        print("Usage: python scripts/extract_docx_text.py <filename>")
        return 1

    filename = sys.argv[1]
    src = TEMPLATES_DIR / filename
    if not src.exists():
        print(f"File not found: {src}")
        return 1

    DUMPS_DIR.mkdir(parents=True, exist_ok=True)
    out = DUMPS_DIR / (filename.rsplit(".", 1)[0] + ".txt")

    doc = Document(str(src))
    lines = []

    lines.append(f"=== {filename} ===")
    lines.append(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    lines.append("")

    lines.append("=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            lines.append(f"[P{i}] {p.text}")

    lines.append("")
    lines.append("=== TABLES ===")
    for ti, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        lines.append(f"\n--- Table {ti}: {rows} rows, {cols} cols ---")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    if len(text) > 400:
                        text = text[:400] + "..."
                    lines.append(f"  [T{ti}/R{ri}/C{ci}] {text}")

    out.write_text("\n".join(lines), encoding="utf-8")
    print(f"Saved to: {out}")
    print(f"Total lines: {len(lines)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())