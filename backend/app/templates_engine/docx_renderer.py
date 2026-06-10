"""
Рендер DOCX-шаблонов через docxtpl.

Bank statement рендерится особым способом: после стандартного docxtpl-рендера
(подставляет шапку с балансами и периодом), мы открываем результат через
python-docx, находим строку-образец с маркерами __TX_*__ и клонируем её
для каждой транзакции, заменяя маркеры на реальные данные.

Pack 16.4 changes:
- _replace_markers_in_tr теперь поддерживает многострочные значения —
  если описание содержит '\\n' (например зарплата от компании), для
  каждой строки создаётся отдельный <w:p> в ячейке.
- Добавлена _remove_empty_paragraph_between_tables — убирает пустой
  параграф между таблицей операций и таблицей подписи, чтобы подпись
  поместилась сразу после операций (без перевода на 2-ю страницу
  если на 1-й есть место).

Pack 33.0 changes (10.05.2026):
- Постпроцессинг render_contract: блок «Адреса и реквизиты Сторон» всегда
  начинается с новой страницы через <w:pageBreakBefore/>. Word сам
  решает, где физически разорвать страницу — никаких пустых параграфов
  и пробелов. Применяется ко ВСЕМ шаблонам из contracts_registry
  (default + 7 company-specific). Идемпотентно: если флаг уже стоит,
  повторно не добавляется. Защита от пустой первой страницы: если
  заголовок реквизитов — самый первый непустой параграф документа,
  разрыв НЕ ставится. См. _apply_page_break_before_requisites.
"""

import io
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from sqlmodel import Session
import lxml.etree as etree

from app.models import Application, Applicant, Bank, Company, Position, SpainAddress
from .context import build_context


TEMPLATES_DIR = Path(__file__).resolve().parents[3] / "templates" / "docx"
REPO_ROOT = Path(__file__).resolve().parents[3]  # Pack 29.0: для resolve_contract_template_path

# Pack 29.0: реестр контрактных шаблонов
from .contracts_registry import (
    resolve_contract_template_path,
    is_template_slug_valid,
    COMPANY_INN_TO_SLUG,
    get_available_template_options,
)
from .employment_contracts_registry import (
    resolve_employment_contract_template_path,
)
from fastapi import HTTPException


class NeedsContractTemplateError(HTTPException):
    """
    Pack 29.0 — поднимается из render_contract когда у компании не выбран
    шаблон договора и ИНН не в COMPANY_INN_TO_SLUG. Frontend ловит 409
    и показывает модалку выбора шаблона.
    """
    def __init__(self, company):
        super().__init__(
            status_code=409,
            detail={
                "code": "NEEDS_CONTRACT_TEMPLATE",
                "message": (
                    f"Для компании '{company.short_name}' (id={company.id}) "
                    f"не выбран шаблон договора. Выберите шаблон в форме компании."
                ),
                "company_id": company.id,
                "company_short_name": company.short_name,
                "available_templates": get_available_template_options(),
            },
        )


class NeedsEmploymentContractTemplateError(HTTPException):
    """
    Pack 50.1-C/G — поднимается из render_employment_contract когда у компании
    не выбран шаблон Трудового договора (ни employment_contract_template_slug,
    ни fallback по ИНН в EMPLOYMENT_COMPANY_INN_TO_SLUG).

    Pack 50.1-G: payload теперь содержит available_templates — список шаблонов
    из EMPLOYMENT_CONTRACT_TEMPLATES_REGISTRY. Фронт открывает модалку
    ContractTemplatePickerModal с kind="employment", менеджер выбирает шаблон,
    мы сохраняем его в company.employment_contract_template_slug.
    """
    def __init__(self, company):
        from app.templates_engine.employment_contracts_registry import (
            get_available_employment_template_options,
        )
        super().__init__(
            status_code=409,
            detail={
                "code": "NEEDS_EMPLOYMENT_CONTRACT_TEMPLATE",
                "message": (
                    f"Для компании '{company.short_name}' (ИНН={company.tax_id_primary}) "
                    f"не выбран шаблон Трудового договора. Выберите шаблон в карточке "
                    f"компании или через модалку."
                ),
                "company_id": company.id,
                "company_short_name": company.short_name,
                "company_inn": company.tax_id_primary,
                "available_templates": get_available_employment_template_options(),
            },
        )

# ============================================================================
# Pack 50.1-H — Post-processor для замены шрифта в DOCX
# ============================================================================
#
# Используется в render_contract: если у компании задан contract_font_family,
# после рендера через docxtpl мы пробегаем по XML и заменяем все <w:rFonts>
# атрибуты ascii/hAnsi/cs/eastAsia на указанный шрифт.

_FONT_WHITELIST_PACK50_1_H = {
    "Times New Roman",
    "Arial",
    "Calibri",
    "Microsoft Sans Serif",
}


def _replace_fonts_in_docx(docx_bytes: bytes, font_name: str) -> bytes:
    """Pack 50.1-H — заменяет все шрифты в word/document.xml на font_name.

    Шаблон может содержать сотни тегов <w:rFonts w:ascii="..." w:hAnsi="..."
    w:cs="..." w:eastAsia="..."/>. Мы пробегаем по всему document.xml и
    заменяем значения этих 4 атрибутов на единый font_name.

    Если font_name не в whitelist — возвращает исходные байты без изменений
    (безопасный fallback на случай мусора в БД).
    """
    import io
    import re
    import zipfile

    if not font_name or font_name not in _FONT_WHITELIST_PACK50_1_H:
        return docx_bytes

    # Открываем zip, читаем document.xml
    with io.BytesIO(docx_bytes) as buf:
        with zipfile.ZipFile(buf, "r") as zin:
            doc_xml = zin.read("word/document.xml").decode("utf-8")
            other_files = {
                name: zin.read(name)
                for name in zin.namelist()
                if name != "word/document.xml"
            }

    # Регулярка ищет атрибуты w:ascii/w:hAnsi/w:cs/w:eastAsia внутри тега
    # <w:rFonts ...> и заменяет их значения.
    # Используем единый паттерн: w:(ascii|hAnsi|cs|eastAsia)="..."
    pattern = re.compile(r'(w:(?:ascii|hAnsi|cs|eastAsia))="[^"]*"')
    doc_xml_new = pattern.sub(rf'\1="{font_name}"', doc_xml)

    # Собираем zip обратно
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        zout.writestr("word/document.xml", doc_xml_new.encode("utf-8"))
        for name, data in other_files.items():
            zout.writestr(name, data)

    return out_buf.getvalue()


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


# ============================================================================
# Pack 33.0 — page break before "Адреса и реквизиты Сторон"
# ============================================================================

# Регэксп ищет заголовок раздела реквизитов в любом виде, который встречался
# в шаблонах: "8. Адреса и реквизиты Сторон", "Адреса и реквизиты сторон",
# "Реквизиты и адреса сторон", "Реквизиты Сторон" и т.д. Не зависит от
# номера раздела (8., 9., и т.д.) и регистра.
_REQUISITES_HEADING_RE = re.compile(
    r"(адреса\s+и\s+реквизиты|реквизиты\s+и\s+адреса|реквизиты\s+сторон)",
    re.IGNORECASE,
)


def _ensure_page_break_before(p_element) -> bool:
    """
    Добавляет <w:pageBreakBefore/> в pPr параграфа.
    Возвращает True если добавлено, False если уже было.
    pPr должен быть первым дочерним элементом <w:p>;
    pageBreakBefore должен идти в начале pPr (по схеме OOXML).
    """
    ppr = p_element.find("w:pPr", NS)
    if ppr is None:
        ppr = etree.Element(f"{W_NS}pPr")
        p_element.insert(0, ppr)

    if ppr.find("w:pageBreakBefore", NS) is not None:
        return False

    page_break = etree.Element(f"{W_NS}pageBreakBefore")
    ppr.insert(0, page_break)
    return True


def _force_left_align_in_table(table_element) -> int:
    """
    Pack 34.6 — пробегает по всем <w:p> внутри таблицы и насильно ставит
    <w:jc w:val="left"/>. Если jc отсутствует — создаёт.

    Используется для таблицы реквизитов сторон в разделе 8 договора:
    параграфы в шаблоне имеют jc=both (justify), и после Pack 34.5 (NBSP
    в адресах) Word слишком сильно растягивает оставшиеся обычные пробелы.

    Возвращает количество изменённых параграфов.
    Идемпотентно: повторный запуск не меняет уже left-параграфы.
    """
    modified = 0
    for p in table_element.findall(".//w:p", NS):
        ppr = p.find("w:pPr", NS)
        if ppr is None:
            ppr = etree.Element(f"{W_NS}pPr")
            p.insert(0, ppr)

        jc = ppr.find("w:jc", NS)
        if jc is None:
            jc = etree.SubElement(ppr, f"{W_NS}jc")
            jc.set(f"{W_NS}val", "left")
            modified += 1
        else:
            current = jc.get(f"{W_NS}val")
            if current != "left":
                jc.set(f"{W_NS}val", "left")
                modified += 1
    return modified


def _apply_page_break_before_requisites(docx_bytes: bytes) -> bytes:
    """
    Pack 33.0 — постпроцессинг отрендеренного договора.

    Находит первый параграф верхнего уровня документа (НЕ внутри таблицы),
    текст которого матчится на _REQUISITES_HEADING_RE («Адреса и реквизиты
    Сторон» и подобные), и добавляет ему <w:pageBreakBefore/>.

    Логика:
      1. Идём по прямым детям <w:body> (только <w:p>, не лезем в таблицы)
      2. Считаем сколько было непустых параграфов ДО найденного заголовка
      3. Если найденный заголовок — самый первый непустой параграф документа
         (count == 0), разрыв НЕ ставим (иначе пустая страница в начале)
      4. Если в pPr уже есть pageBreakBefore — ничего не делаем (идемпотентно)
      5. Если заголовок не найден — возвращаем bytes без изменений
         (например, шаблон с другой структурой или одностраничный)

    Все случаи безопасны: при ошибке/отсутствии заголовка возвращается
    исходный документ без модификаций.
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception:
        # Если по какой-то причине не парсится — возвращаем как есть
        return docx_bytes

    body = doc.element.body
    meaningful_paragraphs_before = 0
    target_p = None

    # Идём только по прямым детям body (параграфы верхнего уровня).
    # Таблицы пропускаем — заголовок раздела 8 в шаблонах всегда отдельный <w:p>,
    # а не строка таблицы. Если у кого-то иначе — фикс просто не сработает,
    # документ отдастся без модификации.
    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag != "p":
            continue

        # Собираем весь текст параграфа из всех <w:t>
        ts = child.findall(".//w:t", NS)
        text = "".join(t.text or "" for t in ts)
        stripped = text.strip()

        if not stripped:
            # Пустые параграфы (включая остатки от подстановки) пропускаем
            # и НЕ считаем их «значимыми». Это защищает от ситуации,
            # когда шаблон начинается с пустых отступов.
            continue

        if _REQUISITES_HEADING_RE.search(stripped):
            target_p = child
            break

        meaningful_paragraphs_before += 1

    if target_p is None:
        # Заголовок не нашли — возвращаем как есть
        return docx_bytes

    if meaningful_paragraphs_before == 0:
        # Заголовок — самый первый непустой параграф, разрыв создаст
        # пустую первую страницу. Не трогаем.
        return docx_bytes

    _ensure_page_break_before(target_p)

    # Pack 34.6 — найти первую таблицу ПОСЛЕ заголовка реквизитов
    # и сделать все её параграфы left-aligned. Это убирает justify-артефакт
    # «растянутых пробелов» когда NBSP (Pack 34.5) уменьшил число точек
    # разрыва на строке.
    requisites_table = None
    found_heading = False
    for child in body.iterchildren():
        if child is target_p:
            found_heading = True
            continue
        if not found_heading:
            continue
        if etree.QName(child).localname == "tbl":
            requisites_table = child
            break

    if requisites_table is not None:
        _force_left_align_in_table(requisites_table)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _stdr_strip_empty_rows(doc) -> None:
    """Pack 50.27 — удаляет пустые строки данных в таблицах СТД-Р.

    Таблица 0 (после 2019): данные с 5-й строки (индекс 4) — заголовки R0..R3.
    Таблица 1 (до 2019): данные с 3-й строки (индекс 2) — заголовки R0..R1.
    Строка считается пустой, если ВСЕ её ячейки без текста.
    Документ заканчивается на последней реальной записи (нет пустых страниц).
    """
    specs = [(0, 4), (1, 2)]  # (table_index, data_start_row)
    for tbl_idx, data_start in specs:
        if tbl_idx >= len(doc.tables):
            continue
        table = doc.tables[tbl_idx]
        for row in list(table.rows[data_start:]):
            if not any(c.text.strip() for c in row.cells):
                table._tbl.remove(row._tr)


def _render(template_name: str, context: dict, post_process=None) -> bytes:
    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    template = DocxTemplate(str(template_path))
    template.render(context)
    if post_process is not None:  # Pack 50.27
        post_process(template.docx)

    buffer = io.BytesIO()
    template.save(buffer)
    return buffer.getvalue()


def _render_from_repo_path(repo_relative_path: str, context: dict) -> bytes:
    """
    Pack 29.0 — рендер шаблона по пути относительно корня репо
    (например 'templates/docx/contracts/by_company/sk10/contract_template.docx').
    Используется для контрактных шаблонов, выбираемых через contracts_registry.
    """
    template_path = REPO_ROOT / repo_relative_path
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    template = DocxTemplate(str(template_path))
    template.render(context)

    buffer = io.BytesIO()
    template.save(buffer)
    return buffer.getvalue()


def render_contract(application: Application, session: Session) -> bytes:
    """
    Pack 29.0/29.3.1 — выбор шаблона по company.contract_template_slug:
      1. Если slug задан и валиден → шаблон из contracts_registry.
      2. Иначе если ИНН компании в COMPANY_INN_TO_SLUG → fallback по ИНН.
      3. Иначе → 409 NEEDS_CONTRACT_TEMPLATE (фронт показывает модалку).

    Pack 29.3.1 fix: Application не имеет relationship 'company',
    только foreign key company_id. Загружаем Company явно через session.

    Pack 33.0: после рендера применяем _apply_page_break_before_requisites —
    раздел «Адреса и реквизиты Сторон» всегда стартует с новой страницы.
    Применяется ко всем шаблонам из contracts_registry детерминированно.
    """
    if not application.company_id:
        raise ValueError(
            f"Application id={application.id} has no company_id assigned"
        )
    company = session.get(Company, application.company_id)
    if not company:
        raise ValueError(
            f"Company id={application.company_id} not found for "
            f"application id={application.id}"
        )

    if not is_template_slug_valid(getattr(company, 'contract_template_slug', None)):
        if (company.tax_id_primary or '') not in COMPANY_INN_TO_SLUG:
            raise NeedsContractTemplateError(company)

    context = build_context(application, session)
    # Pack 41.0-G — для договора override паспортных полей на выбранный
    # менеджером passport_id_for_ru_docs (если задан). Для всех остальных
    # русских/испанских форм паспорт остаётся primary через скаляр-зеркало.
    # Pack 41.0-G fix2 — passports[] хранит issue_date как ISO-строку,
    # нужно конвертить в date перед передачей в fmt_date_ru().
    from app.services.applicant_passports import get_passport_dict_for_ru_docs
    from datetime import date as _date
    _applicant = application.applicant
    _ru_passport = get_passport_dict_for_ru_docs(_applicant)
    if _ru_passport.get("number"):
        from app.templates_engine.context import (
            _parse_passport,
            _resolve_passport_issuer_for_template_from_dict,
            fmt_date_ru,
        )
        # Конвертим issue_date: либо ISO-string из passports[], либо уже date из legacy
        _raw_issue_date = _ru_passport.get("issue_date")
        if isinstance(_raw_issue_date, str) and _raw_issue_date:
            try:
                _issue_date = _date.fromisoformat(_raw_issue_date)
            except ValueError:
                _issue_date = None
        else:
            _issue_date = _raw_issue_date  # уже date или None
        _pdata = _parse_passport(_ru_passport["number"], _applicant.nationality)
        context["applicant"]["passport_number"] = _ru_passport["number"]
        context["applicant"]["passport_series"] = _pdata["series"]
        context["applicant"]["passport_number_only"] = _pdata["number_only"]
        context["applicant"]["passport_formatted"] = _pdata["formatted"]
        context["applicant"]["passport_issue_date"] = _issue_date
        context["applicant"]["passport_issue_date_str"] = fmt_date_ru(_issue_date)
        context["applicant"]["passport_issuer"] = _resolve_passport_issuer_for_template_from_dict(
            _ru_passport, _applicant.nationality
        )
    relative_path = resolve_contract_template_path(company)
    rendered = _render_from_repo_path(relative_path, context)
    # Pack 50.1-H — если у компании задан contract_font_family,
    # подменяем все шрифты в DOCX через post-processor.
    font_family = getattr(company, "contract_font_family", None)
    if font_family:
        rendered = _replace_fonts_in_docx(rendered, font_family)
    # Pack 33.0
    return _apply_page_break_before_requisites(rendered)


def render_act(application: Application, session: Session, sequence_number: int) -> bytes:
    context = build_context(application, session)
    months = context.get("monthly_documents", [])
    target = next((m for m in months if m["sequence_number"] == sequence_number), None)
    if not target:
        raise ValueError(f"No monthly document with sequence {sequence_number}")
    context["act"] = target
    return _render("act_template.docx", context)


def render_invoice(application: Application, session: Session, sequence_number: int) -> bytes:
    context = build_context(application, session)
    months = context.get("monthly_documents", [])
    target = next((m for m in months if m["sequence_number"] == sequence_number), None)
    if not target:
        raise ValueError(f"No monthly document with sequence {sequence_number}")
    context["invoice"] = target
    return _render("invoice_template.docx", context)


def render_employer_letter(application: Application, session: Session) -> bytes:
    context = build_context(application, session)
    return _render("employer_letter_template.docx", context)


def render_employer_letter_naim(application: Application, session: Session) -> bytes:
    """Pack 50.11-B — Письмо работодателя для НАЙМА (трудовой договор).

    Тот же контекст что и письмо самозанятого (build_context), но другой
    шаблон: 'работает в компании по трудовому договору', бессрочно, без
    фраз про независимого подрядчика.
    """
    context = build_context(application, session)
    return _render("employer_letter_naim_template.docx", context)


def render_cv(application: Application, session: Session) -> bytes:
    context = build_context(application, session)
    return _render("cv_template.docx", context)


def render_tech_opinion(application: Application, session: Session) -> bytes:
    """Pack 40.0 — Техническое заключение о дистанционном характере деятельности."""
    context = build_context(application, session)
    return _render("tech_opinion_template.docx", context)


def render_business_trip_order(application: Application, session: Session) -> bytes:
    """Pack 50.7-C — Приказ Т-9 о направлении работника в командировку (найм)."""
    context = build_context(application, session)
    return _render("business_trip_order_template.docx", context)

def render_ndfl_2(application: Application, session: Session) -> bytes:
    """Pack 50.8-B — Справка о доходах и суммах налога физического лица (2-НДФЛ).

    Форма КНД 1175018, приказ ФНС России от 19.09.2023 № ЕД-7-11/649@.

    Рендерится из общего build_context (включая блок 'ndfl_2'), один шаблон
    на всех — форма ФНС универсальна.
    """
    context = build_context(application, session)
    return _render("ndfl_2_template.docx", context)


def render_stdr(application: Application, session: Session) -> bytes:
    """Pack 50.9-B — Сведения о трудовой деятельности из СФР (СТД-Р).

    Форма по приказу Минтруда РФ от 20.01.2020 № 23н.

    Контекст 'stdr' собирается через build_stdr_context (отдельная функция,
    не часть build_context, т.к. требует position для okz_code).
    Парсит applicant.work_history и заполняет 15 слотов Таблицы 1 (события
    с 01.01.2020+) и 8 слотов Таблицы 2 (периоды до 31.12.2019).
    """
    from .context import build_stdr_context

    if not application.applicant_id or not application.company_id or not application.position_id:
        raise ValueError(
            f"Application id={application.id} not fully assigned (need applicant/company/position)"
        )
    applicant = session.get(Applicant, application.applicant_id)
    company = session.get(Company, application.company_id)
    position = session.get(Position, application.position_id)
    if not applicant or not company or not position:
        raise ValueError(f"Application id={application.id}: applicant/company/position not found")

    # Общий контекст (на случай если в шаблоне нужны applicant.*/company.* как-то ещё)
    context = build_context(application, session)
    # Перезаписываем/добавляем блок stdr специфичным контекстом
    stdr_ctx = build_stdr_context(application, applicant, company, position, session)
    context["stdr"] = stdr_ctx
    # Также пробрасываем applicant_stdr (uppercase ФИО + birth_date_long + snils)
    # как алиас на верхнем уровне для удобства шаблона
    context["applicant_stdr"] = stdr_ctx["applicant"]
    return _render("stdr_template.docx", context, post_process=_stdr_strip_empty_rows)  # Pack 50.27


def render_soo(application: Application, session: Session) -> bytes:
    """Pack 50.12-B — Свидетельство об отъезде (СОО).

    Справка СФР по договору РФ-Испания о соцобеспечении 1994/1995.
    Контекст 'soo' собирается через build_soo_context (отдельная функция,
    т.к. требует position + spain_address и пишет soo_number в БД).
    """
    from .context import build_soo_context

    if not application.applicant_id or not application.company_id or not application.position_id:
        raise ValueError(
            f"Application id={application.id} not fully assigned (need applicant/company/position)"
        )
    applicant = session.get(Applicant, application.applicant_id)
    company = session.get(Company, application.company_id)
    position = session.get(Position, application.position_id)
    if not applicant or not company or not position:
        raise ValueError(f"Application id={application.id}: applicant/company/position not found")

    spain_address = (
        session.get(SpainAddress, application.spain_address_id)
        if application.spain_address_id else None
    )

    context = build_context(application, session)
    soo_ctx = build_soo_context(application, applicant, company, position, spain_address, session)
    context["soo"] = soo_ctx
    return _render("soo_template.docx", context)



def render_payslip(application: Application, session: Session, month_idx: int) -> bytes:
    """Pack 50.10-B — Расчётный листок за один из 3 предыдущих месяцев.

    Args:
        application: Application
        session: Session
        month_idx: 0, 1 или 2 (0 = самый ранний месяц, 2 = последний)

    anchor = application.stdr_issue_date or date.today()
    Месяцы рассчитываются как 3 предыдущих относительно anchor.

    Контекст 'payslip' собирается через build_payslip_context.
    """
    from .context import build_payslip_context

    if not application.applicant_id or not application.company_id or not application.position_id:
        raise ValueError(
            f"Application id={application.id} not fully assigned (need applicant/company/position)"
        )
    applicant = session.get(Applicant, application.applicant_id)
    company = session.get(Company, application.company_id)
    position = session.get(Position, application.position_id)
    if not applicant or not company or not position:
        raise ValueError(f"Application id={application.id}: applicant/company/position not found")

    context = build_context(application, session)
    payslip_ctx = build_payslip_context(
        application, applicant, company, position, month_idx, session
    )
    context["payslip"] = payslip_ctx["payslip"]
    return _render("payslip_template.docx", context)



def render_employment_contract(application: Application, session: Session) -> bytes:
    """Pack 50.1-C — Трудовой договор (найм).

    Резолвит шаблон по ИНН компании через employment_contracts_registry.
    Если для компании нет шаблона — поднимает 409 NEEDS_EMPLOYMENT_CONTRACT_TEMPLATE.
    """
    if not application.company_id:
        raise ValueError(
            f"Application id={application.id} has no company_id assigned"
        )
    company = session.get(Company, application.company_id)
    if not company:
        raise ValueError(
            f"Company id={application.company_id} not found for "
            f"application id={application.id}"
        )

    template_path = resolve_employment_contract_template_path(company)
    if template_path is None:
        raise NeedsEmploymentContractTemplateError(company)

    context = build_context(application, session)
    rendered = _render_from_repo_path(template_path, context)
    # Pack 50.1-G — если у компании задан employment_contract_font_family,
    # подменяем все шрифты в DOCX через тот же post-processor что у
    # обычного договора (Pack 50.1-H).
    font_family = getattr(company, "employment_contract_font_family", None)
    if font_family:
        rendered = _replace_fonts_in_docx(rendered, font_family)
    return rendered



def _resolve_bank_statement_template_path(
    application: Application,
    session: Session,
) -> "Path":
    """
    Pack 47.0: резолв шаблона выписки по applicant.bank_id.
    Pack 52:   для Альфы переключение v1 (legacy, Трофимова) ↔ v2 (Ч/Б, Агеева)
               по application.bank_template_legacy_v1.

    Логика:
      1. Если у applicant есть bank_id с собственным шаблоном
         (templates/docx/bank_statement_template_<bik>.docx) — используем его.
         У Сбера/ТБанка/ВТБ свои шаблоны и v1/v2 на них не влияет.
      2. Иначе — Альфа default. Выбор v1 vs v2 по
         application.bank_template_legacy_v1:
           False (default для новых) → bank_statement_template_v2.docx (если файл есть)
           True (миграция выставила для существующих) → bank_statement_template.docx (v1)
    """
    def _alfa_default() -> "Path":
        # Pack 52: v1/v2 переключение для Альфы
        use_v1 = bool(getattr(application, "bank_template_legacy_v1", True))
        if not use_v1:
            v2 = TEMPLATES_DIR / "bank_statement_template_v2.docx"
            if v2.exists():
                return v2
        return TEMPLATES_DIR / "bank_statement_template.docx"

    if not application.applicant_id:
        return _alfa_default()

    applicant = session.get(Applicant, application.applicant_id)
    if applicant is None or not applicant.bank_id:
        return _alfa_default()

    bank = session.get(Bank, applicant.bank_id)
    if bank is None or not bank.bik:
        return _alfa_default()

    # Pack 54: для конкретного банка тоже поддерживаем _v2.docx суффикс
    # (Sber v2 = bank_statement_template_044525225_v2.docx — Ч/Б + подпись Кирьянова + печать)
    use_v2 = not bool(getattr(application, "bank_template_legacy_v1", True))
    if use_v2:
        candidate_v2 = TEMPLATES_DIR / f"bank_statement_template_{bank.bik}_v2.docx"
        if candidate_v2.exists():
            return candidate_v2

    candidate = TEMPLATES_DIR / f"bank_statement_template_{bank.bik}.docx"
    if candidate.exists():
        return candidate

    return _alfa_default()


def render_bank_statement(
    application: Application,
    session: Session,
    *,
    for_translation: bool = False,  # Pack 53
) -> bytes:
    """
    Двухфазный рендер:
    1. docxtpl подставляет шапку (период, балансы) через Jinja
    2. python-docx клонирует строку-образец таблицы для каждой транзакции

    Pack 47.0: шаблон выбирается по applicant.bank_id (см. _resolve_bank_statement_template_path).
    """
    template_path = _resolve_bank_statement_template_path(application, session)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    context = build_context(application, session)
    # Pack 41.0-K — override паспорта для русских + внутреннего теперь
    # делается централизованно в build_context, отдельный блок здесь
    # больше не нужен (Pack 41.0-H/I откачены).
    bank_data = context.get("bank", {})
    transactions = bank_data.get("transactions", [])

    # === ФАЗА 1: рендер шапки через docxtpl ===
    template = DocxTemplate(str(template_path))
    template.render(context)
    buffer = io.BytesIO()
    template.save(buffer)
    buffer.seek(0)

    # === ФАЗА 2: клонирование строк через python-docx ===
    doc = Document(buffer)

    # Находим таблицу транзакций (по маркеру в первой ячейке второй строки)
    tx_table = None
    template_row = None
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        second_row = table.rows[1]
        if second_row.cells and "__TX_DATE__" in second_row.cells[0].text:
            tx_table = table
            template_row = second_row
            break

    if tx_table is None or template_row is None:
        result_buffer = io.BytesIO()
        doc.save(result_buffer)
        return result_buffer.getvalue()

    # Клонируем образцовую строку для каждой транзакции
    template_tr_xml = template_row._tr
    parent = template_tr_xml.getparent()
    insert_position = list(parent).index(template_tr_xml)

    last_row = None
    for idx, tx in enumerate(transactions):
        new_tr = deepcopy(template_tr_xml)
        _replace_markers_in_tr(new_tr, tx)

        # Pack 16.5: серый фон у строк дохода (зарплата от компании).
        # Pack 47.3: серая заливка и жирная сумма — стиль эталона Алиева (Альфа).
        # Применяется ТОЛЬКО для дефолтного шаблона Альфы (v1 или v2). Сбер и
        # другие банки со своим шаблоном имеют белый фон строк-доходов.
        # Pack 52: v2 Ч/Б шаблон тоже считается "default" — gray shading и
        # bold у поступлений сохраняются как в эталоне.
        _is_default_template = template_path.name in (
            "bank_statement_template.docx",
            "bank_statement_template_v2.docx",
        )
        amount = tx.get("amount")
        if amount is not None and _is_default_template:
            try:
                amount_val = float(amount)
            except (TypeError, ValueError):
                amount_val = 0
            if amount_val > 0:
                _apply_gray_shading_to_row(new_tr)
                # Pack 25.0: жирная сумма у поступлений (как в эталонной выписке Алиева)
                _apply_bold_to_amount_cell(new_tr)
                # Pack 25.3: НЕ применяем tcMar — у Алиева его нет.
                # Воздух в серых ячейках обеспечивается через spacing
                # последнего параграфа описания (before=40 after=40),
                # см. _replace_marker_with_multiline.

        # Pack 16.5b: <w:cantSplit/> — запрет разрыва строки между страницами.
        _set_cant_split(new_tr)

        parent.insert(insert_position + idx, new_tr)
        last_row = new_tr

    # Удаляем оригинальную строку-образец
    parent.remove(template_tr_xml)

    # Pack 16.5c: keepNext на последнюю операцию + параграф между таблицами,
    # чтобы подпись не оставалась одна на странице. Если последняя операция
    # не помещается с подписью на 1-й странице — обе уйдут на 2-ю вместе.
    if last_row is not None:
        _set_keep_next_on_row(last_row)
        # Pack 25.1: возвращаем нижнюю границу таблицы — Pack 25.0 убрал
        # <w:bottom> со всех строк (чтобы не было двойной линии между ними).
        # Но последняя строка должна закрывать таблицу снизу, как в Алиеве.
        _add_bottom_border_to_row(last_row)
    _set_keep_next_on_paragraph_between_tables(doc)

    # Pack 47.15: ФАЗА 3 — замена маркера __EP_BADGE__ на runtime-сгенерированную
    # картинку плашки ЭП. Применяется только для Sber-шаблона (только он имеет
    # этот маркер). Для Альфы и других шаблонов без маркера — no-op.
    _replace_ep_badge_marker(doc, bank_data)

    # Pack 52: ФАЗА 3.5 — для v2-шаблона Альфы вставляем 3 PNG-печати
    # (подпись Агеевой + штамп ДО + круглая печать «Альфа-Банк») в маркеры
    # __STAMP_SIGNATURE__ / __STAMP_EMPLOYEE__ / __STAMP_BANK__. Для других
    # шаблонов (включая v1) — no-op, маркеров нет.
    # Pack 52/54: диспетчер по template name — Альфа vs Сбер.
    # mode="markers_only" при подготовке к переводу (без PNG, только чистка маркеров).
    _v2_name = template_path.name
    _v2_mode = "markers_only" if for_translation else "full"
    if _v2_name == "bank_statement_template_v2.docx":
        # Pack 52 — Альфа v2
        _insert_v2_signature_images(doc, mode=_v2_mode)
    elif _v2_name.endswith("_v2.docx") and "044525225" in _v2_name:
        # Pack 54 — Sber v2 (BIK 044525225)
        _insert_v2_sber_signatures(doc, mode=_v2_mode)

    # Pack 47.19: ФАЗА 4 — гарантия что каждая <w:tc> заканчивается на <w:p>.
    # OOXML schema требует это; Word иначе ругается "Обнаружено неоднозначное
    # сопоставление ячеек". Мои функции _strip_empty_paragraphs_before_tables
    # и _replace_ep_badge_marker могли удалить параграф который оказался
    # последним в ячейке.
    _ensure_paragraphs_at_tc_end(doc)

    result_buffer = io.BytesIO()
    doc.save(result_buffer)
    return result_buffer.getvalue()


def _ensure_paragraphs_at_tc_end(doc) -> None:
    """
    Pack 47.19: проходит по всем <w:tc> в документе и если последний дочерний
    элемент НЕ <w:p>, добавляет пустой <w:p/> в конец ячейки.

    Гарантирует валидность OOXML — Word требует чтобы каждая ячейка таблицы
    заканчивалась параграфом.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    for tc in doc.element.iter(_qn("w:tc")):
        children = list(tc)
        if not children:
            # Ячейка вообще пустая — добавляем <w:p/>
            tc.append(_OxmlElement("w:p"))
            continue
        last = children[-1]
        if last.tag != _qn("w:p"):
            # Последний элемент не параграф — добавляем пустой <w:p/>
            tc.append(_OxmlElement("w:p"))


# ============================================================================
# Pack 52 — v2 шаблон Альфы (Ч/Б + PNG-печати + PDF)
# ============================================================================

# Pack 52-fix17: helper для floating-якоря.
def _add_floating_picture(paragraph, png_path, width_mm, x_offset_mm=0, y_offset_mm=0):
    """
    Вставляет PNG как floating anchor.
    relativeFrom="column" для X, "paragraph" для Y, layoutInCell=0.

    1мм = 36000 EMU.
    """
    import random
    import lxml.etree as etree
    from docx.shared import Mm
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    run.add_picture(str(png_path), width=Mm(width_mm))
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    ext = inline.find(qn("wp:extent"))
    cx, cy = ext.get("cx"), ext.get("cy")
    graphic_el = inline.find(qn("a:graphic"))
    graphic_xml_str = etree.tostring(graphic_el, encoding="unicode")
    x_emu = int(x_offset_mm * 36000)
    y_emu = int(y_offset_mm * 36000)
    rid = random.randint(10000, 99999)
    anchor_xml = (
        f'<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
        f'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="{rid}" '
        f'behindDoc="0" locked="0" layoutInCell="0" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{rid}" name="Floating {rid}"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'{graphic_xml_str}'
        f'</wp:anchor>'
    )
    new_anchor = parse_xml(anchor_xml)
    drawing.remove(inline)
    drawing.append(new_anchor)


def _insert_v2_signature_images(doc, *, mode: str = "full") -> None:
    """
    Pack 52-fix17 + Pack 53: ВСЕ 3 картинки якорятся к параграфу прямоугольной печати (R0C2).
    Прямоугольная сидит inline на линии — её параграф находится РОВНО на линии,
    что делает её идеальной точкой отсчёта для подписи и круглой печати.

    Pack 53: параметр mode:
      - "full" (default): вставляем 3 PNG (подпись + штамп + круглая печать).
      - "markers_only": только чистим текст маркеров __STAMP_*__ из R0,
        PNG НЕ вставляем. Используется при подготовке docx под перевод
        (испанская версия выписки идёт без печатей; лейблы R1 переводятся
        в "(firma del empleado AO «ALFA-BANK»)" / ...).

    Стратегия:
      __STAMP_EMPLOYEE__ (R0C2):
        → INLINE 55мм, садится на линию (cell vAlign=BOTTOM в шаблоне)

      __STAMP_SIGNATURE__ (был в R0C0):
        → FLOATING 38мм, якорится к R0C2 параграфу
        → x_off=-60 column-relative (R0C2 column ~79мм → ~19мм от страничного
          левого края = на левой стороне страницы, как в эталоне)
        → y_off=-5 (5мм выше линии = пересекает её снизу)

      __STAMP_BANK__ (был в R0C3):
        → FLOATING 35мм, якорится к R0C2 параграфу
        → x_off=+80 column-relative (R0C2 column ~79мм + 80 = ~159мм от
          левого края = на правой стороне страницы)
        → y_off=-5 (5мм выше линии)

    Старые маркеры __STAMP_SIGNATURE__ (R0C0) и __STAMP_BANK__ (R0C3) просто
    очищаются — их параграфы остаются пустыми (визуально не мешает).
    """
    from docx.shared import Mm

    # Pack 53: markers_only — чистим только текст маркеров (без PNG).
    # Используется в render_bank_statement_for_translation: испанская версия
    # сигнатур-таблицы остаётся, но без печатей.
    if mode == "markers_only":
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if any(
                            m in p.text
                            for m in ("__STAMP_SIGNATURE__", "__STAMP_EMPLOYEE__", "__STAMP_BANK__")
                        ):
                            for r in list(p.runs):
                                r._element.getparent().remove(r._element)
        return

    assets_dir = TEMPLATES_DIR / "assets" / "v2"

    employee_png  = assets_dir / "stamp_employee.png"
    signature_png = assets_dir / "signature.png"
    bank_png      = assets_dir / "stamp_bank.png"

    # Этап 1. Найти параграф с маркером __STAMP_EMPLOYEE__ (= точка отсчёта)
    target_p = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "__STAMP_EMPLOYEE__" in p.text:
                        target_p = p
                        break
                if target_p:
                    break
            if target_p:
                break
        if target_p:
            break

    if target_p is None:
        # Шаблон без маркеров — no-op (другие шаблоны)
        return

    # Этап 2. Чистим target_p и вставляем прямоугольную INLINE
    for r in list(target_p.runs):
        r._element.getparent().remove(r._element)
    if employee_png.exists():
        try:
            run = target_p.add_run()
            run.add_picture(str(employee_png), width=Mm(55))
        except Exception as e:
            import logging
            logging.warning("Pack 52-fix17: не удалось inline %s: %s", employee_png.name, e)

    # Этап 3. Подпись floating, якорь = target_p, x_off=-60, y_off=+2
    # Pack 52-fix22: y_off с +1 на +2 — опустил на 1мм
    if signature_png.exists():
        try:
            _add_floating_picture(target_p, signature_png, 38, x_offset_mm=-60, y_offset_mm=2)
        except Exception as e:
            import logging
            logging.warning("Pack 52-fix22: не удалось floating signature: %s", e)

    # Этап 4. Круглая печать floating, якорь = target_p, x_off=+80, y_off=-5
    if bank_png.exists():
        try:
            _add_floating_picture(target_p, bank_png, 35, x_offset_mm=80, y_offset_mm=-5)
        except Exception as e:
            import logging
            logging.warning("Pack 52-fix17: не удалось floating bank: %s", e)

    # Этап 5. Чистим маркеры __STAMP_SIGNATURE__ и __STAMP_BANK__ (R0C0, R0C3)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "__STAMP_SIGNATURE__" in p.text or "__STAMP_BANK__" in p.text:
                        for r in list(p.runs):
                            r._element.getparent().remove(r._element)


def render_bank_statement_to_pdf(
    application: Application,
    session: Session,
    timeout_sec: int = 60,
) -> bytes:
    """
    Pack 52: рендерит банковскую выписку в PDF через LibreOffice headless.

    Поток:
      1. render_bank_statement(application, session) → DOCX bytes
      2. сохраняем во временный файл
      3. soffice --headless --convert-to pdf
      4. читаем PDF, возвращаем bytes

    ТРЕБОВАНИЕ: на сервере должен быть установлен LibreOffice (soffice в PATH).
    На Railway добавляется через nixpacks.toml или Aptfile:
        # nixpacks.toml
        [phases.setup]
        aptPkgs = ["libreoffice", "libreoffice-writer"]
    """
    import subprocess
    import tempfile  # Pack 52-fix1
    import os         # Pack 52-fix1

    docx_bytes = render_bank_statement(application, session)

    with tempfile.TemporaryDirectory(prefix="vk_pdf_") as tmpdir:
        docx_path = os.path.join(tmpdir, "statement.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # soffice сохраняет PDF рядом с input, имя = базовое имя input + .pdf
        try:
            result = subprocess.run(
                [
                    "soffice", "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True,
                timeout=timeout_sec,
            )
        except FileNotFoundError:
            raise RuntimeError(
                "Pack 52: LibreOffice (soffice) не найден в PATH. "
                "Установите libreoffice на сервер (на Railway — через nixpacks.toml)."
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"Pack 52: LibreOffice конвертация превысила {timeout_sec} сек"
            )

        if result.returncode != 0:
            stderr = result.stderr.decode("utf-8", errors="replace")
            raise RuntimeError(f"Pack 52: LibreOffice не смог сконвертировать DOCX → PDF: {stderr}")

        pdf_path = os.path.join(tmpdir, "statement.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError(
                f"Pack 52: PDF не появился в {tmpdir} после конвертации. "
                f"stdout: {result.stdout.decode(errors='replace')[:500]}"
            )

        with open(pdf_path, "rb") as f:
            return f.read()


def _replace_ep_badge_marker(doc, bank_data: dict) -> None:
    """
    Pack 47.15: ищет в документе ячейку с текстом "__EP_BADGE__", очищает её
    и вставляет inline-картинку плашки ЭП Сбербанка.

    Картинка генерируется через ep_badge_renderer.render_ep_badge_png с
    актуальной датой подписи (bank.statement_date_formatted).

    Если маркер не найден — функция ничего не делает (back-compat для Альфы
    и других банков).
    """
    from .ep_badge_renderer import render_ep_badge_png
    from docx.shared import Mm

    MARKER = "__EP_BADGE__"
    statement_date = bank_data.get("statement_date_formatted", "") if bank_data else ""

    # Перебираем все ячейки всех таблиц (включая вложенные)
    target_cell = None
    target_paragraph = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if MARKER in cell.text:
                    for p in cell.paragraphs:
                        if MARKER in p.text:
                            target_cell = cell
                            target_paragraph = p
                            break
                if target_cell:
                    break
            if target_cell:
                break
        if target_cell:
            break

    if target_cell is None or target_paragraph is None:
        return  # Маркера нет — нечего делать

    # Очищаем параграф (удаляем все runs)
    for r in list(target_paragraph.runs):
        r._r.getparent().remove(r._r)

    # Генерируем PNG плашки
    png_bytes = render_ep_badge_png(statement_date_str=statement_date)

    # Вставляем как inline-картинку шириной 80mm (соответствует ширине
    # right-колонки sig_tbl в шаблоне Сбера).
    import io as _io
    png_io = _io.BytesIO(png_bytes)
    run = target_paragraph.add_run()
    pic = run.add_picture(png_io, width=Mm(80))

    # Pack 47.17 FIX (extended in Pack 47.18): python-docx ставит
    # <pic:cNvPr id="0"> по умолчанию для каждой add_picture. Word ругается
    # на ЛЮБОЙ id="0" (OOXML schema требует id >= 1).
    # 
    # Pack 47.18: проходим по ВСЕМУ документу, ищем <pic:cNvPr id="0"> и
    # ставим уникальные id (1002, 1003, ...). Это исправляет не только нашу
    # вставленную картинку, но и sber_logo.png из шаблона.
    from docx.oxml.ns import qn as _qn
    next_id = 1002
    for el in doc.element.iter():
        if el.tag.endswith("}cNvPr") and el.get("id") == "0":
            el.set("id", str(next_id))
            next_id += 1


def _replace_markers_in_tr(tr_element, tx: dict):
    """
    Заменяет маркеры __TX_*__ на значения транзакции в строке таблицы.

    Pack 16.4: если значение содержит '\\n' (как в описании зарплаты —
    Плательщик / ИНН / Счёт / Назначение платежа), разбивает его на
    отдельные параграфы в ячейке. Word игнорирует '\\n' в <w:t> тегах —
    для реального переноса нужны отдельные <w:p>.
    """
    marker_to_value = {
        "__TX_DATE__": tx.get("date_formatted", ""),
        "__TX_CODE__": tx.get("code", ""),
        "__TX_DESCRIPTION__": tx.get("description", "") or "",
        "__TX_AMOUNT__": tx.get("amount_formatted", ""),
        # Pack 47.2: мульти-банк маркеры. Шаблоны без них (Альфа) — игнорят:
        # _replace_markers_in_tr ищет маркер в тексте параграфа, нет — нет замены.
        "__TX_CATEGORY__": tx.get("category", "") or "",
        "__TX_BALANCE__": tx.get("running_balance_formatted", "") or "",
        # Pack 48.0: маркеры для шаблона ТБанка.
        # __TX_DATE_SETTLE__ — "дата+время списания" (правый столбец дат у ТБанка);
        #   значение может содержать \n → multiline-логика разобьёт на 2 параграфа.
        #   Fallback на date_formatted если settle_date_formatted не задан
        #   (применимо к Сбер/Альфа-tx где этого поля нет — но в их шаблонах
        #   нет и маркера, так что fallback срабатывает только при ошибках данных).
        # __TX_AMOUNT_CARD__ — "сумма операции в валюте карты". Для рублёвых
        #   счетов = amount_formatted (та же сумма в той же валюте). Поле может
        #   появиться в Tx-модели позже для валютных операций.
        # __TX_CARD__ — 4 цифры "номера карты". Заполняет _apply_tbank_postprocess
        #   детерминированно по bank_account. Для не-ТБанк — пустая строка.
        "__TX_DATE_SETTLE__": tx.get("settle_date_formatted") or tx.get("date_formatted", ""),
        "__TX_AMOUNT_CARD__": tx.get("amount_card_formatted") or tx.get("amount_formatted", ""),
        "__TX_CARD__": tx.get("card_number", "") or "",
    }

    cells = tr_element.findall('.//w:tc', NS)

    for cell in cells:
        paragraphs = cell.findall('.//w:p', NS)

        for p in paragraphs:
            ts = p.findall('.//w:t', NS)
            full_text = "".join(t.text or "" for t in ts)

            for marker, value in marker_to_value.items():
                if marker in full_text:
                    if '\n' in value:
                        _replace_marker_with_multiline(cell, p, marker, value)
                    else:
                        _replace_marker_inline(p, marker, value)
                    break


def _replace_marker_inline(p_element, marker: str, value: str):
    """Простая замена маркера в текстах параграфа."""
    for t in p_element.findall('.//w:t', NS):
        if t.text and marker in t.text:
            t.text = t.text.replace(marker, value)


def _force_left_align_paragraph(p_element):
    """
    Pack 34.4 — гарантирует left-align параграфа.

    Маркер-строка в bank_statement_template.docx унаследовала <w:jc w:val="both"/>
    (justify) от исходника Алиева. Пока контент влезал в одну строку — justify
    не проявлялся. Но для очень длинных строк (которые всё-таки переносятся
    внутри ячейки) Word растягивает первую строку по ширине, оставляя хвост
    одиноким на следующей строке.

    Решение — насильно ставить <w:jc w:val="left"/> на ВСЕ параграфы клонов
    строк описания. Это работает и для русской выписки, и для испанского
    перевода (LLM-pipeline не трогает XML-форматирование).

    Идемпотентно: если jc уже left — ничего не делает; если other — перезаписывает.
    """
    ppr = p_element.find('w:pPr', NS)
    if ppr is None:
        ppr = etree.SubElement(p_element, f'{W_NS}pPr')
        # pPr должен быть первым элементом параграфа
        p_element.remove(ppr)
        p_element.insert(0, ppr)

    jc = ppr.find('w:jc', NS)
    if jc is None:
        jc = etree.SubElement(ppr, f'{W_NS}jc')
    jc.set(f'{W_NS}val', 'left')


def _replace_marker_with_multiline(cell_element, p_element, marker: str, multiline_value: str):
    """
    Заменяет маркер на многострочное значение, разбивая на отдельные параграфы.

    Стратегия:
    - Первая строка значения подставляется в существующий <w:p>
    - Для остальных строк создаются deepcopy этого <w:p>, текст заменяется
    - Новые параграфы вставляются после оригинального в ячейке

    Это сохраняет форматирование (отступы, стиль, размер шрифта).
    """
    lines = multiline_value.split('\n')
    if not lines:
        _replace_marker_inline(p_element, marker, "")
        return

    # Заменяем маркер в первом параграфе на первую строку
    _replace_marker_inline(p_element, marker, lines[0])
    # Pack 34.4: страхуем left-align (на случай если когда-то контент перенесётся)
    _force_left_align_paragraph(p_element)

    # Для остальных строк создаём копии параграфа
    parent_of_p = p_element.getparent()
    p_index = list(parent_of_p).index(p_element)
    insert_position = p_index + 1

    for line in lines[1:]:
        new_p = deepcopy(p_element)
        # В копии текст содержит lines[0] — заменим на текущую line
        ts_in_new = new_p.findall('.//w:t', NS)
        for t in ts_in_new:
            if t.text and lines[0] in t.text:
                t.text = t.text.replace(lines[0], line, 1)
                break

        # Pack 16.5d: для длинных строк (например «Назначение платежа: ...»),
        # которые при переносе должны выравниваться по началу первой строки,
        # а не по новому отступу firstLine, поправляем ind:
        # left += firstLine, firstLine = 0.
        # В оригинале Алиева у параграфа «Назначение платежа» именно такая
        # структура: left=388, без firstLine (199 + 195 ? 388 + округление).
        if line.startswith("Назначение платежа"):
            ppr = new_p.find('w:pPr', NS)
            if ppr is not None:
                ind = ppr.find('w:ind', NS)
                if ind is not None:
                    left_attr = f'{W_NS}left'
                    firstLine_attr = f'{W_NS}firstLine'
                    cur_left = int(ind.get(left_attr, '0') or '0')
                    cur_first = int(ind.get(firstLine_attr, '0') or '0')
                    if cur_first > 0:
                        ind.set(left_attr, str(cur_left + cur_first))
                        # Удаляем firstLine
                        if firstLine_attr in ind.attrib:
                            del ind.attrib[firstLine_attr]

                # Pack 25.3: у Алиева ПОСЛЕДНИЙ параграф (Назначение платежа)
                # имеет spacing="before=40 after=40" вместо обычного
                # "before=54 line=244 lineRule=auto". Это создаёт визуально
                # одинаковый воздух сверху и снизу серого блока ячейки.
                # Pack 25.4: after=80 — Word "съедает" часть space-after
                # последнего параграфа в табличной ячейке, поэтому компенсируем
                # удвоением (40 → 80) для визуального равенства верх/низ.
                spacing = ppr.find('w:spacing', NS)
                if spacing is None:
                    spacing = etree.SubElement(ppr, f'{W_NS}spacing')
                # Чистим все атрибуты spacing и ставим before=40 after=80
                for attr in list(spacing.attrib.keys()):
                    del spacing.attrib[attr]
                spacing.set(f'{W_NS}before', '40')
                spacing.set(f'{W_NS}after', '80')

        # Pack 34.4: страхуем left-align для каждого клона
        _force_left_align_paragraph(new_p)
        parent_of_p.insert(insert_position, new_p)
        insert_position += 1


def _remove_empty_paragraph_between_tables(doc):
    """
    Pack 16.4: убирает пустой параграф между таблицей операций и таблицей подписи,
    чтобы подпись могла поместиться сразу после операций.
    """
    body = doc.element.body
    children = list(body)

    for i in range(len(children) - 2):
        if etree.QName(children[i]).localname != 'tbl':
            continue
        if etree.QName(children[i + 1]).localname != 'p':
            continue
        if etree.QName(children[i + 2]).localname != 'tbl':
            continue

        ts = children[i + 1].findall('.//w:t', NS)
        full_text = "".join(t.text or "" for t in ts).strip()

        if not full_text:
            body.remove(children[i + 1])
            break


def _apply_gray_shading_to_row(tr_element):
    """
    Pack 16.5: добавляет серый фон (E8E8E8) каждой ячейке строки таблицы.

    В оригинальной выписке Алиева строки с доходом (зарплата от компании) имеют
    серую заливку. Мы применяем тот же стиль к строкам с положительной суммой.

    Если в ячейке уже есть <w:shd>, она заменяется. Иначе создаётся новый.
    """
    cells = tr_element.findall('.//w:tc', NS)
    for cell in cells:
        tcPr = cell.find('w:tcPr', NS)
        if tcPr is None:
            tcPr = etree.SubElement(cell, f'{W_NS}tcPr')
            # tcPr должен идти первым в tc — переместим его
            cell.remove(tcPr)
            cell.insert(0, tcPr)

        # Удаляем старый shd если есть
        old_shd = tcPr.find('w:shd', NS)
        if old_shd is not None:
            tcPr.remove(old_shd)

        # Создаём новый
        shd = etree.SubElement(tcPr, f'{W_NS}shd')
        shd.set(f'{W_NS}val', 'clear')
        shd.set(f'{W_NS}color', 'auto')
        shd.set(f'{W_NS}fill', 'E8E8E8')


def _apply_bold_to_amount_cell(tr_element):
    """
    Pack 25.0: делает сумму в 4-й (последней) ячейке строки жирной.

    В эталонной выписке Алиева суммы поступлений (300 000,00 RUR) выделены
    жирным шрифтом — это визуально подчёркивает поступления для проверяющего.
    Применяется ТОЛЬКО к строкам поступлений (вызов после проверки amount_val > 0).

    Сумма находится в последней ячейке строки таблицы (где маркер __TX_AMOUNT__).
    Делаем жирными все runs в этой ячейке.
    """
    cells = tr_element.findall('.//w:tc', NS)
    if not cells:
        return
    amount_cell = cells[-1]  # последняя ячейка — сумма

    for run in amount_cell.findall('.//w:r', NS):
        rPr = run.find('w:rPr', NS)
        if rPr is None:
            rPr = etree.Element(f'{W_NS}rPr')
            run.insert(0, rPr)
        # <w:b/> делает run жирным; должен идти в начале rPr
        if rPr.find('w:b', NS) is None:
            b = etree.Element(f'{W_NS}b')
            rPr.insert(0, b)


def _add_bottom_border_to_row(tr_element):
    """
    Pack 25.1: добавляет <w:bottom> в <w:tcBorders> каждой ячейки строки.

    Pack 25.0 убрал <w:bottom> из шаблонной маркер-строки — между строками
    операций больше нет двойной линии (как в Алиеве). Но это удалило линию
    и под последней строкой таблицы — таблица не «закрывается» снизу.

    Эту функцию вызываем ТОЛЬКО для последней строки операций — у Алиева
    таблица операций имеет нижнюю границу под последней операцией.
    """
    cells = tr_element.findall('.//w:tc', NS)
    for cell in cells:
        tcPr = cell.find('w:tcPr', NS)
        if tcPr is None:
            continue
        tc_borders = tcPr.find('w:tcBorders', NS)
        if tc_borders is None:
            tc_borders = etree.SubElement(tcPr, f'{W_NS}tcBorders')

        # Если <w:bottom> уже есть — пропускаем
        if tc_borders.find('w:bottom', NS) is not None:
            continue

        bottom = etree.SubElement(tc_borders, f'{W_NS}bottom')
        bottom.set(f'{W_NS}val', 'single')
        bottom.set(f'{W_NS}sz', '4')
        bottom.set(f'{W_NS}space', '0')
        bottom.set(f'{W_NS}color', '7E7E7E')


def _add_vertical_padding_to_cells(tr_element, top_dxa: int = 80, bottom_dxa: int = 120):
    """
    Pack 25.1: добавляет вертикальный padding ячейкам строки через <w:tcMar>.
    Pack 25.2: top=80, bottom=120 — асимметрия для визуального равенства.

    В Word'e <w:spacing w:after="40"/> ПОСЛЕДНЕГО параграфа в табличной ячейке
    игнорируется (известная особенность Word). Поэтому верхний воздух получает
    "бонус" от space-before первого параграфа (40 dxa), а нижний — нет.
    Компенсируем bottom += 40 = 120, чтобы визуальный воздух был одинаков.

    Эффективный верхний воздух: tcMar top (80) + spacing before (40) = 120 dxa
    Эффективный нижний воздух: tcMar bottom (120) + 0 (after игнорится) = 120 dxa

    Применяется ТОЛЬКО к строкам поступлений (после _apply_gray_shading_to_row).
    """
    cells = tr_element.findall('.//w:tc', NS)
    for cell in cells:
        tcPr = cell.find('w:tcPr', NS)
        if tcPr is None:
            tcPr = etree.SubElement(cell, f'{W_NS}tcPr')
            cell.remove(tcPr)
            cell.insert(0, tcPr)

        # Удаляем старый tcMar если есть
        old_mar = tcPr.find('w:tcMar', NS)
        if old_mar is not None:
            tcPr.remove(old_mar)

        # Создаём новый tcMar с верхним и нижним padding
        tc_mar = etree.SubElement(tcPr, f'{W_NS}tcMar')
        for side, value in [('top', top_dxa), ('bottom', bottom_dxa)]:
            elem = etree.SubElement(tc_mar, f'{W_NS}{side}')
            elem.set(f'{W_NS}w', str(value))
            elem.set(f'{W_NS}type', 'dxa')


def _set_cant_split(tr_element):
    """
    Pack 16.5b: добавляет <w:cantSplit/> в <w:trPr> строки таблицы.

    Это запрещает Word разрывать строку между страницами — если строка
    не помещается на текущей странице, она ЦЕЛИКОМ переносится на
    следующую (стандарт банковских выписок).
    """
    trPr = tr_element.find('w:trPr', NS)
    if trPr is None:
        # Создаём trPr и кладём его первым (после положения tblPrEx)
        trPr = etree.Element(f'{W_NS}trPr')
        # Найдём куда вставить — trPr должен быть до <w:tc>
        tc_idx = None
        for i, child in enumerate(tr_element):
            if etree.QName(child).localname == 'tc':
                tc_idx = i
                break
        if tc_idx is not None:
            tr_element.insert(tc_idx, trPr)
        else:
            tr_element.append(trPr)

    # Проверим — может уже есть cantSplit
    existing = trPr.find('w:cantSplit', NS)
    if existing is None:
        cant_split = etree.SubElement(trPr, f'{W_NS}cantSplit')


def _set_keep_next_on_row(tr_element):
    """
    Pack 16.5c: добавляет <w:cantSplit/> и устанавливает на параграфы внутри ячеек
    атрибут keepNext через pPr — чтобы строка «прилипла» к следующему контенту.

    На уровне строки таблицы Word не понимает <w:keepNext/>. Чтобы строка
    держалась с подписью, ставим keepNext на ВСЕ параграфы в ячейках строки —
    это эквивалентный приём.
    """
    # На каждой ячейке строки — на каждом параграфе — добавляем <w:keepNext/>
    cells = tr_element.findall('.//w:tc', NS)
    for cell in cells:
        for p in cell.findall('.//w:p', NS):
            ppr = p.find('w:pPr', NS)
            if ppr is None:
                ppr = etree.Element(f'{W_NS}pPr')
                p.insert(0, ppr)

            if ppr.find('w:keepNext', NS) is None:
                # keepNext должен идти в начале pPr (после pStyle)
                keep_next = etree.SubElement(ppr, f'{W_NS}keepNext')


def _set_keep_next_on_paragraph_between_tables(doc):
    """
    Pack 16.5c: ставит keepNext на все параграфы между Table 0 (операции)
    и Table 1 (подпись), чтобы они не отрывались от подписи.
    """
    body = doc.element.body
    children = list(body)

    table_indexes = [i for i, c in enumerate(children) if etree.QName(c).localname == 'tbl']
    if len(table_indexes) < 2:
        return

    for i in range(table_indexes[0] + 1, table_indexes[1]):
        p = children[i]
        if etree.QName(p).localname != 'p':
            continue
        ppr = p.find('w:pPr', NS)
        if ppr is None:
            ppr = etree.Element(f'{W_NS}pPr')
            p.insert(0, ppr)
        if ppr.find('w:keepNext', NS) is None:
            etree.SubElement(ppr, f'{W_NS}keepNext')



# ============================================================================
# Pack 53 — перевод выписки на испанский (отдельный flow от orchestrator)
# ============================================================================

def render_bank_statement_for_translation(
    application: Application,
    session: Session,
) -> bytes:
    """
    Pack 53: рендерит русскую DOCX-выписку в варианте для перевода.

    Отличие от render_bank_statement: НЕ вставляет PNG-печати (только чистит
    маркеры __STAMP_*__ из R0). Сигнатур-таблица остаётся, лейблы R1 на месте
    — переводятся в испанский ("(firma del empleado AO «ALFA-BANK»)" / ...).

    Используется в POST /admin/applications/{id}/bank-statement/translate
    как источник перед translate_docx().
    """
    return render_bank_statement(application, session, for_translation=True)


def render_bank_statement_combined_to_pdf(
    application: Application,
    session: Session,
    es_docx_bytes: bytes,
    timeout_sec: int = 60,
) -> bytes:
    """
    Pack 53: объединяет RU PDF (с печатями) и ES PDF (без печатей) в один PDF.

    Поток:
      1. render_bank_statement_to_pdf(application, session) → RU PDF
      2. es_docx_bytes → tempfile → soffice --headless --convert-to pdf → ES PDF
      3. pypdf merge: RU pages + ES pages → объединённый PDF
      4. Возврат bytes

    es_docx_bytes — байты переведённой испанской DOCX из R2 (storage.read).

    Используется в /download-file/bank_statement когда у заявки есть
    Application.bank_statement_translation_storage_key.
    """
    import subprocess
    import tempfile as _tempfile
    import os as _os

    # Этап 1: русская PDF (с печатями) — текущий путь
    ru_pdf = render_bank_statement_to_pdf(application, session, timeout_sec=timeout_sec)

    # Этап 2: испанская DOCX → PDF через soffice
    with _tempfile.TemporaryDirectory(prefix="vk_es_pdf_") as tmpdir:
        es_docx_path = _os.path.join(tmpdir, "statement_es.docx")
        with open(es_docx_path, "wb") as f:
            f.write(es_docx_bytes)

        try:
            result = subprocess.run(
                [
                    "soffice", "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    es_docx_path,
                ],
                capture_output=True,
                timeout=timeout_sec,
            )
        except FileNotFoundError:
            raise RuntimeError(
                "Pack 53: LibreOffice (soffice) не найден в PATH."
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"Pack 53: LibreOffice ES конвертация превысила {timeout_sec} сек"
            )

        if result.returncode != 0:
            stderr = result.stderr.decode("utf-8", errors="replace")
            raise RuntimeError(
                f"Pack 53: LibreOffice не смог сконвертировать ES DOCX → PDF: {stderr}"
            )

        es_pdf_path = _os.path.join(tmpdir, "statement_es.pdf")
        if not _os.path.exists(es_pdf_path):
            raise RuntimeError(
                f"Pack 53: ES PDF не появился в {tmpdir} после конвертации. "
                f"stdout: {result.stdout.decode(errors='replace')[:500]}"
            )

        with open(es_pdf_path, "rb") as f:
            es_pdf = f.read()

    # Этап 3: merge через pypdf
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    for pdf_bytes in (ru_pdf, es_pdf):
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    out_buf = io.BytesIO()
    writer.write(out_buf)
    return out_buf.getvalue()



# ============================================================================
# Pack 54 — Sber v2: подпись + круглая печать
# ============================================================================

def _insert_v2_sber_signatures(doc, *, mode: str = "full") -> None:
    """
    Pack 54: вставка PNG-печатей в Sber v2 шаблон.

    Структура Table[3] в v2-шаблоне:
      Row 0: «Дата формирования» | {{ statement_date }} | «Подпись» | __STAMP_SIGNATURE__
             (R0C3 имеет bottom-border = линия подписи)
      Row 1-4: блок «Сотрудник, ФИО, должность» (слева) + «Структурное, Территориальный,
               Номер, Адрес» (справа). Все статика — в шаблоне прописана.

    Стратегия (по аналогии с Альфа Pack 52-fix17):
      __STAMP_SIGNATURE__ → INLINE в R0C3 (подпись Кирьянова, ~35мм, садится на линию).
      Круглая печать Сбера → FLOATING, якорь = параграф подписи (тот же target_p),
        чтобы пересекала линию справа сверху.

    Параметры mode:
      "full"          — вставляем PNG (подпись inline + печать floating)
      "markers_only"  — только чистим маркеры __STAMP_SIGNATURE__ (для перевода без печатей)
    """
    from docx.shared import Mm

    assets_dir = TEMPLATES_DIR / "assets" / "v2_sber"

    # Этап 1. Найти параграф с маркером __STAMP_SIGNATURE__ (R0C3 = signature cell)
    target_p = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "__STAMP_SIGNATURE__" in p.text:
                        target_p = p
                        break
                if target_p:
                    break
            if target_p:
                break
        if target_p:
            break

    if target_p is None:
        # Маркера нет — не наш шаблон
        return

    # Pack 53: режим markers_only — только чистим маркер, PNG не вставляем
    if mode == "markers_only":
        for r in list(target_p.runs):
            r._element.getparent().remove(r._element)
        return

    signature_png = assets_dir / "signature.png"
    bank_png      = assets_dir / "stamp_bank.png"

    # Этап 2. Чистим target_p и вставляем подпись INLINE
    # Pack 54.0-fix3: ширина 35мм → 25мм. По эталону Сбера подпись маленькая
    # (~25мм). При 35мм Row 0 растягивался до 20мм высотой → большой провал
    # между «Подпись» и «Сотрудник, ...». С 25мм + vAlign=BOTTOM в R0C3
    # (см. fix3 в шаблоне) подпись сидит низом на линии подписи.
    for r in list(target_p.runs):
        r._element.getparent().remove(r._element)
    if signature_png.exists():
        try:
            run = target_p.add_run()
            run.add_picture(str(signature_png), width=Mm(25))
        except Exception as e:
            import logging
            logging.warning("Pack 54: не удалось inline %s: %s", signature_png.name, e)

    # Этап 3. Круглая печать Сбера floating, якорь = target_p
    # Pack 52 урок: маленькие y_off (±5..15мм), не доверять локальному превью —
    # позицию подгонять на проде. Начальные значения консервативные.
    if bank_png.exists():
        try:
            # Pack 54.0-fix2: x_off c -15 на +140, width c 50 на 35, y_off c -15 на -5.
            # Корень бага: relativeFrom="column" в _add_floating_picture — X отсчитывается
            # от ЛЕВОГО ПОЛЯ страницы, а не от ячейки якорного параграфа. С x_off=-15
            # печать оказывалась на 0..50мм от поля = поверх C0-C1 («Дата / ФИО»).
            # Геометрия Sber v2 (cols 2400/2700/2900/2204 dxa): C3 = 141..180мм от поля.
            # Подпись Кирьянова inline 35мм в R0C3. Центр печати = 140 + 35/2 ≈ 158мм
            # = середина R0C3 = пересекает подпись. y_off=-5 — рабочее значение Альфы
            # Pack 52, мелкое (Правило 72 / Инцидент 50).
            _add_floating_picture(
                target_p, bank_png, 35,
                x_offset_mm=140,
                y_offset_mm=-5,
            )
        except Exception as e:
            import logging
            logging.warning("Pack 54.0-fix2: не удалось floating bank: %s", e)

