"""
Pack 15 — DOCX translator (v4 — Pack 33.9 hardened prompts + static lookup).

Pack 33.9 fix (10.05.2026):
- Добавлен STATIC_SINGLE_WORD_GLOSSARY — словарь точных переводов для
  одиночных русских слов которые встречаются как изолированные фрагменты
  (заголовки таблиц, метки textbox-ячеек). Если фрагмент после strip()
  совпадает с ключом — переводим БЕЗ обращения к LLM.

  Зачем: на коротких фрагментах вроде "Описание" LLM игнорировала
  glossary в SYSTEM_PROMPT и отвечала разговорно:
    "Por favor, proporcione el texto en ruso que desea traducir. 
     Solo ha enviado la palabra «Описание»..."
  Static lookup исключает любую вероятность подобной аномалии для
  критичных однословных заголовков выписки и договоров.

- SYSTEM_PROMPT усилен абсолютным запретом отвечать вопросом / просить
  контекста / refuse'иться — добавлено правило 0 в самом начале блока
  ABSOLUTE RULES.

- single_system в _translate_one_by_one (fallback) усилен: добавлены те же
  правила + критический минимум glossary (банк/договорные заголовки).
  Раньше там был только generic «Translate Russian to Spanish» — на
  одиночных словах модель срывалась.

Pack 15.8 fix (заменяет Pack 15.7):
- _set_paragraph_text теперь МЕРДЖИТ w:firstLine в w:left у переведённого
  параграфа: new_left = old_left + firstLine, firstLine удаляется.

Pack 15.6 fix:
- _collect_all_paragraphs обходит <w:txbxContent> внутри <w:drawing>
  (textbox-ы внутри плавающих фигур). До этого пропускалась карточка
  реквизитов клиента в bank_statement_template (textbox-ы) и шапка
  таблицы операций (textbox-ы внутри header1.xml).
- Корректно обрабатываем <mc:AlternateContent>: переводим только
  <mc:Choice>, пропускаем <mc:Fallback>.

Pack 15.2 changes vs Pack 15.1:
- Промпт явно требует переводить ИНН → NIF, КПП → KPP, ОГРН → OGRN,
  БИК → BIC (раньше говорил «keep as-is»).
- Few-shot пример с реквизитной таблицей.
"""

import io
import json
import logging
import re
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from app.services.llm import get_llm_client

from .name_substitution import SubstitutionDict

log = logging.getLogger(__name__)


BATCH_SIZE = 30

_SKIP_PATTERNS = [
    re.compile(r"^\s*$"),
    re.compile(r"^[\d\s.,\-/:]+$"),
    re.compile(r"^[A-Z]{2,5}\s*\d{6,}$"),
]


# Pack 15.6: namespaces для обхода textbox-ов
_W_TXBX = qn('w:txbxContent')
_W_P = qn('w:p')
_MC_FALLBACK = '{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback'


# ============================================================================
# Pack 33.9: static lookup для single-word фрагментов
# ============================================================================
# Если после strip() текст параграфа совпадает с одним из ключей —
# подставляем готовый перевод и НЕ дёргаем LLM. Защита от случаев когда
# модель на короткой строке без контекста отвечает разговорно вместо
# перевода (видели на «Описание» в шапке таблицы выписки).
#
# Покрывает:
# - заголовки таблиц банковской выписки (textbox-ы в header1.xml)
# - короткие метки в карточках реквизитов
# - служебные слова из договоров/актов/счетов
#
# Все ключи и значения — точные jurada-style варианты из основного
# SYSTEM_PROMPT glossary. Изменения только синхронно с глоссарием.

STATIC_SINGLE_WORD_GLOSSARY: dict[str, str] = {
    # Банковская выписка — заголовки и метки
    "Описание": "Descripción",
    "Сумма": "Importe",
    "Дата": "Fecha",
    "Код": "Código",
    "Назначение": "Concepto",
    "Плательщик": "Ordenante",
    "Получатель": "Beneficiario",
    "Клиент": "Cliente",
    "Паспорт": "Pasaporte",
    "Адрес": "Dirección",
    "Поступления": "Ingresos",
    "Расходы": "Gastos",
    "Страница": "Página",
    "из": "de",
    "Дебет": "Débito",
    "Кредит": "Crédito",
    "Баланс": "Saldo",
    "Остаток": "Saldo",
    "Валюта": "Moneda",

    # Договорные/актовые служебные слова
    "Исполнитель": "Contratista",
    "Заказчик": "Cliente",
    "Сторона": "Parte",
    "Стороны": "Partes",
    "Город": "Ciudad",
    "Подпись": "Firma",
    "Печать": "Sello",
    "Должность": "Cargo",
    "ФИО": "Nombre completo",
    "Реквизиты": "Datos bancarios",
    "Услуги": "Servicios",
    "Период": "Período",
    "Итого": "Total",
    "Всего": "Total",

    # Аббревиатуры-метки (часто в одиночных ячейках)
    "ИНН": "NIF",
    "КПП": "KPP",
    "ОГРН": "OGRN",
    "БИК": "BIC",
    "СНИЛС": "SNILS",
    "ОКПО": "OKPO",
}


def _try_static_translation(text: str) -> Optional[str]:
    """Pack 33.9: возвращает готовый перевод если text — single-word из glossary.

    Стрипает whitespace и проверяет точное совпадение. None если не нашли
    (тогда вызывающая сторона отправит фрагмент в LLM как обычно).
    """
    stripped = text.strip()
    if not stripped:
        return None
    return STATIC_SINGLE_WORD_GLOSSARY.get(stripped)


# Pack 15.2: jurada-style промпт со словарём из реальных подач.
# Pack 33.9: добавлено правило 0 — запрет refuse / ask for clarification.
SYSTEM_PROMPT = """You are translating Russian business/legal documents to Spanish as a DRAFT for a sworn translator (traductor jurado MAE) who will review and finalize. Your goal: match the established Spanish jurada conventions so the jurado has minimal corrections to make.

Input: JSON array of text fragments from a Russian document.
Output: JSON array of Spanish translations, EXACTLY same length and order. No commentary. No markdown fences.

═══ ABSOLUTE RULES ═══

0. NEVER refuse to translate. NEVER ask for clarification. NEVER add commentary, apologies, or explanations to the output. If a fragment looks short, ambiguous, or like a single word — TRANSLATE IT ANYWAY using the glossary below or your best literal Spanish equivalent. The output array MUST contain pure translations, NOT meta-text like "Por favor, proporcione el texto" or "Could you clarify". Each item in the output JSON array is the Spanish translation of the input item at the same index — no exceptions.

1. NEVER MODIFY any Latin-script text already in the input — names of people, companies, addresses, banks. They have been pre-substituted from official documents (passports, EGRYL extracts) and must be preserved character-for-character. If you see "Yuksel Vedat", "INZHGEOSERVIS", "SBERBANK" — output them unchanged.

2. NEVER MODIFY: numbers (1234.56, 100 000), dates in any format (04.05.2025, 2025-05-04, «04» мая 2025 г. — translate ONLY the Russian month name, keep digits/structure), tax IDs and account numbers (the digit strings — but DO translate the LABEL before them, see rule 4), email addresses, phone numbers, document codes (#2026-0003).

3. If a fragment is already in Spanish or English — return it unchanged. If a fragment is a Jinja template marker like {{ var }} — return unchanged.

4. ALWAYS TRANSLATE THESE LABELS (jurada convention — Russian abbreviations are NEVER kept in Cyrillic in the final document):
   - ИНН → NIF (always — it's the equivalent for "tax ID")
   - КПП → KPP (transliteration — Spanish has no equivalent, keep this Latin form)
   - ОГРН → OGRN (transliteration)
   - БИК → BIC (Spanish has its own BIC, same letters)
   - СНИЛС → SNILS (transliteration)
   - ОКПО → OKPO
   - р/с / Р/с / расчётный счёт → c/c (cuenta corriente) or "Cuenta corriente"
   - к/с / К/с / корреспондентский счёт → c/corr or "Cuenta corresponsal"
   - БИК банка → BIC del banco

═══ JURADA GLOSSARY (use these exact translations) ═══

Document types:
- Договор оказания услуг → Contrato de prestación de servicios remunerados
- Акт об оказании услуг / Акт оказанных услуг → Acta de servicios prestados
- Счёт на оплату → Factura
- Резюме → Curriculum Vitae
- Письмо от компании / Письмо-поручение → Carta de la empresa
- Выписка по счёту → Extracto de cuenta
- Справка → Certificado

Bank statement headers (Pack 15.6 — appear as standalone fragments from textbox columns):
- Номер счета → Número de cuenta
- Дата открытия счета → Fecha de apertura de la cuenta
- Валюта счета → Moneda de la cuenta
- Тип счета → Tipo de cuenta
- Текущий счёт → Cuenta corriente
- Дата формирования выписки → Fecha de emisión del extracto
- Клиент → Cliente
- Адрес регистрации → Dirección de registro
- Паспорт → Pasaporte
- Дата проводки → Fecha de contabilización
- Код операции → Código de operación
- Описание → Descripción
- Сумма в валюте счета / Сумма в валюте счёта → Importe en la divisa de la cuenta
- Операции по счету → Operaciones en cuenta
- Выписка по счету → Extracto de cuenta
- Страница → Página, из (in "Страница X из Y") → de
- Входящий остаток → Saldo entrante
- Исходящий остаток → Saldo saliente
- Поступления → Ingresos
- Расходы → Gastos
- Уполномоченное лицо → persona autorizada
- (подпись сотрудника АО «АЛЬФА-БАНК») → (firma del empleado de «ALFA-BANK», S.A.)
- Ф.И.О. сотрудника АО «АЛЬФА-БАНК» → Nombre completo del empleado de «ALFA-BANK», S.A.

Common geographic/address terms:
- ул. / улица → c/
- г. / город → ciudad de
- д. / дом → №
- кв. / квартира → piso
- область → región
- проспект → avenida
- Москва → Moscú
- Санкт-Петербург → San Petersburgo
- Other Russian city/street names — TRANSLITERATE (Краснодонская → Krasnodonskaya)
- Юрид. адрес / Юридический адрес → Domicilio social
- Почт. адрес / Почтовый адрес → Dirección postal

Months (translate within dates):
- января → enero, февраля → febrero, марта → marzo, апреля → abril
- мая → mayo, июня → junio, июля → julio, августа → agosto
- сентября → septiembre, октября → octubre, ноября → noviembre, декабря → diciembre

Markers:
- (подпись) → (firma:)
- (печать) / М.П. → (Sello:)

═══ STYLE ═══

- Use formal Spanish (Spain Spanish).
- Use "0,1%" style (comma decimal separator).
- Use «» for inner quotes when copying from Russian.

═══ FEW-SHOT EXAMPLES ═══

Example 1 (body):
Input: ["Граждане Российской Федерации Иванов Иван Иванович, в дальнейшем именуемый «Исполнитель», и ООО «Ромашка», в лице Генерального директора Петрова П.П., именуемое в дальнейшем «Заказчик»"]
Output: ["el ciudadano de la Federación de Rusia D. Ivanov Ivan Ivanovich, en lo sucesivo denominado el «Contratista», y la Sociedad Limitada «Romashka», representada por el Director General D. Petrov P.P., en lo sucesivo denominada el «Cliente»"]

Example 2 (requisites — multi-line cell):
Input: ["Заказчик\\nИНЖГЕОСЕРВИС\\nИНН 2320219620, КПП 232001001\\nЮрид. адрес: ул. Ленина, д. 5, г. Сочи\\nР\\\\с: 40702810000000000001\\nСбербанк\\nБИК банка: 044525225\\nК\\\\с: 30101810400000000225"]
Output: ["El Cliente\\n«INZHGEOSERVIS», S.L.\\nNIF 2320219620, KPP 232001001\\nDomicilio social: c/ Lenina, № 5, ciudad de Sochi\\nc/c: 40702810000000000001\\n«SBERBANK», S.A.\\nBIC del banco: 044525225\\nc/corr: 30101810400000000225"]

Example 3 (requisites with empty fields):
Input: ["Исполнитель:\\nYuksel Vedat\\nПаспорт U23616456,\\nвыдан 08.10.2020 ELAZIĞ\\nNIF\\n—\\n\\nc/c\\nв\\nBIC del banco:\\nc/corr:"]
Output: ["el Contratista:\\nYuksel Vedat\\nPasaporte U23616456,\\nexpedido el 08.10.2020 ELAZIĞ\\nNIF\\n—\\n\\nc/c\\nen\\nBIC del banco:\\nc/corr:"]

Example 4 (Pack 33.9 — single-word table headers):
Input: ["Описание", "Сумма", "Дата проводки", "Код операции"]
Output: ["Descripción", "Importe", "Fecha de contabilización", "Código de operación"]

Note: Latin-script names (Yuksel Vedat) stay unchanged. Russian abbreviation labels (ИНН, КПП, БИК, ОГРН) ALWAYS become Latin labels (NIF, KPP, BIC, OGRN). Empty values stay as "—". Single-word fragments are STILL translated using the glossary above — never refused, never asked about.
"""


def _should_skip(text: str) -> bool:
    """Быстрая проверка: нужно ли отправлять текст в LLM."""
    if not text or not text.strip():
        return True
    for pattern in _SKIP_PATTERNS:
        if pattern.match(text):
            return True
    if not re.search(r"[А-Яа-яЁё]", text):
        return True
    return False


def _extract_paragraph_text(p: Paragraph) -> str:
    return "".join(run.text for run in p.runs)


def _set_paragraph_text(p: Paragraph, new_text: str) -> None:
    if not p.runs:
        return
    p.runs[0].text = new_text
    for run in p.runs[1:]:
        run.text = ""
    # Pack 15.8: мерджим w:firstLine в w:left.
    # На русском короткие описания умещаются в 1 строку и firstLine незаметен.
    # На испанском описания длиннее, переносятся, и firstLine ломает выравнивание
    # (первая строка отступает на firstLine, остальные — нет; получается лесенка).
    # Просто удалить firstLine нельзя: в шаблоне выписки соседние параграфы
    # одного блока (Ordenante/NIF/Cuenta vs Concepto) рассчитаны на одну
    # визуальную позицию, но достигают её разными способами (left=199+firstLine=195
    # vs left=394). Поэтому мерджим: new_left = old_left + firstLine, firstLine удаляем.
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            first_line = ind.get(qn('w:firstLine'))
            if first_line is not None:
                try:
                    fl = int(first_line)
                    left = int(ind.get(qn('w:left')) or '0')
                    ind.set(qn('w:left'), str(left + fl))
                    del ind.attrib[qn('w:firstLine')]
                except ValueError:
                    # Не number — просто удалим firstLine на всякий случай
                    del ind.attrib[qn('w:firstLine')]


def _is_inside_mc_fallback(elem) -> bool:
    """Pack 15.6: True если элемент внутри <mc:Fallback>.
    Word дублирует drawings в <mc:Choice> (modern WPS) и <mc:Fallback> (legacy VML).
    Word рендерит только Choice; при сохранении сам синхронизирует Fallback с Choice.
    Поэтому мы переводим только Choice — иначе LLM может выдать чуть разные переводы
    и Fallback будет расходиться с Choice."""
    parent = elem.getparent()
    while parent is not None:
        if parent.tag == _MC_FALLBACK:
            return True
        parent = parent.getparent()
    return False


def _iter_txbx_paragraphs(part_element, doc: Document):
    """Pack 15.6: возвращает все <w:p> внутри <w:txbxContent> данного part-элемента,
    пропуская дубликаты в <mc:Fallback>. part_element — это doc.element.body для
    основного документа или header._element / footer._element для колонтитулов."""
    for txbx in part_element.iter(_W_TXBX):
        if _is_inside_mc_fallback(txbx):
            continue
        for p_elem in txbx.iter(_W_P):
            yield Paragraph(p_elem, doc.element.body)


def _collect_all_paragraphs(doc: Document) -> list[Paragraph]:
    """Pack 15.6: добавлен обход <w:txbxContent> в body, headers и footers.

    Без этого обхода пропускались textbox-ы в bank_statement_template:
    - левая карточка реквизитов клиента (Номер счета, Дата открытия счета, ...)
    - шапка таблицы операций в header1.xml (Дата проводки, Код операции, ...)
    - нумерация страниц в footer1.xml (Страница X из Y)
    """
    paragraphs: list[Paragraph] = []
    paragraphs.extend(doc.paragraphs)

    def _walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
                    _walk_tables(cell.tables)

    _walk_tables(doc.tables)

    # Pack 15.6: textbox-ы в body
    paragraphs.extend(_iter_txbx_paragraphs(doc.element.body, doc))

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is None:
                continue
            paragraphs.extend(header.paragraphs)
            _walk_tables(header.tables)
            # Pack 15.6: textbox-ы в header
            paragraphs.extend(_iter_txbx_paragraphs(header._element, doc))
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is None:
                continue
            paragraphs.extend(footer.paragraphs)
            _walk_tables(footer.tables)
            # Pack 15.6: textbox-ы в footer
            paragraphs.extend(_iter_txbx_paragraphs(footer._element, doc))

    return paragraphs


async def _translate_batch(texts: list[str]) -> list[str]:
    if not texts:
        return []

    client = get_llm_client()
    user_payload = json.dumps(texts, ensure_ascii=False)

    try:
        response = await client.complete(
            system=SYSTEM_PROMPT,
            user=user_payload,
            max_tokens=8192,
            temperature=0.0,
        )
    except Exception as e:
        log.warning(f"[translation] LLM call failed for batch of {len(texts)}: {e}")
        return list(texts)

    cleaned = response.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```\s*$", "", cleaned)

    try:
        translated = json.loads(cleaned)
    except json.JSONDecodeError as e:
        log.warning(f"[translation] LLM returned non-JSON: {e}. Response head: {cleaned[:300]}")
        return await _translate_one_by_one(texts)

    if not isinstance(translated, list) or len(translated) != len(texts):
        log.warning(
            f"[translation] LLM returned wrong shape: "
            f"expected list of {len(texts)}, got {type(translated).__name__} of "
            f"len={len(translated) if hasattr(translated, '__len__') else '?'}"
        )
        return await _translate_one_by_one(texts)

    result = []
    for i, item in enumerate(translated):
        if isinstance(item, str):
            result.append(item)
        else:
            log.warning(f"[translation] item {i} is not a string: {type(item).__name__}, falling back")
            result.append(texts[i])
    return result


async def _translate_one_by_one(texts: list[str]) -> list[str]:
    """Pack 33.9: усилен single_system промпт — добавлены запрет на refuse и
    минимально необходимый glossary критичных однословных меток. Раньше тут
    был только generic «Translate Russian to Spanish» — модель на одиночных
    словах вроде «Описание» отвечала просьбой контекста («Por favor,
    proporcione...») вместо перевода.
    """
    log.info(f"[translation] Falling back to one-by-one translation for {len(texts)} fragments")
    client = get_llm_client()
    results = []
    single_system = (
        "You are translating Russian business/legal text to formal Spanish (Spain) "
        "as a DRAFT for a sworn translator. "
        "\n\n"
        "ABSOLUTE RULES:\n"
        "0. NEVER refuse to translate. NEVER ask for clarification. NEVER add "
        "commentary, apologies, or explanations. If the input is a single word, "
        "translate that single word using the glossary or your best literal Spanish "
        "equivalent. Output ONLY the translation — no preamble, no quotes, no "
        "markdown.\n"
        "1. Keep all Latin-script text (names, companies) unchanged.\n"
        "2. Keep all numbers, dates, account numbers, emails, phones unchanged.\n"
        "3. Translate Russian abbreviation labels: "
        "ИНН → NIF, КПП → KPP, ОГРН → OGRN, БИК → BIC, СНИЛС → SNILS, ОКПО → OKPO.\n"
        "\n"
        "GLOSSARY (use these exact translations for single-word fragments):\n"
        "Описание → Descripción\n"
        "Сумма → Importe\n"
        "Дата → Fecha\n"
        "Код → Código\n"
        "Назначение → Concepto\n"
        "Плательщик → Ordenante\n"
        "Получатель → Beneficiario\n"
        "Клиент → Cliente\n"
        "Паспорт → Pasaporte\n"
        "Адрес → Dirección\n"
        "Поступления → Ingresos\n"
        "Расходы → Gastos\n"
        "Исполнитель → Contratista\n"
        "Заказчик → Cliente\n"
        "Дата проводки → Fecha de contabilización\n"
        "Код операции → Código de operación\n"
        "Сумма в валюте счета → Importe en la divisa de la cuenta\n"
        "Операции по счету → Operaciones en cuenta\n"
        "Входящий остаток → Saldo entrante\n"
        "Исходящий остаток → Saldo saliente\n"
        "\n"
        "Output ONLY the Spanish translation. No commentary."
    )
    for text in texts:
        try:
            response = await client.complete(
                system=single_system,
                user=text,
                max_tokens=1024,
                temperature=0.0,
            )
            results.append(response.strip())
        except Exception as e:
            log.warning(f"[translation] one-by-one failed for text {text[:50]!r}: {e}")
            results.append(text)
    return results



# ============================================================================
# Pack 35.9: разбиение шапки «город + дата» на 2 параграфа
# ============================================================================

_SPANISH_DATE_RE = re.compile(
    r'«?\s*\d{1,2}\s*»?\s+de\s+\w+\s+de\s+\d{4}',
    re.IGNORECASE,
)


def _split_city_date_paragraphs(doc) -> int:
    """
    После перевода ищет параграфы вида «город ... дата» (с табом или большим
    блоком пробелов) и разбивает на 2 параграфа: город и дата, оба слева.

    Покрывает:
    - Акт: <w:jc=both> + <w:tab w:val="right"> + <w:tab/> в runs
    - Договор: <w:jc=center> + строка с 3+ подряд пробелами

    Не трогает:
    - Параграфы без испанской даты (русские шаблоны, тело документа)
    - Параграфы где до даты слишком много текста (>30 chars) — это не шапка
    - Параграфы где упоминается «Contrato» (это ссылка на договор-номер, не шапка)

    Returns: количество разбитых параграфов.
    """
    from docx.oxml.ns import qn
    from copy import deepcopy

    splits = 0

    # Собираем все параграфы (включая textbox-ы)
    all_paragraphs = _collect_all_paragraphs(doc)

    for paragraph in all_paragraphs:
        try:
            text = paragraph.text
        except Exception:
            continue

        if not text or not text.strip():
            continue

        # Не трогаем «Contrato n.º X de fecha Y» — это ссылка на договор
        if "Contrato" in text and "de fecha" in text:
            continue

        # Ищем испанскую дату
        date_match = _SPANISH_DATE_RE.search(text)
        if not date_match:
            continue

        date_str = date_match.group(0).strip()
        before_date = text[:date_match.start()].strip()

        # Если до даты слишком много текста — это не шапка
        if len(before_date) > 30:
            continue

        # Если до даты ничего — параграф только с датой, не трогаем
        if not before_date:
            continue

        # Это шапка «город + дата». Разбиваем.
        p_elem = paragraph._element

        # Сохраняем pPr оригинала для клонирования стиля (но почистим)
        old_pPr = p_elem.find(qn('w:pPr'))

        # Pack 35.10: сохраняем rPr из первого run перед удалением
        # (чтобы новые run'ы унаследовали шрифт/размер)
        old_runs = list(p_elem.findall(qn('w:r')))
        saved_rPr = None
        if old_runs:
            first_rPr = old_runs[0].find(qn('w:rPr'))
            if first_rPr is not None:
                from copy import deepcopy as _deepcopy
                saved_rPr = _deepcopy(first_rPr)

        # Удаляем все runs из оригинального параграфа
        for r in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r)
        # Удаляем все hyperlinks если есть
        for hl in list(p_elem.findall(qn('w:hyperlink'))):
            p_elem.remove(hl)

        # Очищаем форматирование pPr: убираем tabs, меняем jc на left
        if old_pPr is not None:
            # Удаляем <w:tabs>
            tabs_elem = old_pPr.find(qn('w:tabs'))
            if tabs_elem is not None:
                old_pPr.remove(tabs_elem)
            # Меняем <w:jc> на left
            jc_elem = old_pPr.find(qn('w:jc'))
            if jc_elem is not None:
                jc_elem.set(qn('w:val'), 'left')
            # Убираем firstLine indent (если есть)
            ind_elem = old_pPr.find(qn('w:ind'))
            if ind_elem is not None:
                # Не удаляем целиком — может остаться полезный отступ
                if ind_elem.get(qn('w:firstLine')):
                    del ind_elem.attrib[qn('w:firstLine')]

        # Создаём новый run с городом
        from docx.oxml import OxmlElement
        r_city = OxmlElement('w:r')
        # Pack 35.10: наследуем шрифт/размер из старого первого run
        if saved_rPr is not None:
            r_city.append(_deepcopy(saved_rPr))
        t_city = OxmlElement('w:t')
        t_city.text = before_date
        t_city.set(qn('xml:space'), 'preserve')
        r_city.append(t_city)
        p_elem.append(r_city)

        # Создаём НОВЫЙ параграф для даты, копируя pPr оригинала
        new_p = OxmlElement('w:p')
        if old_pPr is not None:
            new_p.append(deepcopy(old_pPr))
        r_date = OxmlElement('w:r')
        # Pack 35.10: наследуем rPr и в date run
        if saved_rPr is not None:
            r_date.append(_deepcopy(saved_rPr))
        t_date = OxmlElement('w:t')
        t_date.text = date_str
        t_date.set(qn('xml:space'), 'preserve')
        r_date.append(t_date)
        new_p.append(r_date)

        # Вставляем новый параграф сразу после оригинального
        p_elem.addnext(new_p)

        splits += 1
        log.info(f"[Pack 35.9] split city+date: {before_date!r} / {date_str!r}")

    return splits



async def translate_docx(
    docx_bytes: bytes,
    substitutions: Optional[SubstitutionDict] = None,
) -> bytes:
    """
    Главная функция: переводит DOCX с русского на испанский.

    Pack 15.2: pre-substitution теперь покрывает метки ИНН/КПП/ОГРН/БИК → NIF/KPP/etc
    + applicant.full_name_native + company.short_name + None → —.
    Pack 15.6: обходим textbox-ы (карточка реквизитов в выписке, шапка операций в header).
    Pack 33.9: static lookup для single-word фрагментов (защита от LLM-разговорчивости
    на коротких изолированных строках вроде «Описание»).
    """
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = _collect_all_paragraphs(doc)

    # Pack 15.6: дедупликация по id <w:p> элемента — на случай если один параграф
    # попал в список несколько раз (gridSpan-ячейки в таблицах python-docx).
    seen_elem_ids: set[int] = set()
    deduped: list[Paragraph] = []
    for p in paragraphs:
        elem_id = id(p._element)
        if elem_id in seen_elem_ids:
            continue
        seen_elem_ids.add(elem_id)
        deduped.append(p)
    paragraphs = deduped

    log.info(
        f"[translation] DOCX has {len(paragraphs)} paragraphs"
        + (f", {len(substitutions)} pre-substitutions" if substitutions else "")
    )

    targets: list[tuple[int, str]] = []
    static_hits = 0
    for idx, p in enumerate(paragraphs):
        text = _extract_paragraph_text(p)
        original_text = text  # для сравнения с результатом подстановки

        if substitutions:
            text = substitutions.apply(text)

        # Pack 15.5: КРИТИЧНЫЙ ФИКС.
        # Если pre-substitution что-то заменила, но текст после неё уже не нужно
        # отправлять в LLM (нет кириллицы) — мы ОБЯЗАНЫ записать его в DOCX
        # сейчас. Иначе параграф останется в исходном русском виде, потому что
        # _set_paragraph_text() позже вызывается только для targets (тех что
        # пошли в LLM).
        if text != original_text and _should_skip(text):
            _set_paragraph_text(p, text)
            log.info(f"[translation] Pre-sub-only update for paragraph {idx}")
            continue

        if not _should_skip(text):
            # Pack 33.9: static lookup ПЕРЕД отправкой в LLM.
            # Защищает от случая когда модель на короткой строке
            # ("Описание") отвечает разговорно вместо перевода.
            static_translation = _try_static_translation(text)
            if static_translation is not None:
                _set_paragraph_text(p, static_translation)
                static_hits += 1
                log.info(
                    "[translation] Pack 33.9 static lookup: %r -> %r (paragraph %d)",
                    text.strip(), static_translation, idx,
                )
                continue

            targets.append((idx, text))

    if static_hits > 0:
        log.info(
            "[translation] Pack 33.9: %d single-word fragments translated via static lookup "
            "(skipped LLM call)",
            static_hits,
        )

    if not targets:
        log.info("[translation] Nothing to translate in this DOCX")
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    log.info(f"[translation] {len(targets)} fragments to translate")

    all_translations: list[str] = []
    for batch_start in range(0, len(targets), BATCH_SIZE):
        batch = targets[batch_start:batch_start + BATCH_SIZE]
        batch_texts = [t[1] for t in batch]
        translated = await _translate_batch(batch_texts)
        all_translations.extend(translated)
        log.info(
            f"[translation] Batch {batch_start // BATCH_SIZE + 1} "
            f"({batch_start + 1}-{batch_start + len(batch)} / {len(targets)}) done"
        )

    for (paragraph_idx, _original), translated_text in zip(targets, all_translations):
        _set_paragraph_text(paragraphs[paragraph_idx], translated_text)

    # Pack 35.9: разбить шапку «город + дата» на 2 параграфа
    try:
        splits = _split_city_date_paragraphs(doc)
        if splits > 0:
            log.info(f"[Pack 35.9] split city+date paragraphs: {splits}")
    except Exception as e:
        log.warning(f"[Pack 35.9] split city+date failed: {e}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
