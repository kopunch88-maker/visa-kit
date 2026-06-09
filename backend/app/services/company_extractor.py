"""
Pack 26.0 — Извлечение реквизитов компании из DOCX (или plaintext).

Используется для UI «Загрузить реквизиты компании» в админке.

Pipeline:
1. DOCX → читаем все параграфы и таблицы python-docx → собираем plaintext
2. Plaintext → отправляем LLM с COMPANY_REQUISITES_PROMPT
3. Возвращаем dict готовый к создаию/обновлению Company через CompanyCreate/CompanyPatch

Преимущества vs Vision OCR:
- Текст из DOCX чистый (без OCR-шума, ошибок типа 7707038236 vs 7707038266)
- Дешевле в 5-10× (~500 vs ~3000 токенов)
- Быстрее (~2с vs ~10с)

Только DOCX в первой итерации. PDF/JPG — в следующих пакетах через Vision-путь.
"""
import io
import json
import logging
import re
from typing import Optional

from app.services.llm import get_llm_client
from app.services.ocr.prompts import COMPANY_REQUISITES_PROMPT

log = logging.getLogger(__name__)


class CompanyExtractError(Exception):
    """Не удалось извлечь реквизиты компании."""
    pass


def _read_docx_to_text(docx_bytes: bytes) -> str:
    """
    Читает DOCX и возвращает plaintext всех параграфов и ячеек таблиц.

    Сохраняет порядок чтения (как видит пользователь): параграфы → таблицы по строкам.
    Игнорирует пустые строки.
    """
    try:
        from docx import Document
    except ImportError:
        raise CompanyExtractError(
            "python-docx not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise CompanyExtractError(f"Failed to read DOCX: {e}")

    parts: list[str] = []

    # Параграфы верхнего уровня
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Таблицы — каждая ячейка как отдельная "строка"
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts[-3:]:  # дедупликация ячеек, повторяющихся в merged
                    parts.append(t)

    text = "\n".join(parts).strip()

    if not text:
        raise CompanyExtractError("DOCX file appears to be empty")

    if len(text) < 30:
        raise CompanyExtractError(
            f"DOCX content too short ({len(text)} chars) — likely not a requisites document"
        )

    return text


def _strip_json_fence(s: str) -> str:
    """Убирает ```json ... ``` обёртку если LLM её добавил, и ищет {...}."""
    s = s.strip()
    if s.startswith("```"):
        # снять ```json или ``` в начале и ``` в конце
        lines = s.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        s = "\n".join(lines).strip()

    first_brace = s.find("{")
    last_brace = s.rfind("}")
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        s = s[first_brace : last_brace + 1]
    return s


def _normalize_extracted_fields(data: dict) -> dict:
    """
    Минимальная нормализация полей после LLM:
    - тримминг строк
    - удаление пробелов внутри ИНН/КПП/ОГРН/счетов/БИК
    - конверсия пустых строк в null

    LLM может возвращать "  " вместо null, или ИНН с пробелом — чистим.
    """
    cleaned: dict = {}
    for key, value in data.items():
        if value is None:
            cleaned[key] = None
            continue
        if isinstance(value, str):
            v = value.strip()
            if v == "":
                cleaned[key] = None
                continue
            # Числовые поля — убрать пробелы и не-цифры
            if key in ("ogrn", "inn", "kpp", "bank_account", "bank_bic", "bank_correspondent_account"):
                v = re.sub(r"[^0-9]", "", v)
                if v == "":
                    cleaned[key] = None
                    continue
            cleaned[key] = v
        else:
            cleaned[key] = value
    return cleaned


# Pack 26.x — подсказка LLM для двуязычных выписок (испанский перевод + русский оригинал).
_BILINGUAL_HINT = (
    "ВАЖНО: документ может содержать ИСПАНСКИЙ перевод и РУССКИЙ оригинал "
    "выписки ЕГРЮЛ в одном файле (часто сначала идёт испанский перевод, затем "
    "русский оригинал). Все поля — особенно ФИО директора во всех падежах "
    "(склонения) и юридический адрес — бери ИЗ РУССКОЙ части документа. "
    "Испанский текст используй только если соответствующего поля нет в русской "
    "части. Числовые реквизиты (ИНН, КПП, ОГРН, расчётный счёт, корр. счёт, БИК) "
    "одинаковы в обеих частях — бери как есть."
)


async def _extract_fields_from_text(text: str, *, bilingual_hint: bool = False) -> dict:
    """
    Pack 26.x — общий шаг «plaintext → поля компании через LLM».

    Используется и DOCX-путём (Pack 26.0), и PDF-путём (Pack 26.x).
    bilingual_hint=True добавляет инструкцию про испанский+русский в одном файле.
    """
    client = get_llm_client()

    hint = f"\n\n{_BILINGUAL_HINT}" if bilingual_hint else ""

    try:
        response_text = await client.complete(
            system="You are a precise data extraction assistant. Always return strict JSON.",
            user=f"{COMPANY_REQUISITES_PROMPT}{hint}\n\n--- DOCUMENT TEXT ---\n{text}\n--- END ---",
            max_tokens=2048,
            temperature=0.0,
        )
    except Exception as e:
        log.error(f"LLM error in company extraction: {e}", exc_info=True)
        raise CompanyExtractError(f"LLM error: {e}")

    cleaned_response = _strip_json_fence(response_text)

    try:
        parsed = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        log.error(f"Failed to parse LLM JSON: {e}\nResponse head: {response_text[:500]}")
        raise CompanyExtractError(
            f"LLM returned invalid JSON. Head: {response_text[:200]}"
        )

    if not isinstance(parsed, dict):
        raise CompanyExtractError(f"LLM returned non-dict: {type(parsed).__name__}")

    normalized = _normalize_extracted_fields(parsed)

    log.info(
        f"Pack 26.x: extracted {sum(1 for v in normalized.values() if v is not None)}"
        f"/{len(normalized)} fields"
    )

    return normalized


async def _read_pdf_to_text(pdf_bytes: bytes) -> str:
    """
    Pack 26.x — PDF → plaintext через гибридный экстрактор Pack 39.0-C.

    Текстовый PDF → pypdf (бесплатно). Скан-PDF → Vision (claude-sonnet-4-5).
    Импорт ленивый: модуль тянет pypdfium2 и не нужен на DOCX-пути.
    """
    from app.services.final_submission.extractor import extract_document_text

    result = await extract_document_text(pdf_bytes, "company.pdf", "application/pdf")

    if result.method == "failed" or not result.text:
        raise CompanyExtractError(
            f"Не удалось прочитать PDF: {result.error or 'пустой текст'}"
        )

    text = result.text.strip()
    if len(text) < 30:
        raise CompanyExtractError(
            f"PDF content too short ({len(text)} chars) — likely not a requisites document"
        )

    log.info(
        f"Pack 26.x: PDF extracted {len(text)} chars via '{result.method}' "
        f"(pages={result.page_count}, cost=${result.cost_usd})"
    )
    return text


async def extract_company_from_docx(docx_bytes: bytes) -> dict:
    """
    Pack 26.0 — извлечение реквизитов компании из DOCX.

    Args:
        docx_bytes: содержимое .docx файла

    Returns:
        dict с полями совместимыми с CompanyCreate (плюс доп. charter_capital в notes)

    Raises:
        CompanyExtractError при любой проблеме (плохой файл, плохой ответ LLM, и т.д.)
    """
    text = _read_docx_to_text(docx_bytes)
    log.info(f"Pack 26.0: extracted {len(text)} chars from DOCX, calling LLM...")
    return await _extract_fields_from_text(text, bilingual_hint=False)


async def extract_company_from_file(
    content: bytes,
    filename: str,
    mime_type: str = "",
) -> dict:
    """
    Pack 26.x — единая точка входа: DOCX или PDF.

    .docx → python-docx (как Pack 26.0).
    .pdf  → гибрид pypdf/Vision + двуязычная подсказка (исп. перевод + рус. оригинал).

    Args:
        content: байты файла
        filename: оригинальное имя (для определения расширения)
        mime_type: content-type из UploadFile (fallback к расширению)

    Returns:
        dict с полями совместимыми с CompanyCreate

    Raises:
        CompanyExtractError при неподдерживаемом типе или проблеме извлечения
    """
    name = (filename or "").lower()

    if name.endswith(".docx") or "wordprocessingml" in (mime_type or ""):
        return await extract_company_from_docx(content)

    if name.endswith(".pdf") or (mime_type or "") == "application/pdf":
        text = await _read_pdf_to_text(content)
        log.info(f"Pack 26.x: extracted {len(text)} chars from PDF, calling LLM...")
        return await _extract_fields_from_text(text, bilingual_hint=True)

    raise CompanyExtractError(
        f"Неподдерживаемый тип файла: {filename}. Нужен .docx или .pdf"
    )
