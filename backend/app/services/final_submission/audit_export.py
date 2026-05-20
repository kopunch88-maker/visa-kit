# -*- coding: utf-8 -*-
"""
Pack 39.0-F — Final Submission Audit DOCX export.

Генерирует DOCX-отчёт из FinalSubmissionAuditReport + всех FinalSubmissionFinding
для скачивания менеджером. Включает полную историю — open, acknowledged, dismissed.
Acknowledged/dismissed рендерятся зачёркнутыми + цветной плашкой статуса.

Структура DOCX:
1. Заголовок + субтитр
2. Метаданные кейса (таблица 2×N)
3. Баннер вердикта с цветным фоном
4. Inspector summary (резюме от ИИ)
5. Счётчики severity
6. Findings по 8 категориям A-H
7. Footer с датой генерации

Использует python-docx (уже в requirements.txt).
"""
import io
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sqlmodel import Session, select

from app.models import (
    FinalSubmissionAuditReport,
    FinalSubmissionFinding,
    FinalSubmissionVerdict,
    FinalSubmissionCategory,
    FinalSubmissionSeverity,
    FinalSubmissionFindingStatus,
    Application,
    Applicant,
)

log = logging.getLogger(__name__)


# ====================================================================
# Цветовая схема (синхронизирована с Pack 37.0 audit_export.py)
# ====================================================================

COLOR_CRITICAL = RGBColor(0xDC, 0x26, 0x26)
COLOR_WARNING = RGBColor(0xD9, 0x77, 0x06)
COLOR_INFO = RGBColor(0x25, 0x63, 0xEB)
COLOR_SUCCESS = RGBColor(0x16, 0xA3, 0x4A)
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
COLOR_DARK = RGBColor(0x11, 0x18, 0x27)

VERDICT_LABELS = {
    FinalSubmissionVerdict.PASS_: ("ПАКЕТ ГОТОВ К ПОДАЧЕ", COLOR_SUCCESS, "✓"),
    FinalSubmissionVerdict.WARN: ("ЕСТЬ ПРЕДУПРЕЖДЕНИЯ", COLOR_WARNING, "⚠"),
    FinalSubmissionVerdict.FAIL: ("НАЙДЕНЫ КРИТИЧЕСКИЕ ОШИБКИ", COLOR_CRITICAL, "✗"),
}

VERDICT_BG = {
    FinalSubmissionVerdict.PASS_: "ECFDF5",   # green-50
    FinalSubmissionVerdict.WARN: "FFFBEB",    # amber-50
    FinalSubmissionVerdict.FAIL: "FEF2F2",    # red-50
}

SEVERITY_LABELS = {
    FinalSubmissionSeverity.CRITICAL: ("КРИТИЧНО", COLOR_CRITICAL),
    FinalSubmissionSeverity.WARNING: ("ПРЕДУПРЕЖДЕНИЕ", COLOR_WARNING),
    FinalSubmissionSeverity.INFO: ("ЗАМЕЧАНИЕ", COLOR_INFO),
}

CATEGORY_LABELS = {
    FinalSubmissionCategory.A_IDENTITY: "A. Личные данные",
    FinalSubmissionCategory.B_NUMERIC: "B. Суммы и числа",
    FinalSubmissionCategory.C_DATES: "C. Даты",
    FinalSubmissionCategory.D_COMPANY: "D. Реквизиты компании",
    FinalSubmissionCategory.E_TRANSLATION: "E. Переводы jurada",
    FinalSubmissionCategory.F_COMPLETENESS: "F. Комплектность пакета",
    FinalSubmissionCategory.G_QUALITY: "G. Качество сканов",
    FinalSubmissionCategory.H_STALE: "H. Хвосты прошлых клиентов",
}

CATEGORY_ORDER = [
    FinalSubmissionCategory.A_IDENTITY,
    FinalSubmissionCategory.B_NUMERIC,
    FinalSubmissionCategory.C_DATES,
    FinalSubmissionCategory.D_COMPANY,
    FinalSubmissionCategory.E_TRANSLATION,
    FinalSubmissionCategory.F_COMPLETENESS,
    FinalSubmissionCategory.G_QUALITY,
    FinalSubmissionCategory.H_STALE,
]

STATUS_LABELS = {
    FinalSubmissionFindingStatus.OPEN: ("⚠ Требует исправления", COLOR_WARNING),
    FinalSubmissionFindingStatus.ACKNOWLEDGED: ("✓ Учтено — идёт исправление", COLOR_SUCCESS),
    FinalSubmissionFindingStatus.DISMISSED: ("✗ Отклонено как false positive", COLOR_MUTED),
}


# ====================================================================
# Helpers для docx стилей (копия из audit_export.py)
# ====================================================================

def _set_cell_bg(cell, color_hex: str):
    """Заливка ячейки таблицы цветом (color_hex без #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             size: Optional[int] = None, color: Optional[RGBColor] = None,
             font: Optional[str] = None, strike: bool = False):
    """Добавить run с форматированием. strike=True — зачёркнутый текст."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font is not None:
        run.font.name = font
    if strike:
        # python-docx не поддерживает strike напрямую — добавляем через XML
        rPr = run._element.get_or_add_rPr()
        strike_el = OxmlElement("w:strike")
        strike_el.set(qn("w:val"), "true")
        rPr.append(strike_el)
    return run


def _add_horizontal_line(doc: Document):
    """Тонкая горизонтальная линия."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _format_duration_ru(ms: Optional[int]) -> str:
    if ms is None:
        return "—"
    seconds = ms // 1000
    if seconds < 60:
        return f"{seconds} сек"
    minutes = seconds // 60
    rem = seconds % 60
    return f"{minutes} мин {rem:02d} сек" if rem else f"{minutes} мин"


def _format_datetime_ru(dt: Optional[datetime]) -> str:
    if dt is None:
        return "—"
    months_ru = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря",
    ]
    return f"{dt.day} {months_ru[dt.month - 1]} {dt.year}, {dt.hour:02d}:{dt.minute:02d}"


# ====================================================================
# Сборка имени файла
# ====================================================================

def build_filename(applicant: Optional[Applicant], report: FinalSubmissionAuditReport) -> str:
    """
    final_check_<Фамилия>_<Имя>_<YYYY-MM-DD>.docx
    """
    date_str = (report.started_at or datetime.utcnow()).strftime("%Y-%m-%d")
    if applicant and (applicant.last_name_native or applicant.first_name_native):
        last = (applicant.last_name_native or "").strip().replace(" ", "_")
        first = (applicant.first_name_native or "").strip().replace(" ", "_")
        safe = "".join(c for c in f"{last}_{first}" if c.isalnum() or c in "_-")
        return f"final_check_{safe}_{date_str}.docx"
    return f"final_check_report_{report.application_id}_{date_str}.docx"


# ====================================================================
# Главная функция
# ====================================================================

def build_final_submission_audit_docx(
    report_id: int, session: Session
) -> tuple[bytes, str]:
    """
    Собирает DOCX отчёта финальной проверки.

    Returns:
        (docx_bytes, filename)
    """
    report = session.get(FinalSubmissionAuditReport, report_id)
    if not report:
        raise ValueError(f"FinalSubmissionAuditReport {report_id} not found")

    application = session.get(Application, report.application_id) if report.application_id else None
    applicant = session.get(Applicant, report.applicant_id) if report.applicant_id else None

    findings = session.exec(
        select(FinalSubmissionFinding)
        .where(FinalSubmissionFinding.report_id == report_id)
    ).all()

    # Сортировка: critical → warning → info, внутри — по sort_order
    sev_priority = {
        FinalSubmissionSeverity.CRITICAL: 0,
        FinalSubmissionSeverity.WARNING: 1,
        FinalSubmissionSeverity.INFO: 2,
    }
    findings_sorted = sorted(
        findings,
        key=lambda f: (sev_priority.get(f.severity, 99), f.sort_order, f.id),
    )

    doc = Document()

    # Поля документа
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Базовый шрифт
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ========================================================================
    # 1. Шапка
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "ФИНАЛЬНАЯ ПРОВЕРКА ПАКЕТА ДОКУМЕНТОВ",
             bold=True, size=18, color=COLOR_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, "Симуляция приёма физических документов в консульстве",
             italic=True, size=11, color=COLOR_MUTED)

    doc.add_paragraph()

    # ========================================================================
    # 2. Метаданные кейса (таблица 2×N)
    # ========================================================================
    applicant_name = "—"
    if applicant:
        parts = [
            applicant.last_name_native or "",
            applicant.first_name_native or "",
            applicant.middle_name_native or "",
        ]
        applicant_name = " ".join(p for p in parts if p).strip() or "—"

    meta_rows = [
        ("Заявка:", f"#{report.application_id}"),
        ("Заявитель:", applicant_name),
        ("Дата проверки:", _format_datetime_ru(report.started_at)),
        ("Длительность:", _format_duration_ru(report.duration_ms)),
        ("Модель ИИ:", (report.model_used or "—").replace("anthropic/", "")),
    ]
    if report.cost_usd is not None:
        meta_rows.append(("Стоимость прогона:", f"${report.cost_usd}"))

    # Снэпшот документов: сколько и каких категорий
    if report.document_categories_snapshot:
        snap = report.document_categories_snapshot
        total_docs = sum(snap.values()) if isinstance(snap, dict) else 0
        if total_docs:
            meta_rows.append(("Документов в пакете:", str(total_docs)))

    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.autofit = False
    for i, (k, v) in enumerate(meta_rows):
        c1 = meta_table.cell(i, 0)
        c2 = meta_table.cell(i, 1)
        c1.width = Cm(5)
        c2.width = Cm(12)
        c1.text = ""
        c2.text = ""
        _add_run(c1.paragraphs[0], k, bold=True, size=10, color=COLOR_MUTED)
        _add_run(c2.paragraphs[0], v, size=10, color=COLOR_DARK)

    doc.add_paragraph()

    # ========================================================================
    # 3. Баннер вердикта (таблица 1×1 с заливкой)
    # ========================================================================
    verdict_text, verdict_color, verdict_icon = VERDICT_LABELS[report.verdict]
    light_bg = VERDICT_BG[report.verdict]

    verdict_table = doc.add_table(rows=1, cols=1)
    cell = verdict_table.cell(0, 0)
    _set_cell_bg(cell, light_bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"{verdict_icon}  ВЕРДИКТ: {verdict_text}",
             bold=True, size=20, color=verdict_color)

    # Счётчики
    counts = report.summary_counts or {}
    total = counts.get("total", len(findings))
    critical = counts.get("critical", 0)
    warning = counts.get("warning", 0)
    info = counts.get("info", 0)

    summary_p = cell.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(summary_p, f"Всего найдено: {total}  ·  ",
             size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Критично: {critical}",
             bold=True, size=11, color=COLOR_CRITICAL)
    _add_run(summary_p, "  ·  ", size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Предупреждений: {warning}",
             bold=True, size=11, color=COLOR_WARNING)
    _add_run(summary_p, "  ·  ", size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Замечаний: {info}",
             bold=True, size=11, color=COLOR_INFO)

    doc.add_paragraph()

    # ========================================================================
    # 4. Inspector summary
    # ========================================================================
    if report.inspector_summary:
        h = doc.add_paragraph()
        _add_run(h, "РЕЗЮМЕ ОТ ВИЗОВОГО ИНСПЕКТОРА",
                 bold=True, size=12, color=COLOR_DARK)
        _add_horizontal_line(doc)
        p = doc.add_paragraph()
        _add_run(p, report.inspector_summary, size=10, color=COLOR_DARK, italic=True)
        doc.add_paragraph()

    # ========================================================================
    # 5. Findings — группируем по категориям A-H
    # ========================================================================
    findings_by_category = {}
    for f in findings_sorted:
        if f.category not in findings_by_category:
            findings_by_category[f.category] = []
        findings_by_category[f.category].append(f)

    has_any = False
    for category in CATEGORY_ORDER:
        items = findings_by_category.get(category)
        if not items:
            continue
        has_any = True

        section_h = doc.add_paragraph()
        _add_run(section_h, f"{CATEGORY_LABELS[category]} ({len(items)})",
                 bold=True, size=14, color=COLOR_DARK)
        _add_horizontal_line(doc)

        for idx, f in enumerate(items, 1):
            _render_finding(doc, idx, f)

        doc.add_paragraph()

    if not has_any:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, "✓ Несоответствий не найдено",
                 bold=True, size=14, color=COLOR_SUCCESS)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p2, "AI-инспектор проверил пакет и не обнаружил проблем. Можно подавать.",
                 size=11, color=COLOR_MUTED)

    # ========================================================================
    # 6. Footer
    # ========================================================================
    doc.add_paragraph()
    _add_horizontal_line(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        footer,
        f"Отчёт сгенерирован автоматически системой проверки документов "
        f"{_format_datetime_ru(datetime.utcnow())} UTC",
        italic=True, size=8, color=COLOR_MUTED,
    )

    # Сохраняем в bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    filename = build_filename(applicant, report)

    log.info(
        f"[final_audit_export] Generated DOCX for report {report_id}: "
        f"{len(docx_bytes)} bytes, filename={filename}"
    )

    return docx_bytes, filename


def _render_finding(doc: Document, idx: int, f: FinalSubmissionFinding) -> None:
    """Один finding — заголовок, описание, рекомендация, документы, статус.

    Если status != OPEN — текст рисуется зачёркнутым с цветной плашкой статуса.
    """
    sev_label, sev_color = SEVERITY_LABELS[f.severity]
    status_label, status_color = STATUS_LABELS[f.status]
    is_resolved = f.status != FinalSubmissionFindingStatus.OPEN

    # Заголовок: номер + severity + title
    h = doc.add_paragraph()
    _add_run(h, f"{idx}. ", bold=True, size=11, color=sev_color, strike=is_resolved)
    _add_run(h, f"[{sev_label}] ",
             bold=True, size=10, color=sev_color, strike=is_resolved)
    _add_run(h, f.title, bold=True, size=11, color=COLOR_DARK, strike=is_resolved)

    # Описание
    if f.description:
        p = doc.add_paragraph()
        _add_run(p, f.description, size=10, color=COLOR_DARK, strike=is_resolved)

    # Рекомендация
    if f.recommendation:
        p = doc.add_paragraph()
        _add_run(p, "Рекомендация: ", bold=True, size=10, color=COLOR_DARK, strike=is_resolved)
        _add_run(p, f.recommendation, size=10, color=COLOR_DARK, strike=is_resolved)

    # Affected documents
    if f.affected_documents:
        p = doc.add_paragraph()
        _add_run(p, "Документы: ", bold=True, size=9, color=COLOR_MUTED, strike=is_resolved)
        parts = []
        for ad in f.affected_documents:
            if isinstance(ad, dict):
                filename = ad.get("filename", "?")
                page = ad.get("page")
                parts.append(f"{filename}" + (f" (стр. {page})" if page else ""))
        if parts:
            _add_run(p, ", ".join(parts), size=9, color=COLOR_DARK,
                     font="Consolas", strike=is_resolved)

    # Values found — таблица
    if f.values_found and isinstance(f.values_found, dict) and f.values_found:
        t = doc.add_table(rows=len(f.values_found), cols=2)
        for i, (k, v) in enumerate(f.values_found.items()):
            c1 = t.cell(i, 0)
            c2 = t.cell(i, 1)
            c1.width = Cm(5)
            c2.width = Cm(12)
            c1.text = ""
            c2.text = ""
            _add_run(c1.paragraphs[0], str(k) + ":", bold=True, size=9,
                     color=COLOR_MUTED, strike=is_resolved)
            _add_run(c2.paragraphs[0], str(v), size=9, color=COLOR_DARK,
                     font="Consolas", strike=is_resolved)

    # Field name (опционально)
    if f.field_name:
        p = doc.add_paragraph()
        _add_run(p, "Поле: ", bold=True, size=9, color=COLOR_MUTED, strike=is_resolved)
        _add_run(p, f.field_name, size=9, color=COLOR_DARK,
                 font="Consolas", strike=is_resolved)

    # Статус
    sp = doc.add_paragraph()
    _add_run(sp, "Статус: ", bold=True, size=9, color=COLOR_MUTED)
    _add_run(sp, status_label, bold=True, size=9, color=status_color)
    if f.resolved_at:
        _add_run(sp, f"  ({_format_datetime_ru(f.resolved_at)}",
                 size=9, color=COLOR_MUTED)
        if f.resolved_by:
            _add_run(sp, f", {f.resolved_by}", size=9, color=COLOR_MUTED)
        _add_run(sp, ")", size=9, color=COLOR_MUTED)
    if f.resolution_note:
        np = doc.add_paragraph()
        _add_run(np, "Заметка: ", bold=True, size=9, color=COLOR_MUTED)
        _add_run(np, f.resolution_note, italic=True, size=9, color=COLOR_DARK)

    # Разделитель
    _add_horizontal_line(doc)
