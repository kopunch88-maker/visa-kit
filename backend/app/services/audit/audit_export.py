# -*- coding: utf-8 -*-
"""
Pack 37.1 — Audit report DOCX export.

Генерирует DOCX-отчёт из AuditReport + всех его AuditFinding для скачивания
менеджером. Включает полную историю — открытые, принятые, отклонённые,
исправленные вручную findings.

Структура DOCX:
1. Заголовок (имя клиента, ID заявки, дата проверки)
2. Метаданные прогона (длительность, стоимость, модель ИИ)
3. Светофор вердикта с цветной подложкой
4. Сводка (счётчики по severity)
5. Резюме от ИИ
6. Findings по категориям, отсортированные по severity (critical → warning → info)

Используется python-docx (уже в проекте) — без новых зависимостей.
"""
import io
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sqlmodel import Session, select

from app.models import (
    AuditReport,
    AuditFinding,
    AuditVerdict,
    AuditCategory,
    AuditSeverity,
    AuditFindingStatus,
    Application,
    Applicant,
)

log = logging.getLogger(__name__)


# ====================================================================
# Цветовая схема
# ====================================================================

COLOR_CRITICAL = RGBColor(0xDC, 0x26, 0x26)  # red-600
COLOR_WARNING = RGBColor(0xD9, 0x77, 0x06)   # amber-600
COLOR_INFO = RGBColor(0x25, 0x63, 0xEB)      # blue-600
COLOR_SUCCESS = RGBColor(0x16, 0xA3, 0x4A)   # green-600
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)     # gray-500
COLOR_DARK = RGBColor(0x11, 0x18, 0x27)      # gray-900

VERDICT_LABELS = {
    AuditVerdict.PASS_: ("ГОТОВ К ПОДАЧЕ", COLOR_SUCCESS, "✓"),
    AuditVerdict.WARN: ("ПОДАВАТЬ С РИСКОМ", COLOR_WARNING, "⚠"),
    AuditVerdict.FAIL: ("НЕ ГОТОВ К ПОДАЧЕ", COLOR_CRITICAL, "✗"),
}

SEVERITY_LABELS = {
    AuditSeverity.CRITICAL: ("КРИТИЧНО", COLOR_CRITICAL),
    AuditSeverity.WARNING: ("ПРЕДУПРЕЖДЕНИЕ", COLOR_WARNING),
    AuditSeverity.INFO: ("ЗАМЕЧАНИЕ", COLOR_INFO),
}

CATEGORY_LABELS = {
    AuditCategory.IDENTITY: "Личные данные",
    AuditCategory.FINANCIAL: "Финансы",
    AuditCategory.COMPANY: "Компания",
    AuditCategory.EDUCATION: "Образование",
    AuditCategory.SPAIN_PACK: "Испанские документы",
    AuditCategory.FORMAL: "Комплектность пакета",
}

STATUS_LABELS = {
    AuditFindingStatus.OPEN: ("⚠ Требует решения", COLOR_WARNING),
    AuditFindingStatus.ACCEPTED: ("✓ Принято — фикс применён", COLOR_SUCCESS),
    AuditFindingStatus.DISMISSED: ("✗ Отклонено менеджером", COLOR_MUTED),
    AuditFindingStatus.MANUALLY_FIXED: ("✎ Исправлено вручную", COLOR_INFO),
}


# ====================================================================
# Helpers для docx стилей
# ====================================================================

def _set_cell_bg(cell, color_hex: str):
    """Заливка ячейки таблицы цветом (color_hex без #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             size: Optional[int] = None, color: Optional[RGBColor] = None,
             font: Optional[str] = None):
    """Добавить run с форматированием."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font is not None:
        run.font.name = font
    return run


def _add_horizontal_line(doc: Document):
    """Тонкая горизонтальная разделительная линия."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _format_duration_ru(ms: Optional[int]) -> str:
    if ms is None:
        return "—"
    seconds = ms // 1000
    if seconds < 60:
        return f"{seconds} сек"
    minutes = seconds // 60
    rem = seconds % 60
    return f"{minutes} мин {rem:02d} сек" if rem else f"{minutes} мин"


def _format_datetime_ru(dt: Optional[datetime]) -> str:
    if dt is None:
        return "—"
    months_ru = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря",
    ]
    return f"{dt.day} {months_ru[dt.month - 1]} {dt.year}, {dt.hour:02d}:{dt.minute:02d}"


# ====================================================================
# Сборка имени файла
# ====================================================================

def build_filename(applicant: Optional[Applicant], report: AuditReport) -> str:
    """
    audit_<Фамилия>_<Имя>_<YYYY-MM-DD>.docx
    Если applicant отсутствует — audit_<application_id>_<YYYY-MM-DD>.docx
    """
    date_str = (report.started_at or datetime.utcnow()).strftime("%Y-%m-%d")
    if applicant and (applicant.last_name_native or applicant.first_name_native):
        last = (applicant.last_name_native or "").strip().replace(" ", "_")
        first = (applicant.first_name_native or "").strip().replace(" ", "_")
        # Чистим небезопасные символы
        safe = "".join(c for c in f"{last}_{first}" if c.isalnum() or c in "_-")
        return f"audit_{safe}_{date_str}.docx"
    return f"audit_report_{report.application_id}_{date_str}.docx"


# ====================================================================
# Главная функция
# ====================================================================

def build_audit_report_docx(report_id: int, session: Session) -> tuple[bytes, str]:
    """
    Собирает DOCX отчёта аудита.

    Returns:
        (docx_bytes, filename)
    """
    report = session.get(AuditReport, report_id)
    if not report:
        raise ValueError(f"AuditReport {report_id} not found")

    application = session.get(Application, report.application_id)
    applicant = application.applicant if application else None

    findings = session.exec(
        select(AuditFinding)
        .where(AuditFinding.report_id == report_id)
    ).all()

    # Сортировка: critical → warning → info, внутри — по sort_order, потом id
    sev_priority = {
        AuditSeverity.CRITICAL: 0,
        AuditSeverity.WARNING: 1,
        AuditSeverity.INFO: 2,
    }
    findings_sorted = sorted(
        findings,
        key=lambda f: (sev_priority.get(f.severity, 99), f.sort_order, f.id),
    )

    doc = Document()

    # Поля документа
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Базовый шрифт
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ========================================================================
    # 1. Шапка
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "ОТЧЁТ О ПРОВЕРКЕ ПАКЕТА ДОКУМЕНТОВ",
             bold=True, size=18, color=COLOR_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, "Симуляция приёма документов в консульстве",
             italic=True, size=11, color=COLOR_MUTED)

    doc.add_paragraph()  # пустая

    # ========================================================================
    # 2. Метаданные кейса (таблица 2×N)
    # ========================================================================
    applicant_name = "—"
    if applicant:
        parts = [
            applicant.last_name_native or "",
            applicant.first_name_native or "",
            applicant.middle_name_native or "",
        ]
        applicant_name = " ".join(p for p in parts if p).strip() or "—"

    meta_rows = [
        ("Заявка:", f"#{report.application_id}"),
        ("Заявитель:", applicant_name),
        ("Дата проверки:", _format_datetime_ru(report.started_at)),
        ("Длительность:", _format_duration_ru(report.duration_ms)),
        ("Модель ИИ:", report.model_used or "—"),
    ]
    if report.cost_usd is not None:
        meta_rows.append(("Стоимость прогона:", f"${report.cost_usd}"))

    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.autofit = False
    for i, (k, v) in enumerate(meta_rows):
        c1 = meta_table.cell(i, 0)
        c2 = meta_table.cell(i, 1)
        c1.width = Cm(5)
        c2.width = Cm(12)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        _add_run(p1, k, bold=True, size=10, color=COLOR_MUTED)
        p2 = c2.paragraphs[0]
        _add_run(p2, v, size=10, color=COLOR_DARK)

    doc.add_paragraph()

    # ========================================================================
    # 3. Светофор вердикта (большая таблица 1×1 с заливкой)
    # ========================================================================
    verdict_text, verdict_color, verdict_icon = VERDICT_LABELS[report.verdict]
    color_hex = "{:02X}{:02X}{:02X}".format(verdict_color[0], verdict_color[1], verdict_color[2])

    # Светлая заливка — используем альфа-аналог через светлый фон
    light_bg = {
        AuditVerdict.PASS_: "ECFDF5",   # green-50
        AuditVerdict.WARN: "FFFBEB",    # amber-50
        AuditVerdict.FAIL: "FEF2F2",    # red-50
    }[report.verdict]

    verdict_table = doc.add_table(rows=1, cols=1)
    cell = verdict_table.cell(0, 0)
    _set_cell_bg(cell, light_bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Очистим default параграф и добавим свои
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"{verdict_icon}  ВЕРДИКТ: {verdict_text}",
             bold=True, size=20, color=verdict_color)

    # Сводка
    counts = report.summary_counts or {}
    total = counts.get("total", len(findings))
    critical = counts.get("critical", 0)
    warning = counts.get("warning", 0)
    info = counts.get("info", 0)

    summary_p = cell.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(summary_p, f"Всего найдено: {total}  ·  ",
             size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Критично: {critical}",
             bold=True, size=11, color=COLOR_CRITICAL)
    _add_run(summary_p, "  ·  ", size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Предупреждений: {warning}",
             bold=True, size=11, color=COLOR_WARNING)
    _add_run(summary_p, "  ·  ", size=11, color=COLOR_MUTED)
    _add_run(summary_p, f"Замечаний: {info}",
             bold=True, size=11, color=COLOR_INFO)

    doc.add_paragraph()

    # ========================================================================
    # 4. Резюме от ИИ (если есть)
    # ========================================================================
    llm_summary = counts.get("_llm_summary") if isinstance(counts, dict) else None
    if llm_summary:
        h = doc.add_paragraph()
        _add_run(h, "РЕЗЮМЕ", bold=True, size=12, color=COLOR_DARK)
        _add_horizontal_line(doc)
        p = doc.add_paragraph()
        _add_run(p, str(llm_summary), size=10, color=COLOR_DARK)
        doc.add_paragraph()

    # ========================================================================
    # 5. Findings — группируем по severity
    # ========================================================================
    by_severity = {
        AuditSeverity.CRITICAL: [f for f in findings_sorted if f.severity == AuditSeverity.CRITICAL],
        AuditSeverity.WARNING: [f for f in findings_sorted if f.severity == AuditSeverity.WARNING],
        AuditSeverity.INFO: [f for f in findings_sorted if f.severity == AuditSeverity.INFO],
    }

    for severity, items in by_severity.items():
        if not items:
            continue
        sev_label, sev_color = SEVERITY_LABELS[severity]

        section_h = doc.add_paragraph()
        _add_run(section_h, f"{sev_label} ({len(items)})",
                 bold=True, size=14, color=sev_color)
        _add_horizontal_line(doc)

        for idx, f in enumerate(items, 1):
            _render_finding(doc, idx, f)

        doc.add_paragraph()  # отступ между секциями

    # Если совсем findings нет — отдельное сообщение
    if not findings_sorted:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, "✓ Несоответствий не найдено",
                 bold=True, size=14, color=COLOR_SUCCESS)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p2, "ИИ-аудитор проверил пакет и не обнаружил проблем. Можно подавать.",
                 size=11, color=COLOR_MUTED)

    # ========================================================================
    # 6. Footer
    # ========================================================================
    doc.add_paragraph()
    _add_horizontal_line(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        footer,
        f"Отчёт сгенерирован автоматически системой проверки документов "
        f"{_format_datetime_ru(datetime.utcnow())} UTC",
        italic=True, size=8, color=COLOR_MUTED,
    )

    # Сохраняем в bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    filename = build_filename(applicant, report)

    log.info(
        f"[audit_export] Generated DOCX for report {report_id}: "
        f"{len(docx_bytes)} bytes, filename={filename}"
    )

    return docx_bytes, filename


def _render_finding(doc: Document, idx: int, f: AuditFinding) -> None:
    """Один finding — карточка с заголовком, deta diff и статусом."""

    sev_label, sev_color = SEVERITY_LABELS[f.severity]
    status_label, status_color = STATUS_LABELS[f.status]

    # Заголовок: номер + категория + title
    h = doc.add_paragraph()
    _add_run(h, f"{idx}. ", bold=True, size=11, color=sev_color)
    _add_run(h, f"[{CATEGORY_LABELS[f.category]}] ",
             bold=True, size=10, color=COLOR_MUTED)
    _add_run(h, f.title, bold=True, size=11, color=COLOR_DARK)

    # Описание
    if f.description:
        p = doc.add_paragraph()
        _add_run(p, f.description, size=10, color=COLOR_DARK)

    # Поле / текущее / предлагаемое — таблица
    has_diff = (
        f.field_path is not None
        or f.current_value is not None
        or f.suggested_value is not None
    )
    if has_diff:
        rows = []
        if f.field_path:
            rows.append(("Поле:", f.field_path))
        if f.current_value is not None:
            rows.append(("Сейчас:", str(f.current_value) or "(пусто)"))
        if f.suggested_value is not None:
            rows.append(("Предлагается:", str(f.suggested_value)))

        if rows:
            t = doc.add_table(rows=len(rows), cols=2)
            for i, (k, v) in enumerate(rows):
                c1 = t.cell(i, 0)
                c2 = t.cell(i, 1)
                c1.width = Cm(3.5)
                c2.width = Cm(13.5)
                c1.text = ""
                c2.text = ""
                _add_run(c1.paragraphs[0], k, bold=True, size=9, color=COLOR_MUTED)
                # Подкрашиваем "Сейчас" красным, "Предлагается" зелёным
                v_color = COLOR_DARK
                if k == "Сейчас:":
                    v_color = COLOR_CRITICAL
                elif k == "Предлагается:":
                    v_color = COLOR_SUCCESS
                _add_run(c2.paragraphs[0], v, size=9, color=v_color, font="Consolas")

    # Обоснование (evidence) — мелким курсивом
    if f.evidence:
        p = doc.add_paragraph()
        _add_run(p, "Обоснование: ", bold=True, size=9, color=COLOR_MUTED)
        _add_run(p, f.evidence, italic=True, size=9, color=COLOR_MUTED)

    # Статус
    sp = doc.add_paragraph()
    _add_run(sp, "Статус: ", bold=True, size=9, color=COLOR_MUTED)
    _add_run(sp, status_label, bold=True, size=9, color=status_color)
    if f.resolved_at:
        _add_run(sp, f"  ({_format_datetime_ru(f.resolved_at)}", size=9, color=COLOR_MUTED)
        if f.resolved_by:
            _add_run(sp, f", {f.resolved_by}", size=9, color=COLOR_MUTED)
        _add_run(sp, ")", size=9, color=COLOR_MUTED)
    if f.resolution_note:
        np = doc.add_paragraph()
        _add_run(np, "Заметка: ", bold=True, size=9, color=COLOR_MUTED)
        _add_run(np, f.resolution_note, italic=True, size=9, color=COLOR_DARK)

    # Разделитель
    _add_horizontal_line(doc)
