# -*- coding: utf-8 -*-
"""
Pack 37.0-B — Document extractor.

Извлекает plain-text из сгенерированных DOCX/PDF документов пакета для
передачи в LLM-аудитор. Без этого аудитор не увидит, что реально написано
в финальных файлах — только поля applicant в БД.

Стратегия (гибрид):
1. Пробуем взять готовый файл из R2 (быстро, как при «Сгенерировать пакет»)
2. Если файла нет или он устарел (updated_at < applicant.updated_at) — рендерим
   через docx_renderer.py / pdf_forms_engine.builder в памяти
3. Извлекаем текст: docx2txt для .docx, pypdf для .pdf
4. Нормализуем (lstrip/rstrip/collapse whitespace)
5. Возвращаем dict[filename → text_content]

Полный пакет — около 16 файлов:
- 01_Договор.docx
- 02-04_Акт_MM.docx × 3
- 05-07_Счёт_MM.docx × 3
- 08_Письмо.docx
- 09_Резюме.docx
- 10_Выписка.docx
- 11_MI-T.pdf
- 12_Designacion.pdf
- 13_Compromiso.pdf
- 14_Declaracion.pdf
- 15_Справка_НПД.docx
- 15b_Справка_НПД_ЛКН.docx
- 16_Апостиль.docx
- 16_EX-17.pdf (если есть)

Размер выхода: при полном пакете ~30-40k токенов текста, ~120-160KB UTF-8.
"""
import io
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from typing import Dict, Optional, Tuple

from sqlmodel import Session

log = logging.getLogger(__name__)


# ====================================================================
# Конфиг
# ====================================================================

# Максимум символов на документ. Защита от непредвиденно больших файлов
# (например если жёстко зациклится генератор и nasrёт 100 страниц).
MAX_CHARS_PER_DOCUMENT = 20_000

# Сколько секунд считать R2-файл «свежим». Если applicant.updated_at моложе,
# чем этот возраст, считаем что нужно перерендерить.
STALENESS_THRESHOLD_SEC = 60


@dataclass
class ExtractedDocument:
    """Результат извлечения текста из одного файла."""
    filename: str
    source: str  # "r2_cached" | "rendered_fresh" | "missing" | "extract_failed"
    text: str = ""
    char_count: int = 0
    error: Optional[str] = None


@dataclass
class ExtractionResult:
    """Полный результат извлечения по всему пакету."""
    documents: Dict[str, ExtractedDocument] = field(default_factory=dict)
    total_chars: int = 0
    rendered_fresh_count: int = 0
    r2_cached_count: int = 0
    missing_count: int = 0
    failed_count: int = 0

    def to_llm_dict(self) -> Dict[str, str]:
        """Упрощённый dict для передачи в LLM: только filename → text."""
        result = {}
        for fname, doc in self.documents.items():
            if doc.text:
                result[fname] = doc.text
        return result


# ====================================================================
# Нормализация текста
# ====================================================================

_WS_RE = re.compile(r"[ \t\xa0]+")
_EMPTY_LINES_RE = re.compile(r"\n{3,}")


def _normalize_text(text: str) -> str:
    """
    Убираем шум:
    - nbsp → пробел
    - множественные пробелы → один
    - 3+ пустые строки → 2
    - trailing whitespace на каждой строке
    - leading/trailing пустые строки

    LLM лучше читает чистый текст и меньше тратит токенов.
    """
    if not text:
        return ""

    # Сначала trailing whitespace на каждой строке
    lines = [ln.rstrip() for ln in text.split("\n")]
    text = "\n".join(lines)

    # nbsp + множественные пробелы
    text = _WS_RE.sub(" ", text)

    # 3+ переводов строк → 2
    text = _EMPTY_LINES_RE.sub("\n\n", text)

    return text.strip()


def _truncate_if_huge(text: str, filename: str) -> str:
    """Защита от файлов-монстров."""
    if len(text) <= MAX_CHARS_PER_DOCUMENT:
        return text
    log.warning(
        f"[doc_extractor] {filename} text is {len(text)} chars, "
        f"truncating to {MAX_CHARS_PER_DOCUMENT}"
    )
    return text[:MAX_CHARS_PER_DOCUMENT] + "\n\n[... TRUNCATED, see original file ...]"


# ====================================================================
# Извлечение из bytes
# ====================================================================

def _extract_docx_text(file_bytes: bytes) -> str:
    """
    DOCX → plain text через docx2txt (поддерживает таблицы, multi-column,
    что важно для CV-шаблона Pack 20.4 с двумя колонками).
    """
    try:
        import docx2txt
    except ImportError:
        log.error("[doc_extractor] docx2txt not installed; falling back to python-docx")
        return _extract_docx_text_fallback(file_bytes)

    try:
        # docx2txt принимает file-like объект
        text = docx2txt.process(io.BytesIO(file_bytes))
        return text or ""
    except Exception as e:
        log.warning(f"[doc_extractor] docx2txt failed: {e}, trying python-docx")
        return _extract_docx_text_fallback(file_bytes)


def _extract_docx_text_fallback(file_bytes: bytes) -> str:
    """Fallback через python-docx если docx2txt не справился или не установлен."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Параграфы
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # Таблицы
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as e:
        log.error(f"[doc_extractor] python-docx fallback failed: {e}")
        return ""


def _extract_pdf_text(file_bytes: bytes) -> str:
    """
    PDF → plain text через pypdf.

    Для AcroForm PDF после flatten (Pack 36.0) текст уже в content stream,
    извлекается корректно. До flatten — данные были бы в /V полей и
    pypdf.extract_text их не увидел бы.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        log.error("[doc_extractor] pypdf not installed")
        return ""

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"=== PAGE {i+1} ===\n{text}")
            except Exception as e:
                log.warning(f"[doc_extractor] PDF page {i+1} extract failed: {e}")
        return "\n\n".join(parts)
    except Exception as e:
        log.error(f"[doc_extractor] PDF extract failed: {e}")
        return ""


# ====================================================================
# Достать bytes из R2 или сгенерировать заново
# ====================================================================

def _try_fetch_from_r2(
    storage_key: str,
    storage,
) -> Optional[bytes]:
    """
    Пытаемся скачать файл из R2 по ключу. None если файла нет.

    Использует storage.fetch() / storage.get_object() — точное имя метода
    зависит от твоего R2Storage class. Адаптируется ниже под факт.
    """
    if not storage_key:
        return None

    try:
        # Pack 37: проверяем какой метод доступен в storage
        if hasattr(storage, "fetch"):
            data = storage.fetch(storage_key)
        elif hasattr(storage, "download"):
            data = storage.download(storage_key)
        elif hasattr(storage, "get_bytes"):
            data = storage.get_bytes(storage_key)
        elif hasattr(storage, "read"):
            data = storage.read(storage_key)
        else:
            log.error(
                "[doc_extractor] R2 storage has no known fetch method "
                "(fetch/download/get_bytes/read)"
            )
            return None

        if not data:
            return None
        return data if isinstance(data, bytes) else bytes(data)
    except Exception as e:
        # Файла нет в R2 или ошибка сети — не критично, попробуем render
        log.debug(f"[doc_extractor] R2 fetch failed for {storage_key}: {e}")
        return None


def _is_fresh_enough(
    file_modified_at: Optional[datetime],
    applicant_updated_at: Optional[datetime],
) -> bool:
    """
    Файл «свежий» если он был создан/обновлён ПОСЛЕ последнего апдейта applicant.

    Логика: если applicant.updated_at > file_modified_at, значит менеджер
    что-то правил в applicant ПОСЛЕ генерации файла → данные устарели → 
    нужно перерендерить.

    Если хотя бы один из timestamps отсутствует — считаем не свежим.
    """
    if not file_modified_at or not applicant_updated_at:
        return False

    # Допуск 60 секунд: если applicant.updated_at и file_modified_at почти
    # одновременные — считаем свежим (это нормальный кейс render→save).
    delta = (applicant_updated_at - file_modified_at).total_seconds()
    return delta < STALENESS_THRESHOLD_SEC


# ====================================================================
# Главный entry-point
# ====================================================================

def extract_application_documents(
    application_id: int,
    session: Session,
    force_render: bool = False,
) -> ExtractionResult:
    """
    Извлекает текст из всех сгенерированных DOCX/PDF документов заявки.

    Args:
        application_id: ID заявки
        session: SQLModel session
        force_render: если True — игнорируем R2 кеш, всегда рендерим заново.
                      Полезно для тестов и при debugging «почему текст устарел».

    Returns:
        ExtractionResult с словарём документов и статистикой.
    """
    from app.models import Application, GeneratedDocument

    application = session.get(Application, application_id)
    if not application:
        raise ValueError(f"Application {application_id} not found")

    result = ExtractionResult()
    log_prefix = f"[doc_extractor:app#{application_id}]"

    # === 1. Получаем storage клиент ===
    try:
        from app.services.storage.factory import get_storage
        storage = get_storage()
    except Exception as e:
        log.error(f"{log_prefix} Cannot get storage: {e}")
        return result

    # === 2. Собираем DOCX из GeneratedDocument (то что лежит в R2 после
    #        «Сгенерировать пакет») ===
    generated_docs = []
    try:
        from sqlmodel import select
        generated_docs = session.exec(
            select(GeneratedDocument)
            .where(GeneratedDocument.application_id == application_id)
        ).all()
    except Exception as e:
        log.warning(f"{log_prefix} Cannot load GeneratedDocument: {e}")

    log.info(
        f"{log_prefix} Found {len(generated_docs)} generated documents "
        f"in DB (force_render={force_render})"
    )

    # === 3. Извлекаем текст из каждого ===
    applicant_updated_at = (
        application.applicant.updated_at if application.applicant else None
    )

    for gen_doc in generated_docs:
        # Pack 37.0-B.1: реальные имена полей в GeneratedDocument — filename и s3_key
        filename = getattr(gen_doc, "filename", None)
        storage_key = getattr(gen_doc, "s3_key", None)

        if not filename:
            log.warning(f"{log_prefix} GeneratedDocument {gen_doc.id} has no filename")
            continue

        extracted = ExtractedDocument(filename=filename, source="missing")

        # 3a. Пытаемся из R2 (если не force_render и storage_key есть)
        file_bytes: Optional[bytes] = None
        if not force_render and storage_key:
            file_bytes = _try_fetch_from_r2(storage_key, storage)
            if file_bytes:
                # Проверка свежести
                gen_updated_at = getattr(gen_doc, "updated_at", None) or getattr(
                    gen_doc, "created_at", None
                )
                if _is_fresh_enough(gen_updated_at, applicant_updated_at):
                    extracted.source = "r2_cached"
                    result.r2_cached_count += 1
                else:
                    log.info(
                        f"{log_prefix} {filename} is stale "
                        f"(file={gen_updated_at}, applicant={applicant_updated_at}), "
                        f"will re-render"
                    )
                    file_bytes = None  # игнорим устаревший
                    extracted.source = "rendered_fresh"

        # 3b. Если из R2 не вышло — пробуем перерендерить
        # ВНИМАНИЕ: re-render логика зависит от твоего docx_renderer.
        # В части B я оставляю эту ветку как TODO с описанием, в части C
        # реализуем полный re-render через build_full_package(application, session).
        if file_bytes is None:
            log.info(
                f"{log_prefix} {filename}: render-on-the-fly not yet implemented "
                f"(will be added when audit calls build_full_package). "
                f"For now skipping."
            )
            extracted.source = "missing"
            extracted.error = "Not in R2 and re-render not implemented yet"
            result.missing_count += 1
            result.documents[filename] = extracted
            continue

        # 3c. Извлекаем текст по расширению
        lower = filename.lower()
        try:
            if lower.endswith(".docx"):
                raw_text = _extract_docx_text(file_bytes)
            elif lower.endswith(".pdf"):
                raw_text = _extract_pdf_text(file_bytes)
            else:
                log.warning(f"{log_prefix} Unknown extension for {filename}, skip")
                extracted.error = f"Unknown extension: {filename}"
                extracted.source = "extract_failed"
                result.failed_count += 1
                result.documents[filename] = extracted
                continue

            cleaned = _normalize_text(raw_text)
            cleaned = _truncate_if_huge(cleaned, filename)

            extracted.text = cleaned
            extracted.char_count = len(cleaned)
            result.total_chars += len(cleaned)

            log.info(
                f"{log_prefix} {filename}: extracted {len(cleaned)} chars "
                f"(source={extracted.source})"
            )

        except Exception as e:
            log.error(f"{log_prefix} Extract failed for {filename}: {e}", exc_info=True)
            extracted.error = str(e)[:300]
            extracted.source = "extract_failed"
            result.failed_count += 1

        result.documents[filename] = extracted

    log.info(
        f"{log_prefix} Done: total_chars={result.total_chars}, "
        f"r2_cached={result.r2_cached_count}, rendered={result.rendered_fresh_count}, "
        f"missing={result.missing_count}, failed={result.failed_count}"
    )

    return result


# ====================================================================
# Утилита для тестирования
# ====================================================================

def smoke_test_extract(application_id: int) -> None:
    """
    Быстрый smoke-test из консоли:

        python -c "from app.services.audit.document_extractor import smoke_test_extract; smoke_test_extract(10)"
    """
    from app.db.session import engine
    from sqlmodel import Session

    with Session(engine) as session:
        result = extract_application_documents(application_id, session)

    print(f"\n=== Documents extracted for application {application_id} ===")
    print(f"Total chars: {result.total_chars}")
    print(f"From R2 cache: {result.r2_cached_count}")
    print(f"Re-rendered:   {result.rendered_fresh_count}")
    print(f"Missing:       {result.missing_count}")
    print(f"Failed:        {result.failed_count}")
    print()
    for fname, doc in result.documents.items():
        preview = doc.text[:100].replace("\n", " ")
        print(f"  {fname:40s} [{doc.source:18s}] {doc.char_count:6d} chars: {preview}...")
