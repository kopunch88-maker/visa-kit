r"""
Pack 50.23-T — Патч шаблона stdr_template.docx:
  1. Разрыв страницы перед заголовком «...за периоды до 31 декабря 2019 года»
     (таблица до-2019 всегда с новой страницы, как в эталоне Орлова).
  2. Нумерация страниц «Страница X из Y» в footer по центру.

Берёт текущий templates/docx/stdr_template.docx, патчит, сохраняет на место
(с backup). Идемпотентен (проверяет page_break_before и наличие footer-поля).

Запуск (из корня репо или backend):
    python patch_stdr_template_pagebreak.py
"""
from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _find_template() -> Path:
    here = Path(__file__).resolve().parent
    for cand in (
        here / "templates" / "docx" / "stdr_template.docx",
        here.parent / "templates" / "docx" / "stdr_template.docx",
        Path("D:/VISA/visa_kit/templates/docx/stdr_template.docx"),
    ):
        if cand.exists():
            return cand
    raise FileNotFoundError("stdr_template.docx не найден в templates/docx/")


def _add_field(paragraph, instr: str):
    """Добавляет Word-поле (например PAGE / NUMPAGES) в параграф."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    run2 = paragraph.add_run()
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldEnd)


def main() -> int:
    path = _find_template()
    print(f"Шаблон: {path}")

    doc = Document(str(path))

    # ---- 1. Разрыв страницы перед заголовком "до 31 декабря 2019" ----
    changed = False
    p12 = None
    for p in doc.paragraphs:
        if "до 31 декабря 2019" in p.text:
            p12 = p
            break
    if p12 is None:
        raise RuntimeError("Заголовок 'до 31 декабря 2019' не найден в шаблоне")

    if p12.paragraph_format.page_break_before:
        print("   ⏭️  разрыв страницы уже стоит")
    else:
        p12.paragraph_format.page_break_before = True
        changed = True
        print("   ✅ добавлен разрыв страницы перед 'до 31 декабря 2019'")

    # ---- 2. Footer: "Страница X из Y" ----
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    existing = " ".join(p.text for p in footer.paragraphs)
    if "из" in existing and any("PAGE" in (r.text or "") for p in footer.paragraphs for r in p.runs):
        print("   ⏭️  нумерация в footer уже есть")
    else:
        # очистим первый параграф footer и наполним
        fp = footer.paragraphs[0]
        for r in list(fp.runs):
            r._r.getparent().remove(r._r)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run("Страница ")
        _add_field(fp, "PAGE")
        fp.add_run(" из ")
        _add_field(fp, "NUMPAGES")
        changed = True
        print("   ✅ добавлена нумерация 'Страница X из Y' в footer")

    if changed:
        backup = path.with_suffix(
            f".docx.bak_pre_pack50_23T_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        shutil.copy2(path, backup)
        doc.save(str(path))
        print(f"\n✅ Сохранено. Backup: {backup.name}")
    else:
        print("\n(изменений нет — всё уже применено)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
