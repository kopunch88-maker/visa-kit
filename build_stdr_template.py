r"""
Pack 50.9-C — Генератор шаблона stdr_template.docx (СТД-Р).

Создаёт DOCX-шаблон СФР "Сведения о трудовой деятельности" с плейсхолдерами
docxtpl для подстановки данных из build_stdr_context().

Структура:
  1. Шапка (центрированный заголовок "Сведения о трудовой деятельности...")
  2. §1 "Сведения о зарегистрированном лице" (таблица 5×2: ФИО / ДР / СНИЛС)
  3. Подзаголовок "Сведения о трудовой деятельности зарегистрированного лица"
  4. §2 Таблица 1 — события с 2020+ (15 фикс. слотов × 11 колонок)
       С tblLayout="fixed" + cell.width на каждой ячейке (Pack 50.8-C урок).
  5. Подзаголовок "...за периоды до 31 декабря 2019 года"
  6. §3 Таблица 2 — периоды до 2019 (8 фикс. слотов × 4 колонки)
  7. Подпись + дата + М.П.
  8. Блок ЭЦП (фиксированный текст про УКЭП)

Ориентация: альбомная (29.7 × 21.0 cm).

Запуск:
    python build_stdr_template.py
"""

from __future__ import annotations

import os
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


# ============================================================================
# Constants
# ============================================================================

OUTPUT_DIR = Path(__file__).resolve().parent / "templates" / "docx"
OUTPUT_PATH = OUTPUT_DIR / "stdr_template.docx"

# Альбомная: 29.7 × 21.0 cm, поля 1.5 cm с каждой стороны
PAGE_WIDTH_CM = 29.7
PAGE_HEIGHT_CM = 21.0
MARGIN_CM = 1.5
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM  # = 26.7

FONT_FAMILY = "Arial"
FONT_SIZE_PT = 8
FONT_SIZE_SMALL_PT = 7  # шапки таблиц
FONT_SIZE_TITLE_PT = 10
FONT_SIZE_HEADER_PT = 9

STDR_TABLE1_SLOTS = 15
STDR_TABLE2_SLOTS = 8


# ============================================================================
# Helpers
# ============================================================================

def _set_run(run, *, size=FONT_SIZE_PT, bold=False, italic=False, color=None):
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color)
    # Для русских символов нужно явно указать восточно-европейский шрифт
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    rFonts.set(qn("w:cs"), FONT_FAMILY)
    rFonts.set(qn("w:eastAsia"), FONT_FAMILY)


def _add_paragraph(doc, text, *, size=FONT_SIZE_PT, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before_pt=0, space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if text:
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, italic=italic)
    return p


def _set_cell_borders(cell, *, top=True, bottom=True, left=True, right=True,
                      size_val="4"):
    """Управление каждой стороной отдельно. size_val — толщина линии в 1/8 pt."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    sides = [("top", top), ("left", left), ("bottom", bottom), ("right", right)]
    for side_name, enabled in sides:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        border = OxmlElement(f"w:{side_name}")
        if enabled:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size_val)
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "nil")
        tcBorders.append(border)


def _set_table_borders_all(table, size_val="4"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, size_val=size_val)


def _cell_text(cell, text, *, size=FONT_SIZE_PT, bold=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical="center"):
    """Очищает ячейку и пишет текст с настройками шрифта."""
    cell.text = ""  # сбрасываем default-параграф
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[vertical]
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Multi-line text — каждая строка в отдельный run с переносом
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold)


def _set_table_fixed_layout(table):
    """Pack 50.8-C урок: tblLayout w:type='fixed' через XML.

    Иначе Word игнорирует column.width и подгоняет колонки по содержимому.
    """
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _apply_widths(table, widths_cm: list):
    """Pack 50.8-C урок: ставим ширину на колонки И на каждую ячейку явно."""
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def _set_landscape(doc):
    """Альбомная ориентация для всей секции."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Меняем местами ширину и высоту
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)


# ============================================================================
# Builder
# ============================================================================

def build_template() -> Document:
    doc = Document()
    _set_landscape(doc)

    # Шрифт по умолчанию для документа
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_SIZE_PT)

    # ========================================================================
    # 1. ШАПКА — название документа
    # ========================================================================
    _add_paragraph(
        doc,
        "Сведения о трудовой деятельности, предоставляемые из "
        "информационных ресурсов Фонда пенсионного и социального "
        "страхования Российской Федерации",
        size=FONT_SIZE_TITLE_PT,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    # ========================================================================
    # 2. §1 — Сведения о зарегистрированном лице (таблица 5×2)
    # ========================================================================
    _add_paragraph(
        doc, "Сведения о зарегистрированном лице:",
        size=FONT_SIZE_HEADER_PT, bold=True,
        space_before_pt=4, space_after_pt=2,
    )

    t1 = doc.add_table(rows=5, cols=2)
    t1.autofit = False
    _set_table_fixed_layout(t1)
    # Колонка label узкая, значение широкое
    widths_t1 = [4.0, USABLE_WIDTH_CM - 4.0]  # 4.0 + 22.7 = 26.7
    _apply_widths(t1, widths_t1)

    rows_section1 = [
        ("Фамилия", "{{ applicant_stdr.last_name_upper }}"),
        ("Имя", "{{ applicant_stdr.first_name_upper }}"),
        ("Отчество", "{{ applicant_stdr.middle_name_upper }}"),
        ("Дата Рождения", "{{ applicant_stdr.birth_date_long }}"),
        ("СНИЛС", "{{ applicant_stdr.snils }}"),
    ]
    for i, (label, value) in enumerate(rows_section1):
        _cell_text(t1.rows[i].cells[0], label, size=FONT_SIZE_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text(t1.rows[i].cells[1], value, size=FONT_SIZE_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_table_borders_all(t1)

    # ========================================================================
    # 3. Подзаголовок Таблицы 1
    # ========================================================================
    _add_paragraph(
        doc, "Сведения о трудовой деятельности зарегистрированного лица",
        size=FONT_SIZE_HEADER_PT, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=8, space_after_pt=4,
    )

    # ========================================================================
    # 4. §2 — Таблица 1 — события с 2020+ (15 слотов × 11 колонок)
    # ========================================================================
    # Колонки (всего 26.7 cm):
    #   1: №            0.7
    #   2: Работодатель + рег.№ СФР     4.5
    #   3: Дата                          1.7
    #   4: Сведения                      2.0
    #   5: Наименование                  3.5
    #   6: Основание                     1.5
    #   7: Код функции                   1.4
    #   8: Причины увольнения            5.0
    #   9: Док. наименование             1.7
    #  10: Док. дата                     1.7
    #  11: Док. номер                    1.5
    #  12: Признак отмены                1.5
    # Итого: 0.7+4.5+1.7+2.0+3.5+1.5+1.4+5.0+1.7+1.7+1.5+1.5 = 26.7 ✅
    # ВАЖНО: 12 колонок (а не 11 как я сначала писал) — последняя "Признак отмены"

    t2_cols = 12
    t2 = doc.add_table(rows=2 + STDR_TABLE1_SLOTS, cols=t2_cols)
    # Row 0: главные заголовки
    # Row 1: номера колонок (1..12)
    # Row 2..16: 15 слотов данных
    t2.autofit = False
    _set_table_fixed_layout(t2)
    widths_t2 = [0.7, 4.5, 1.7, 2.0, 3.5, 1.5, 1.4, 5.0, 1.7, 1.7, 1.5, 1.5]
    assert abs(sum(widths_t2) - USABLE_WIDTH_CM) < 0.01, f"Sum: {sum(widths_t2)} != {USABLE_WIDTH_CM}"
    _apply_widths(t2, widths_t2)

    # Шапка row 0 — заголовки колонок
    headers_t1 = [
        "№ п/п",
        "Работодатель (наименование),\nрегистрационный номер в СФР",
        "Дата (число, месяц, год) приема, перевода, увольнения",
        "Сведения о приеме, переводе, увольнении",
        "Наименование",
        "Основание",
        "Код выполняемой функции\n(при наличии)",
        "Причины увольнения, пункт, часть статьи, статья ТК РФ, федерального закона",
        "Наименование документа",
        "Дата",
        "Номер документа",
        "Признак отмены записи сведений о приеме, переводе, увольнения",
    ]
    for i, h in enumerate(headers_t1):
        _cell_text(t2.rows[0].cells[i], h,
                   size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1: номера колонок 1..12 (как в эталоне СФР)
    for i in range(t2_cols):
        _cell_text(t2.rows[1].cells[i], str(i + 1),
                   size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Слоты с плейсхолдерами
    for slot in range(STDR_TABLE1_SLOTS):
        row_idx = 2 + slot  # начинаем с row 2 (после 2 строк шапки)
        cells = t2.rows[row_idx].cells
        # Поля: index, company_with_sfr, event_date, event_type, position,
        #       basis, okz_code, dismissal_reason, doc_name, doc_date, doc_number, cancellation
        # 12 ячеек, выравнивания:
        #   №  — CENTER, остальное — LEFT (по эталону СФР)
        slot_fields = [
            ("index", WD_ALIGN_PARAGRAPH.CENTER),
            ("company_with_sfr", WD_ALIGN_PARAGRAPH.LEFT),
            ("event_date", WD_ALIGN_PARAGRAPH.CENTER),
            ("event_type", WD_ALIGN_PARAGRAPH.CENTER),
            ("position", WD_ALIGN_PARAGRAPH.LEFT),
            ("basis", WD_ALIGN_PARAGRAPH.LEFT),
            ("okz_code", WD_ALIGN_PARAGRAPH.CENTER),
            ("dismissal_reason", WD_ALIGN_PARAGRAPH.LEFT),
            ("doc_name", WD_ALIGN_PARAGRAPH.LEFT),
            ("doc_date", WD_ALIGN_PARAGRAPH.CENTER),
            ("doc_number", WD_ALIGN_PARAGRAPH.LEFT),
            ("cancellation", WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for col_idx, (field, align) in enumerate(slot_fields):
            placeholder = f"{{{{ stdr.table1_rows[{slot}].{field} }}}}"
            _cell_text(cells[col_idx], placeholder, align=align)

    _set_table_borders_all(t2)

    # ========================================================================
    # 5. Подзаголовок Таблицы 2
    # ========================================================================
    _add_paragraph(
        doc,
        "Сведения о трудовой деятельности зарегистрированного лица "
        "за периоды до 31 декабря 2019 года",
        size=FONT_SIZE_HEADER_PT, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=12, space_after_pt=4,
    )

    # ========================================================================
    # 6. §3 — Таблица 2 — периоды до 2019 (8 слотов × 4 колонки)
    # ========================================================================
    # Колонки (всего 26.7 cm):
    #   1: №            1.0
    #   2: Работодатель + рег.№        18.0
    #   3: с            3.85
    #   4: по           3.85
    # Итого: 1.0 + 18.0 + 3.85 + 3.85 = 26.7 ✅

    t3 = doc.add_table(rows=2 + STDR_TABLE2_SLOTS, cols=4)
    t3.autofit = False
    _set_table_fixed_layout(t3)
    widths_t3 = [1.0, 18.0, 3.85, 3.85]
    assert abs(sum(widths_t3) - USABLE_WIDTH_CM) < 0.01
    _apply_widths(t3, widths_t3)

    # Row 0: главные заголовки (с merged cells эмуляция через 2 строки)
    headers_t2 = [
        "№ п/п",
        "Работодатель (наименование),\nрегистрационный номер в СФР (при наличии)",
        "Периоды работы",
        "Периоды работы",  # объединим в одну "Периоды работы" над двумя ячейками
    ]
    for i, h in enumerate(headers_t2):
        _cell_text(t3.rows[0].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    # Объединяем "Периоды работы" над двумя последними колонками row 0
    t3.rows[0].cells[2].merge(t3.rows[0].cells[3])

    # Row 1: подзаголовки
    subheaders_t2 = ["", "", "с дд.мм.гггг", "по дд.мм.гггг"]
    for i, h in enumerate(subheaders_t2):
        _cell_text(t3.rows[1].cells[i], h, size=FONT_SIZE_SMALL_PT,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    # Для первых двух колонок row 1 — объединяем по вертикали с row 0
    # (т.е. row 0 cells[0] и cells[1] занимают 2 строки)
    t3.rows[0].cells[0].merge(t3.rows[1].cells[0])
    t3.rows[0].cells[1].merge(t3.rows[1].cells[1])

    # Слоты Таблицы 2
    for slot in range(STDR_TABLE2_SLOTS):
        row_idx = 2 + slot
        cells = t3.rows[row_idx].cells
        slot_fields = [
            ("index", WD_ALIGN_PARAGRAPH.CENTER),
            ("company_with_sfr", WD_ALIGN_PARAGRAPH.LEFT),
            ("date_from", WD_ALIGN_PARAGRAPH.CENTER),
            ("date_to", WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for col_idx, (field, align) in enumerate(slot_fields):
            placeholder = f"{{{{ stdr.table2_rows[{slot}].{field} }}}}"
            _cell_text(cells[col_idx], placeholder, align=align)

    _set_table_borders_all(t3)

    # ========================================================================
    # 7. ПОДПИСЬ + ДАТА
    # ========================================================================
    _add_paragraph(
        doc, "",
        space_before_pt=12,
    )

    # Должность уполномоченного лица (пусто) + Подпись / Расшифровка подписи
    t_sig = doc.add_table(rows=2, cols=3)
    t_sig.autofit = False
    _set_table_fixed_layout(t_sig)
    widths_sig = [14.0, 6.35, 6.35]
    assert abs(sum(widths_sig) - USABLE_WIDTH_CM) < 0.01
    _apply_widths(t_sig, widths_sig)

    _cell_text(t_sig.rows[0].cells[0],
               "Должность уполномоченного лица\nтерриториального органа СФР",
               size=FONT_SIZE_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t_sig.rows[0].cells[1], "Подпись",
               size=FONT_SIZE_PT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t_sig.rows[0].cells[2], "Расшифровка подписи",
               size=FONT_SIZE_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_text(t_sig.rows[1].cells[0], "{{ stdr.issue_date_long }}",
               size=FONT_SIZE_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(t_sig.rows[1].cells[1], "(дата)",
               size=FONT_SIZE_SMALL_PT,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t_sig.rows[1].cells[2], "М.П. (при наличии)",
               size=FONT_SIZE_PT, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Подпись — без границ (это просто текстовая структура)
    for row in t_sig.rows:
        for cell in row.cells:
            _set_cell_borders(cell, top=False, bottom=False, left=False, right=False)

    # ========================================================================
    # 8. БЛОК ЭЦП (фиксированный текст)
    # ========================================================================
    _add_paragraph(
        doc, "",
        space_before_pt=12,
    )

    ep_text = (
        "Документ подписан\n"
        "усиленной квалифицированной электронной подписью.\n"
        "Организация: ФОНД ПЕНСИОННОГО И СОЦИАЛЬНОГО СТРАХОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n"
        "Сертификат: 122b085f13c95875c1a786a93d9b4e6f  "
        "Издатель: Федеральное казначейство  "
        "Действителен: c 02.09.2025 по 26.11.2026"
    )

    # ЭЦП блок — небольшая рамка вокруг
    t_ep = doc.add_table(rows=1, cols=1)
    t_ep.autofit = False
    _set_table_fixed_layout(t_ep)
    _apply_widths(t_ep, [USABLE_WIDTH_CM])
    _cell_text(t_ep.rows[0].cells[0], ep_text,
               size=FONT_SIZE_SMALL_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_table_borders_all(t_ep)

    return doc


def main() -> int:
    print(f"Building STDR template → {OUTPUT_PATH}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Backup существующего шаблона
    if OUTPUT_PATH.exists():
        backup = OUTPUT_PATH.with_suffix(
            f".docx.bak_pre_pack50_9_C_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        shutil.copy2(OUTPUT_PATH, backup)
        print(f"   backup: {backup.name}")

    doc = build_template()
    doc.save(str(OUTPUT_PATH))

    size = OUTPUT_PATH.stat().st_size
    print(f"✅ Saved: {OUTPUT_PATH}  ({size / 1024:.1f} KB)")
    print(f"✅ OK — шаблон создан, размер {size} байт")
    return 0


if __name__ == "__main__":
    sys.exit(main())
